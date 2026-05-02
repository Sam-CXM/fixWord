"""
   ______  ____  __   ____  _             _ _
  / ___\ \/ /  \/  | / ___|| |_ _   _  __| (_) ___
 | |    \  /| |\/| | \___ \| __| | | |/ _` | |/ _ \
 | |___ /  \| |  | |  ___) | |_| |_| | (_| | | (_) |
  \____/_/\_\_|  |_| |____/ \__|\__,_|\__,_|_|\___/

开发作者：晨小明
开发日期：2025/06/23
开发版本：v1.1.3_Dev
发布版本：v1.1.3_Release
修改日期：2026/02/28
主要功能：一、支持批量文档处理，输入文件夹路径，自动判断。
         二、读取.docx文件并设置格式；
         三、支持自定义格式设置：字体、字号、页边距、行距
         三、支持添加页码（可选）：4号半角宋体阿拉伯数字，数字左右各加一条4号“一字线”，奇数页在右侧左空一字，偶数页在左侧左空一字
         四、识别文档中的图片并输出（可选）：（注：图片可能会被压缩）
         五、替换功能
            1.符号替换
                将英文状态下的符号替换为中文状态下的相同符号，包含如下：
                "(" --> "（"
                ")" --> "）"
                "," --> "，"
                ":" --> "："
                ";" --> "；"
                "?" --> "？"
                " " --> ""
            2.其他格式
         六、输出文件名称含时间点，方便标记（可选）
         （注，本程序无法处理图片格式，如果图片独立成段，本程序所用API识别到图片会被默认是空段落，为了防止图片删除，只能放弃处理空段落及图片格式）
更新日志：
【新增】支持三级标题的识别；
【修复】标题识别错误的问题；
【优化】段落与正文函数优化，代码更简洁，提高阅读度；
【优化】复制路径功能；
"""

from docx import Document
from docx.shared import Pt, Cm  # 用来设置字体的大小
from docx.oxml.ns import qn  # 控件名称
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT  # 设置对其方式
from docx.oxml import OxmlElement
from os import listdir, path, makedirs, getcwd, startfile
from tkinter import Tk, Entry, Button, Label, filedialog, messagebox, SUNKEN, Radiobutton, Frame, ttk, Listbox, StringVar, END, Toplevel, Canvas, Menu, LabelFrame, Spinbox
from time import localtime, strftime
from PIL import Image, ImageTk
from webbrowser import open as webopen
from configparser import ConfigParser


def margin(docx):
    """ 设置页边距 """
    global data
    for s in docx.sections:
        s.top_margin = Cm(float(data["margin"]["t_value"]))
        s.bottom_margin = Cm(float(data["margin"]["b_value"]))
        s.left_margin = Cm(float(data["margin"]["l_value"]))
        s.right_margin = Cm(float(data["margin"]["r_value"]))


def footer(docx):
    """ 设置页脚，添加页码 """
    # print(len(docx.sections))
    def AddFooterNumber(p):
        t1 = p.add_run("— ")
        font = t1.font
        font.name = '宋体'
        font.size = Pt(14)  # 14号字体
        t1._element.rPr.rFonts.set(qn("w:eastAsia"), '宋体')

        run1 = p.add_run('')
        fldChar1 = OxmlElement('w:fldChar')  # creates a new element
        fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        run1._element.append(fldChar1)

        run2 = p.add_run('')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'PAGE'
        font = run2.font
        font.name = '宋体'
        font.size = Pt(14)  # 14号字体
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), '宋体')
        run2._element.append(instrText)

        run3 = p.add_run('')
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._element.append(fldChar2)

        t2 = p.add_run(" —")
        font = t2.font
        font.name = '宋体'
        font.size = Pt(14)  # 14号字体
        t2._element.rPr.rFonts.set(qn("w:eastAsia"), '宋体')

    for s in docx.sections:
        # print(s.footer)
        footer = s.footer  # 获取第一个节的页脚
        footer.is_linked_to_previous = True  # 编号续前一节
        paragraph = footer.paragraphs[0]  # 获取页脚的第一个段落
        paragraphFun("odd_footer", paragraph)
        AddFooterNumber(paragraph)
        even_footer = s.even_page_footer  # 获取第一个节的页脚
        even_footer.is_linked_to_previous = True  # 编号续前一节
        paragraph = even_footer.paragraphs[0]  # 获取页脚的第一个段落
        paragraphFun("even_footer", paragraph)
        AddFooterNumber(paragraph)


def isLevel1(p):
    """ 判断是否是 1 级标题 """
    index1_list = ["一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、",
                   "十一、", "十二、", "十三、", "十四、", "十五、", "十六、", "十七、", "十八、", "十九、", "二十、"]
    for i in index1_list:
        if i in p.text[:3]:
            if '。' in p.text:
                p.text = p.text.replace('。', '')
            if '？' in p.text:
                p.text = p.text.replace('？', '')
            if '：' in p.text:
                p.text = p.text.replace('：', '')
            if '；' in p.text:
                p.text = p.text.replace('；', '')
            return "level1"
        else:
            continue


def isLevel2(p):
    """ 判断是否是 2 级标题 """
    index2_list = ["（一）", "（二）", "（三）", "（四）", "（五）", "（六）", "（七）", "（八）", "（九）", "（十）",
                   "（十一）", "（十二）", "（十三）", "（十四）", "（十五）", "（十六）", "（十七）", "（十八）", "（十九）", "（二十）"]
    for i in index2_list:
        if i in p.text[:4]:
            if '。' in p.text:
                p.text = p.text.replace('。', '')
            if '？' in p.text:
                p.text = p.text.replace('？', '')
            if '：' in p.text:
                p.text = p.text.replace('：', '')
            if '；' in p.text:
                p.text = p.text.replace('；', '')
            return "level2"
        else:
            continue


def isLevel3(p):
    """ 判断是否是 3 级标题 """
    index3_list = ["1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.", "11.", "12.", "13.", "14.", "15.", "16.", "17.", "18.", "19.", "20."]
    for i in index3_list:
        if i in p.text[:3]:
            return "level3"
        else:
            continue


def paragraphFun(is_title, p, is_level1="", is_level2="", is_level3=""):
    """ 段落函数 """
    global data
    try:
        p.paragraph_format.element.pPr.ind.set(qn("w:leftChars"), '0')  # 左侧缩进
        p.paragraph_format.element.pPr.ind.set(qn("w:rightChars"), '0')  # 右侧缩进
        p.paragraph_format.element.pPr.ind.set(qn("w:left"), '0')  # 缩进(cm)
        p.paragraph_format.element.pPr.ind.set(qn("w:right"), '0')  # 缩进(cm)
    except Exception as e:
        pass
        # print('错误：', e, f'\t{p.text} 不支持设置缩进')
    pgp_almt = data["main"]["pgp_almt"]
    if pgp_almt == "居中":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif pgp_almt == "左对齐":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif pgp_almt == "右对齐":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if is_title == "title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = Pt(float(data["title_font"]["font_ls_vlu"]))  # 行距
        p.paragraph_format.first_line_indent = None
    if is_title == "odd_footer":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.paragraph_format.right_indent = Pt(14)
        p.paragraph_format.line_spacing = Pt(28)
    elif is_title == "even_footer":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.line_spacing = Pt(28)
    else:
        p.paragraph_format.line_spacing = Pt(float(data["mb_font"]["font_mb_ls_value"]))  # 行距
        if is_level1 == "level1":
            p.paragraph_format.line_spacing = Pt(float(data["1title_font"]["font_ls_vlu"]))  # 行距
            p.paragraph_format.space_before = Pt(float(data["spacing"]["b_value"]))
            p.paragraph_format.space_after = Pt(float(data["spacing"]["a_value"]))
        if is_level2 == "level2":
            p.paragraph_format.line_spacing = Pt(float(data["2title_font"]["font_ls_vlu"]))  # 行距
            p.paragraph_format.space_before = Pt(float(data["spacing"]["b_value"]))
            p.paragraph_format.space_after = Pt(float(data["spacing"]["a_value"]))
        if is_level3 == "level3":
            p.paragraph_format.line_spacing = Pt(float(data["3title_font"]["font_ls_vlu"]))  # 行距
        p.paragraph_format.space_before = Pt(float(data["spacing"]["b_value"]))
        p.paragraph_format.space_after = Pt(float(data["spacing"]["a_value"]))
    p.paragraph_format.left_indent = Cm(float(data["indent"]["l_value"]))
    p.paragraph_format.right_indent = Cm(float(data["indent"]["r_value"]))
    first_line_idt = float(data["indent"]["f_value"])
    if first_line_idt > 0:
        p.paragraph_format.first_line_indent = 0
        p.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), f'{first_line_idt * 100}')  # 首行缩进
    else:
        p.paragraph_format.first_line_indent = 0
        p.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '0')  # 首行缩进
    single_crl_value = data["main"]["single_crl_value"]
    if single_crl_value == "1":
        p.paragraph_format.widow_control = True
    else:
        p.paragraph_format.widow_control = False


def text(is_title, is_level1, is_level2, is_level3, is_digit, p, i):
    """ 正文函数 """
    global data
    if is_title == "title":
        run = p.add_run(i)
        if is_digit == "num_or_let":
            run.font.name = data["num_font"]["font_num_name"]
        else:
            run.font.name = data["title_font"]["font_title_name"]
            run._element.rPr.rFonts.set(qn('w:eastAsia'), data["title_font"]["font_title_name"])
        run.font.size = Pt(FONTSIZEDICT[data["title_font"]["font_title_size"]])
    else:
        run_content = p.add_run(i)
        if is_digit == "num_or_let":
            run_content.font.name = data["num_font"]["font_num_name"]
            run_content.font.size = Pt(FONTSIZEDICT[data["num_font"]["font_num_size"]])
        else:
            run_content.font.name = data["mb_font"]["font_mb_name"]
            run_content._element.rPr.rFonts.set(qn('w:eastAsia'), data["mb_font"]["font_mb_name"])
            run_content.font.size = Pt(FONTSIZEDICT[data["mb_font"]["font_mb_size"]])
            if is_level1 == "level1":
                run_content.font.name = data["1title_font"]["font_title_name"]
                run_content._element.rPr.rFonts.set(qn('w:eastAsia'), data["1title_font"]["font_title_name"])
                run_content.font.size = Pt(FONTSIZEDICT[data["1title_font"]["font_title_size"]])
            if is_level2 == "level2":
                run_content.font.name = data["2title_font"]["font_title_name"]
                run_content._element.rPr.rFonts.set(qn('w:eastAsia'), data["2title_font"]["font_title_name"])
                run_content.font.size = Pt(FONTSIZEDICT[data["2title_font"]["font_title_size"]])
            if is_level3 == "level3":
                run_content.font.name = data["3title_font"]["font_title_name"]
                run_content._element.rPr.rFonts.set(qn('w:eastAsia'), data["3title_font"]["font_title_name"])
                run_content.font.size = Pt(FONTSIZEDICT[data["3title_font"]["font_title_size"]])


def replace(p):
    """ 替换函数 """
    # 替换符号
    if '(' in p.text:
        p.text = p.text.replace('(', '（')
    if ')' in p.text:
        p.text = p.text.replace(')', '）')
    if ',' in p.text:
        p.text = p.text.replace(',', '，')
    if ':' in p.text:
        p.text = p.text.replace(':', '：')
    if ';' in p.text:
        p.text = p.text.replace(';', '；')
    if '?' in p.text:
        p.text = p.text.replace('?', '？')
    if '》、' in p.text:
        p.text = p.text.replace('》、', '》')
    if '．' in p.text:  # U+ff0e
        p.text = p.text.replace('．', '.')
    if ' ' in p.text:  # 空格
        p.text = p.text.replace(' ', '')
    if '　' in p.text:  # U+3000
        p.text = p.text.replace('　', '')
    if ' ' in p.text:  # U+2003
        p.text = p.text.replace(' ', '')
    return p


def fixDocx(docx):
    """ 主要格式 """
    lvl = 0
    for p in docx.paragraphs:
        if p.text == '':
            continue
        else:
            lvl += 1
            p = replace(p)
            is_level1 = isLevel1(p)
            is_level2 = isLevel2(p)
            is_level3 = isLevel3(p)
            if lvl == 1:
                paragraphFun("title", p, is_level1, is_level2, is_level3)
                for run_title in p.runs:
                    # print(run_title.text)
                    string = run_title.text
                    # print(string)
                    run_title._element.getparent().remove(run_title._element)
                    for i in string:
                        num_or_let = isNumberOrLetter(i)
                        text("title", is_level1, is_level2, is_level3, num_or_let, p, i)
            else:
                paragraphFun("text", p, is_level1, is_level2, is_level3)
                for run_content in p.runs:
                    # print(run_content.text)
                    string = run_content.text
                    # print(string)
                    run_content._element.getparent().remove(run_content._element)
                    for i in string:  # 遍历字符串
                        num_or_let = isNumberOrLetter(i)
                        text("notitle", is_level1, is_level2, is_level3, num_or_let, p, i)


def isNumberOrLetter(char):
    """ 判断是否为数字或字母 """
    number_and_letter_strs = '0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ'
    if char in number_and_letter_strs:
        return "num_or_let"
    else:
        return False


def picFix(docx, file, output_path):
    """ 图片处理 """
    img_path = output_path + "\image"
    file_name = path.splitext(file)[0]
    parts = docx.part.related_parts
    parts_values = parts.values()
    parts_keys = parts.keys()
    list_val = list(parts_values)
    list_key = list(parts_keys)
    parts_length = len(parts_values)
    if parts_length > 5:
        # print(type(list(parts_values)[-1]))
        k = 0
        for i in range(parts_length):
            # print(type(list_val[i]))
            if 'image' in str(type(list_val[i])):
                if not path.isdir(img_path):
                    makedirs(img_path)
                # print('找到图片数据')
                k += 1
                try:
                    img_data = parts[list_key[i]].image.blob
                    img_type = parts[list_key[i]].image.ext
                    full_path = f'{img_path}\{file_name}_image{k}.{img_type}'
                    print(f"··>提示<·· 正在输出：{full_path}")
                    with open(full_path, 'wb') as f:
                        f.write(img_data)
                except:
                    print(f"··>错误<·· 图片{k}输出失败！")
        if k == 0:
            print(f"··>提示<·· 未找到图片！")


def fixWord(docx_path, file, output_path, time_ipt, page_ipt, img_ipt):
    """ 文档处理 """
    docx = Document(docx_path)

    # 页边距
    margin(docx)

    # 修改格式
    fixDocx(docx)

    # 添加时间后缀
    file_name = path.splitext(file)[0]
    if time_ipt == "1":
        save_time = strftime("%m%d%H%M", localtime())
        save_path = output_path + f"\{file_name}" + save_time + ".docx"
    else:
        save_time = ""
        save_path = output_path + f"\{file_name}" + ".docx"

    # 设置页码
    if page_ipt == "1":
        # 奇偶页不同
        docx.settings.odd_and_even_pages_header_footer = True
        footer(docx)

    # 保存文档中的图片
    if img_ipt == "1":
        picFix(docx, file, output_path)
    try:
        docx.save(save_path)
        writeHistory(save_path)
    except PermissionError:
        writeHistory(f"{file_name} 保存失败！文件已打开，请关闭后重试！")
        messagebox.showerror("错误", f"{file_name} 保存失败！文件已打开，请关闭后重试！")
    return save_time


def writeHistory(text=""):
    """ 写入历史记录 """
    output_time = strftime("%m-%d %H:%M:%S", localtime())
    output_txt = output_time + "    " + text
    play_history_frm_listbox.insert(END, output_txt)
    play_history_frm_listbox.update()
    print(f"··>提示<·· {output_txt}")
    # 设置滚动条位置到最大值，即拖动到最底部
    play_history_frm_listbox.yview_moveto(1)
    # play_history_frm_listbox.xview_moveto(1)


def inputDir():
    """ 输入路径 """
    dir_path = filedialog.askdirectory(title="请选择文件夹")
    if dir_path != "":
        path_entry.delete(0, END)
        path_entry.insert(0, dir_path)


class InitFile():
    """配置文件"""

    def __init__(self):
        global data
        self.config_path = path.join(path.dirname(__file__), "fixWord_D_config.ini")
        self.config = ConfigParser()
        data = getUserInput()
        # print(data)
        # 确保配置文件存在
        if not path.exists(self.config_path):
            self.config['DEFAULT'] = {}
            with open(self.config_path, 'w') as configfile:
                self.config.write(configfile)

    def saveConfig(self):
        """保存配置"""
        # 检查是否有配置文件
        if path.exists(self.config_path):
            isgoon = messagebox.askyesno("提示", "配置文件已存在，点击【是】会替换原有配置，是否继续？")
            if not isgoon:
                writeHistory("取消保存配置！")
                return

        # 解析数据
        for section, options in data.items():
            if not self.config.has_section(section):
                self.config.add_section(section)
            if isinstance(options, dict):
                for key, value in options.items():
                    self.config.set(section, key, str(value))
            else:
                self.config.set("DEFAULT", section, str(options))

        # 保存配置
        with open(self.config_path, "w", encoding="utf-8") as f:
            self.config.write(f)
        writeHistory("保存配置成功！")

    def importConfig(self):
        """导入配置"""
        # 检查是否有配置文件
        if not path.exists(self.config_path):
            writeHistory("错误，配置文件不存在！")
            messagebox.showerror("错误", "配置文件不存在！")
            return

        # 读取配置
        self.config.read(self.config_path, encoding="utf-8")
        # 获取配置
        for section in self.config.sections():
            if section not in data:
                data[section] = {}
            for key, value in self.config.items(section):
                data[section][key] = value
        # print(data)
        font_title_name_frm_combox.set(data['title_font']['font_title_name']), font_title_size_frm_combox.set(data['title_font']['font_title_size']), font_title_ls_frm_combox.set(data['title_font']['font_title_ls']), font_ls_vlu.set(data['title_font']['font_ls_vlu']), font_ls_frm_lbl_b.config(text=f"{data['title_font']['font_ls_lbl_txt']}")
        font_title_name_frm1_combox.set(data['1title_font']['font_title_name']), font_title_size_frm1_combox.set(data['1title_font']['font_title_size']), font_title_ls_frm1_combox.set(data['1title_font']['font_title_ls']), font_ls_vlu1.set(data['1title_font']['font_ls_vlu']), font_ls_frm_lbl_b1.config(text=f"{data['1title_font']['font_ls_lbl_txt']}")
        font_title_name_frm2_combox.set(data['2title_font']['font_title_name']), font_title_size_frm2_combox.set(data['2title_font']['font_title_size']), font_title_ls_frm2_combox.set(data['2title_font']['font_title_ls']), font_ls_vlu2.set(data['2title_font']['font_ls_vlu']), font_ls_frm_lbl_b2.config(text=f"{data['2title_font']['font_ls_lbl_txt']}")
        font_title_name_frm3_combox.set(data['3title_font']['font_title_name']), font_title_size_frm3_combox.set(data['3title_font']['font_title_size']), font_title_ls_frm3_combox.set(data['3title_font']['font_title_ls']), font_ls_vlu3.set(data['3title_font']['font_ls_vlu']), font_ls_frm_lbl_b3.config(text=f"{data['3title_font']['font_ls_lbl_txt']}")
        font_mb_name_frm_combox.set(data['mb_font']['font_mb_name']), font_mb_size_frm_combox.set(data['mb_font']['font_mb_size']), font_mb_ls_frm_combox.set(data['mb_font']['font_mb_ls']), font_mb_ls_vlu.set(data['mb_font']['font_mb_ls_value']), font_mb_ls_frm_lbl_b.config(text=f"{data['mb_font']['font_mb_ls_txt']}")
        font_num_name_frm_combox.set(data['num_font']['font_num_name']), font_num_size_frm_combox.set(data['num_font']['font_num_size']), font_num_ls_frm_combox.set(data['num_font']['font_num_ls']), font_num_ls_vlu.set(data['num_font']['font_num_ls_value']), font_num_ls_frm_lbl_b.config(text=f"{data['num_font']['font_num_ls_txt']}")
        pgp_indent_l_vlu.set(data['indent']['l_value']), pgp_indent_r_vlu.set(data['indent']['r_value']), pgp_indent_f_vlu.set(data['indent']['f_value'])
        spacing_b_ent.set(data['spacing']['b_value']), spacing_a_ent.set(data['spacing']['a_value']), spacing_l_ent.set(data['spacing']['l_value'])
        pgp_almt_frm_combox.set(data['main']['pgp_almt'])
        single_crl_radio_value.set(data['main']['single_crl_value'])
        time_radio_value.set(data['main']['time_ipt'])
        page_radio_value.set(data['main']['page_ipt'])
        img_radio_value.set(data['main']['img_ipt'])
        pgp_margin_t_vlu.set(data['margin']['t_value']), pgp_margin_b_vlu.set(data['margin']['b_value']), pgp_margin_l_vlu.set(data['margin']['l_value']), pgp_margin_r_vlu.set(data['margin']['r_value'])
        writeHistory("配置导入成功！")


def reSet():
    """ 重置 """
    global data
    font_title_name_frm_combox.current(0), font_title_size_frm_combox.current(0), font_title_ls_frm_combox.current(0), font_ls_vlu.set("1"), font_ls_frm_lbl_b.config(text="倍"), font_title_ls_ent.config(state="disabled")
    font_title_name_frm1_combox.current(0), font_title_size_frm1_combox.current(0), font_title_ls_frm1_combox.current(0), font_ls_vlu1.set("1"), font_ls_frm_lbl_b1.config(text="倍"), font_title_ls_ent1.config(state="disabled")
    font_title_name_frm2_combox.current(0), font_title_size_frm2_combox.current(0), font_title_ls_frm2_combox.current(0), font_ls_vlu2.set("1"), font_ls_frm_lbl_b2.config(text="倍"), font_title_ls_ent2.config(state="disabled")
    font_title_name_frm3_combox.current(0), font_title_size_frm3_combox.current(0), font_title_ls_frm3_combox.current(0), font_ls_vlu3.set("1"), font_ls_frm_lbl_b3.config(text="倍"), font_title_ls_ent3.config(state="disabled")
    font_mb_name_frm_combox.current(0), font_mb_size_frm_combox.current(0), font_mb_ls_frm_combox.current(0), font_mb_ls_vlu.set("1"), font_mb_ls_frm_lbl_b.config(text="倍"), font_mb_ls_ent.config(state="disabled")
    font_num_name_frm_combox.current(0), font_num_size_frm_combox.current(0), font_num_ls_frm_combox.current(0), font_num_ls_vlu.set("1"), font_num_ls_frm_lbl_b.config(text="倍"), font_num_ls_ent.config(state="disabled")
    if font_title_ls_frm_combox.cget("state").string == "disabled":
        font_title_ls_frm_combox.config(state="normal")
        font_title_ls_frm1_combox.config(state="normal")
        font_title_ls_frm2_combox.config(state="normal")
        font_title_ls_frm3_combox.config(state="normal")
        font_mb_ls_frm_combox.config(state="normal")
        font_num_ls_frm_combox.config(state="normal")
    pgp_indent_l_vlu.set("0"), pgp_indent_r_vlu.set("0"), pgp_indent_f_vlu.set("0")
    spacing_b_ent.set("0"), spacing_a_ent.set("0"), spacing_l_ent.set("0")
    pgp_almt_frm_combox.current(3)
    pgp_margin_t_vlu.set("0"), pgp_margin_b_vlu.set("0"), pgp_margin_l_vlu.set("0"), pgp_margin_r_vlu.set("0")
    single_crl_radio2.select()
    time_radio2.select()
    page_radio2.select()
    img_radio2.select()
    play_history_frm_listbox.delete(0, END)
    writeHistory("重置成功！")


def on_enter(event):
    event.widget.config(cursor="hand2", fg="blue")


def on_leave(event):
    event.widget.config(cursor="", fg="#888888")


def toMail(event):
    """ 打开邮箱 """
    webopen("mailto:3038693133@qq.com")


def wxTk(event):
    wx_tk = Toplevel(tk)
    original_image = Image.open(wxgzh_path)
    wx_tk.geometry(f"{original_image.width}x460+0+0")
    wx_tk.iconbitmap(icon_path)
    wx_tk.title("微信公众号：晨小明工作室")
    wx_title = Label(wx_tk, text="微信扫一扫关注公众号", font=("微软雅黑", 14))
    wx_title.grid(row=0, column=0, padx=2, pady=2)
    # 创建Canvas
    cv = Canvas(wx_tk, width=original_image.width, height=original_image.height + 30, highlightthickness=0)
    cv.grid(row=1, column=0, padx=2, pady=0)
    # 加载图片
    time_icon = original_image.resize((round(original_image.width / 1), round(original_image.height / 1)))  # 缩放图片到指定大小
    time_icon_new = ImageTk.PhotoImage(time_icon)
    cv.create_image(0, 0, image=time_icon_new, anchor="nw")
    wx_tk.mainloop()


class MyFrame():
    def __init__(self, frm, title_txt, row, col, last_txt):
        self.frm = frm
        self.title_txt = title_txt
        self.row = row
        self.col = col
        self.last_txt = last_txt

    def cFontFrame(self):
        """ 字体标题 """
        font_label = Label(self.frm, font=("Ya Hei", 10, "bold"), text=self.title_txt)
        font_label.grid(row=self.row, column=0, padx=2, pady=5, sticky="e")
        font_name_frm = Frame(self.frm)  # 字体选择下拉框
        font_name_frm.grid(row=self.row, column=1, sticky="n")
        font_name_frm_lbl = Label(font_name_frm, text="字体：", font=("Ya Hei", 10, "bold"))  # 字体文本
        font_name_frm_lbl.grid(row=self.row, column=0, padx=2, pady=5)
        font_name_frm_combox = ttk.Combobox(font_name_frm, width=14, font=("Ya Hei", 10))  # 字体下拉框盒子
        font_name_frm_combox.grid(row=self.row, column=1, padx=2, pady=5)
        font_name_frm_combox['values'] = ("方正小标宋简体",  "黑体", "楷体_GB2312", "仿宋_GB2312", "Times New Roman", "宋体", "微软雅黑")
        font_name_frm_combox.current(0)
        font_size_frm = Frame(self.frm)  # 字号选择下拉框
        font_size_frm.grid(row=self.row, column=2, sticky="n")
        font_size_frm_lbl = Label(font_size_frm, text="字号：", font=("Ya Hei", 10, "bold"))  # 字号文本
        font_size_frm_lbl.grid(row=self.row, column=0, padx=2, pady=5)
        font_size_frm_combox = ttk.Combobox(font_size_frm, width=4, font=("Ya Hei", 10))  # 字号下拉框盒子
        font_size_frm_combox.grid(row=self.row, column=1, padx=2, pady=5)
        font_size_frm_combox['values'] = ("初号", "小初", "一号", "小一", "二号", "小二", "三号", "小三", "四号", "小四", "五号", "小五", "六号", "小六", "七号", "八号")
        font_size_frm_combox.current(0)
        font_ls_frm = Frame(self.frm)  # 行距选择下拉框
        font_ls_frm.grid(row=self.row, column=3, sticky="n")
        font_ls_frm_lbl = Label(font_ls_frm, text="行距：", font=("Ya Hei", 10, "bold"))  # 行距文本
        font_ls_frm_lbl.grid(row=self.row, column=0, padx=2, pady=5)
        font_ls_frm_vlu = StringVar()  # 创建一个StringVar变量来存储数值
        font_ls_frm_combox = ttk.Combobox(font_ls_frm, width=6, font=("Ya Hei", 10), textvariable=font_ls_frm_vlu)  # 行距下拉框盒子
        font_ls_frm_combox.grid(row=self.row, column=1, padx=2, pady=5)
        font_ls_frm_combox['values'] = ("单倍", "1.5倍", "2倍", "最小值", "固定值", "多倍")
        font_ls_frm_combox.current(0)
        font_ls_vlu = StringVar()  # 创建一个StringVar变量来存储数值
        font_ls_vlu.set("1")  # 初始值设置为0
        font_ls_ent = Entry(font_ls_frm, width=4, font=("Ya Hei", 10), relief="solid", textvariable=font_ls_vlu, state="disabled")  # 输入框
        font_ls_ent.grid(row=self.row, column=2, padx=2, pady=5)
        font_ls_frm_lbl_b = Label(font_ls_frm, text="倍", font=("Ya Hei", 10, "bold"))  # 磅文本
        font_ls_frm_lbl_b.grid(row=self.row, column=3, padx=2, pady=5)
        return font_name_frm_combox, font_size_frm_combox, font_ls_frm_combox, font_ls_ent, font_ls_frm_lbl_b, font_ls_frm_vlu, font_ls_vlu

    def cIndentSpacingFrame(self):
        """ 缩进 """
        pgp_indent_frm_lbl = Label(self.frm, text=self.title_txt, font=("Ya Hei", 10, "bold"))  # 左侧缩进文本
        pgp_indent_frm_lbl.grid(row=self.row, column=0, padx=2, pady=5, sticky="e")
        pgp_indent_vlu = StringVar()  # 创建一个StringVar变量来存储数值
        pgp_indent_vlu.set("0")  # 初始值设置为0
        pgp_indent_spinbox = Spinbox(self.frm, from_=0, to=100, increment=0.1, textvariable=pgp_indent_vlu, width=5, font=("Ya Hei", 10), wrap=True)
        pgp_indent_spinbox.grid(row=self.row, column=1, padx=2, pady=5)
        pgp_indent_lbl_b = Label(self.frm, text=self.last_txt, font=("Ya Hei", 10, "bold"))  # 磅文本
        pgp_indent_lbl_b.grid(row=self.row, column=2, padx=2, pady=5)
        return pgp_indent_spinbox, pgp_indent_vlu

    def cRadioFrame(self):
        """ 是否 """
        crl_label = Label(self.frm, font=("Ya Hei", 10, "bold"), text=self.title_txt)
        crl_label.grid(row=self.row, column=0, padx=5, pady=5, sticky="e")
        crl_radio_value = StringVar()
        crl_radio1 = Radiobutton(self.frm, font=("Ya Hei", 10), text="是", variable=crl_radio_value, value=True)
        crl_radio1.grid(row=self.row, column=1, padx=5, pady=5)
        crl_radio2 = Radiobutton(self.frm, font=("Ya Hei", 10), text="否", variable=crl_radio_value, value=False)
        crl_radio2.grid(row=self.row, column=2, padx=5, pady=5)
        return crl_radio_value, crl_radio1, crl_radio2

    def cMarginFrame(self):
        """ 页边距 """
        if self.col > 0:
            self.col = self.col + self.col * 2
        pgp_margin_frm_lbl = Label(self.frm, text=self.title_txt, font=("Ya Hei", 10, "bold"))  # 左侧缩进文本
        pgp_margin_frm_lbl.grid(row=self.row, column=self.col, padx=2, pady=2, sticky="e")
        pgp_margin_vlu = StringVar()  # 创建一个IntVar变量来存储数值
        pgp_margin_vlu.set("0")
        pgp_margin_spinbox = Spinbox(self.frm, from_=0, to=100, increment=0.25, textvariable=pgp_margin_vlu, width=5, font=("Ya Hei", 10), wrap=True)
        pgp_margin_spinbox.grid(row=0, column=self.col + 1, padx=2, pady=2)
        pgp_margin_lbl_b = Label(self.frm, text="cm", font=("Ya Hei", 10, "bold"))  # 磅文本
        pgp_margin_lbl_b.grid(row=0, column=self.col + 2, padx=(2, 20), pady=2)
        return pgp_margin_spinbox, pgp_margin_vlu


def fontTitleLsFrmCombox(ls_ent, ls_frm_lbl_b, ls_frm_vlu):
    """ 标题行距选择事件 """
    # print(ls_frm_vlu.get())
    if ls_frm_vlu.get() == "单倍":
        ls_ent.config(state='normal')
        ls_ent.delete(0, END)  # 删除所有文本
        ls_ent.insert(0, "1")  # 设置新的默认值
        ls_ent.config(state='disabled')
        ls_frm_lbl_b.config(text="倍")
    elif ls_frm_vlu.get() == "1.5倍":
        ls_ent.config(state='normal')
        ls_ent.delete(0, END)  # 删除所有文本
        ls_ent.insert(0, "1.5")  # 设置新的默认值
        ls_ent.config(state='disabled')
        ls_frm_lbl_b.config(text="倍")
    elif ls_frm_vlu.get() == "2倍":
        ls_ent.config(state='normal')
        ls_ent.delete(0, END)  # 删除所有文本
        ls_ent.insert(0, "2")  # 设置新的默认值
        ls_ent.config(state='disabled')
        ls_frm_lbl_b.config(text="倍")
    elif ls_frm_vlu.get() == "最小值":
        ls_ent.config(state='normal')
        ls_ent.delete(0, END)  # 删除所有文本
        ls_ent.insert(0, "12")  # 设置新的默认值
        ls_ent.config(state='normal')
        ls_frm_lbl_b.config(text="磅")
    elif ls_frm_vlu.get() == "固定值":
        ls_ent.config(state='normal')
        ls_ent.delete(0, END)  # 删除所有文本
        ls_ent.insert(0, "12")  # 设置新的默认值
        ls_ent.config(state='normal')
        ls_frm_lbl_b.config(text="磅")
    elif ls_frm_vlu.get() == "多倍":
        ls_ent.config(state='normal')
        ls_ent.delete(0, END)  # 删除所有文本
        ls_ent.insert(0, "1")  # 设置新的默认值
        ls_frm_lbl_b.config(text="倍")
    else:
        ls_ent.config(state='disabled')
        ls_frm_lbl_b.config(text="磅")


def spacingLSpb(spacing_l_ent):
    """ 全文行距统一事件 """
    global data
    l_ent_vlu = float(spacing_l_ent.get())
    if l_ent_vlu > 0:
        font_title_ls_frm_combox.config(state="disabled"), font_title_ls_ent.config(state="disabled")
        font_title_ls_frm1_combox.config(state="disabled"), font_title_ls_ent1.config(state="disabled")
        font_title_ls_frm2_combox.config(state="disabled"), font_title_ls_ent2.config(state="disabled")
        font_title_ls_frm3_combox.config(state="disabled"), font_title_ls_ent3.config(state="disabled")
        font_mb_ls_frm_combox.config(state="disabled"), font_mb_ls_ent.config(state="disabled")
        font_num_ls_frm_combox.config(state="disabled"), font_num_ls_ent.config(state="disabled")
        data["title_font"]["font_ls_vlu"] = data["1title_font"]["font_ls_vlu"] = data["2title_font"]["font_ls_vlu"] = data["3title_font"]["font_ls_vlu"] = data["mb_font"]["font_mb_ls_value"] = data["num_font"]["font_num_ls_value"] = l_ent_vlu
        data["title_font"]["font_ls_lbl_txt"] = data["1title_font"]["font_ls_lbl_txt"] = data["2title_font"]["font_ls_lbl_txt"] = data["3title_font"]["font_ls_lbl_txt"] = data["mb_font"]["font_mb_ls_txt"] = data["num_font"]["font_num_ls_txt"] = "磅"
    else:
        font_title_ls_frm_combox.config(state="normal"), font_title_ls_ent.config(state="normal")
        font_title_ls_frm1_combox.config(state="normal"), font_title_ls_ent1.config(state="normal")
        font_title_ls_frm2_combox.config(state="normal"), font_title_ls_ent2.config(state="normal")
        font_title_ls_frm3_combox.config(state="normal"), font_title_ls_ent3.config(state="normal")
        font_mb_ls_frm_combox.config(state="normal"), font_mb_ls_ent.config(state="normal")
        font_num_ls_frm_combox.config(state="normal"), font_num_ls_ent.config(state="normal")
        data = getUserInput()


def create_popup_menu(event):
    """ 创建右键菜单 """
    # 获取当前选中的条目索引和内容
    selected = play_history_frm_listbox.curselection()
    if selected:
        # 创建一个菜单
        popup_menu = Menu(play_history_frm, tearoff=0, font=("Ya Hei", 10))
        # 添加菜单项
        popup_menu.add_command(label="打开文件", command=lambda: open_folder(1))
        popup_menu.add_command(label="复制路径", command=lambda: copy_selected(play_history_frm_listbox))
        popup_menu.add_command(label="在文件夹中显示", command=lambda: open_folder(2))
        popup_menu.add_separator()  # 添加分隔线
        popup_menu.add_command(label="退出", command=tk.quit)  # 添加退出命令（可选）
        # 显示菜单
        popup_menu.tk_popup(event.x_root, event.y_root)


def open_folder(type):
    # 获取当前选中的条目索引和内容
    selected_index = play_history_frm_listbox.curselection()[0]  # 获取当前选中项的索引
    selected_folder = play_history_frm_listbox.get(selected_index)  # 获取当前选中项的内容（即文件夹路径）
    selected_ = selected_folder.split("    ")[1]
    if path.exists(selected_):
        if type == 1:  # 打开文件
            # 使用系统默认的文件浏览器打开文件夹
            startfile(selected_)  # Windows系统使用此方法
        elif type == 2:    # 打开文件夹
            folder_path = path.dirname(selected_)
            # 使用系统默认的文件浏览器打开文件夹
            startfile(folder_path)  # Windows系统使用此方法
    else:
        messagebox.showwarning("提示", "请指向正确路径！")


def copy_selected(listbox):
    # 获取选中的项
    # 获取当前选中的条目索引和内容
    selected_index = listbox.curselection()[0]  # 获取当前选中项的索引
    selected_folder = listbox.get(selected_index)  # 获取当前选中项的内容（即文件夹路径）
    selected_ = selected_folder.split("    ")[1]
    if path.exists(selected_):
        # 这里可以添加复制到剪贴板的代码，例如使用tkinter的clipboard模块
        if tk.clipboard_get():
            tk.clipboard_clear()  # 清空剪贴板
        tk.clipboard_append(selected_)
        messagebox.showinfo("提示", "已复制到剪贴板！\n使用 【Ctrl+V】 粘贴即可！")
    else:
        messagebox.showwarning("提示", "未检测到有效路径！请重试！")


def done():
    """ 处理完成 """
    merge_button.config(state="normal", cursor="", text="开始处理")
    reset_button.config(state="normal")


def getUserInput():
    """ 获取用户输入 """
    global data
    input_path = path_entry.get().replace("/", "\\")
    output_path = input_path + "\output"
    # 获取数值
    font_title_name, font_title_size, font_title_ls, font_ls_value, font_ls_lbl_txt = font_title_name_frm_combox.get(), font_title_size_frm_combox.get(), font_title_ls_frm_combox.get(), font_ls_vlu.get(), font_ls_frm_lbl_b.cget("text")
    font_title_name1, font_title_size1, font_title_ls1, font_ls_value1, font_ls_lbl_txt1 = font_title_name_frm1_combox.get(), font_title_size_frm1_combox.get(), font_title_ls_frm1_combox.get(), font_ls_vlu1.get(), font_ls_frm_lbl_b1.cget("text")
    font_title_name2, font_title_size2, font_title_ls2, font_ls_value2, font_ls_lbl_txt2 = font_title_name_frm2_combox.get(), font_title_size_frm2_combox.get(), font_title_ls_frm2_combox.get(), font_ls_vlu2.get(), font_ls_frm_lbl_b2.cget("text")
    font_title_name3, font_title_size3, font_title_ls3, font_ls_value3, font_ls_lbl_txt3 = font_title_name_frm3_combox.get(), font_title_size_frm3_combox.get(), font_title_ls_frm3_combox.get(), font_ls_vlu3.get(), font_ls_frm_lbl_b3.cget("text")
    font_mb_name, font_mb_size, font_mb_ls, font_mb_ls_value, font_mb_ls_txt = font_mb_name_frm_combox.get(), font_mb_size_frm_combox.get(), font_mb_ls_frm_combox.get(), font_mb_ls_vlu.get(), font_mb_ls_frm_lbl_b.cget("text")
    font_num_name, font_num_size, font_num_ls, font_num_ls_value, font_num_ls_txt = font_num_name_frm_combox.get(), font_num_size_frm_combox.get(), font_num_ls_frm_combox.get(), font_num_ls_vlu.get(), font_num_ls_frm_lbl_b.cget("text")
    pgp_indent_l_value, pgp_indent_r_value, pgp_indent_f_value = pgp_indent_l_vlu.get(), pgp_indent_r_vlu.get(), pgp_indent_f_vlu.get()
    spacing_b_vlu, spacing_a_vlu, spacing_l_vlu = spacing_b_ent.get(), spacing_a_ent.get(), spacing_l_ent.get()
    pgp_almt = pgp_almt_frm_combox.get()
    single_crl_value = single_crl_radio_value.get()
    pgp_margin_t_value, pgp_margin_b_value, pgp_margin_l_value, pgp_margin_r_value = pgp_margin_t_vlu.get(), pgp_margin_b_vlu.get(), pgp_margin_l_vlu.get(), pgp_margin_r_vlu.get()
    time_ipt = time_radio_value.get()
    page_ipt = page_radio_value.get()
    img_ipt = img_radio_value.get()
    data = {
        "title_font": {
            "font_title_name": font_title_name,
            "font_title_size": font_title_size,
            "font_title_ls": font_title_ls,
            "font_ls_vlu": font_ls_value,
            "font_ls_lbl_txt": font_ls_lbl_txt
        },
        "1title_font": {
            "font_title_name": font_title_name1,
            "font_title_size": font_title_size1,
            "font_title_ls": font_title_ls1,
            "font_ls_vlu": font_ls_value1,
            "font_ls_lbl_txt": font_ls_lbl_txt1
        },
        "2title_font": {
            "font_title_name": font_title_name2,
            "font_title_size": font_title_size2,
            "font_title_ls": font_title_ls2,
            "font_ls_vlu": font_ls_value2,
            "font_ls_lbl_txt": font_ls_lbl_txt2
        },
        "3title_font": {
            "font_title_name": font_title_name3,
            "font_title_size": font_title_size3,
            "font_title_ls": font_title_ls3,
            "font_ls_vlu": font_ls_value3,
            "font_ls_lbl_txt": font_ls_lbl_txt3
        },
        "mb_font": {
            "font_mb_name": font_mb_name,
            "font_mb_size": font_mb_size,
            "font_mb_ls": font_mb_ls,
            "font_mb_ls_value": font_mb_ls_value,
            "font_mb_ls_txt": font_mb_ls_txt
        },
        "num_font": {
            "font_num_name": font_num_name,
            "font_num_size": font_num_size,
            "font_num_ls": font_num_ls,
            "font_num_ls_value": font_num_ls_value,
            "font_num_ls_txt": font_num_ls_txt
        },
        "indent": {
            "l_value": pgp_indent_l_value,
            "r_value": pgp_indent_r_value,
            "f_value": pgp_indent_f_value,
        },
        "spacing": {
            "b_value": spacing_b_vlu,
            "a_value": spacing_a_vlu,
            "l_value": spacing_l_vlu,
        },
        "margin": {
            "t_value": pgp_margin_t_value,
            "b_value": pgp_margin_b_value,
            "l_value": pgp_margin_l_value,
            "r_value": pgp_margin_r_value,
        },
        "main": {
            "pgp_almt": pgp_almt,
            "single_crl_value": single_crl_value,
            "time_ipt": time_ipt,
            "page_ipt": page_ipt,
            "img_ipt": img_ipt,
            "input_path": input_path,
            "output_path": output_path}
    }
    # 判断是否需要统一行距
    l_ent_vlu = float(spacing_l_ent.get())
    if l_ent_vlu > 0:
        data["title_font"]["font_ls_vlu"] = data["1title_font"]["font_ls_vlu"] = data["2title_font"]["font_ls_vlu"] = data["3title_font"]["font_ls_vlu"] = data["mb_font"]["font_mb_ls_value"] = data["num_font"]["font_num_ls_value"] = l_ent_vlu
        data["title_font"]["font_ls_lbl_txt"] = data["1title_font"]["font_ls_lbl_txt"] = data["2title_font"]["font_ls_lbl_txt"] = data["3title_font"]["font_ls_lbl_txt"] = data["mb_font"]["font_mb_ls_txt"] = data["num_font"]["font_num_ls_txt"] = "磅"
    # print(data)
    return data


def main():
    """ 主函数 """
    global data
    input_path = path_entry.get().replace("/", "\\")
    data = getUserInput()
    if input_path == "":
        messagebox.showerror("错误", "请选择文件或文件夹路径！")
    else:
        if not path.isdir(input_path):
            messagebox.showerror("错误", "文件夹路径错误！")
        else:
            print("··>提示<·· 检查路径成功！")
            output_path = data["main"]["output_path"]
            time_ipt = data["main"]["time_ipt"]
            page_ipt = data["main"]["page_ipt"]
            img_ipt = data["main"]["img_ipt"]
            merge_button.config(state="disabled", cursor="wait", text="正在处理")
            reset_button.config(state="disabled")
            print(data)
            have_docx = 0
            for file in listdir(input_path):
                if '~' in file:
                    continue
                elif file.endswith('.docx'):
                    if not path.isdir(output_path):
                        makedirs(output_path)
                    have_docx += 1
                    file_path = path.join(input_path, file)
                    fixWord(file_path, file, output_path, time_ipt, page_ipt, img_ipt)
            if have_docx == 0:
                print("··>错误<·· 没有找到.docx文件")
                messagebox.showwarning("警告", "没有找到.docx文件！")
            else:
                messagebox.showinfo("提示", "全部处理完成！\n输出路径：" + output_path)
            done()


if __name__ == '__main__':
    VERSION = "D_v1.1.3"
    UPDATETIME = "2026年2月28日"
    """
        !!!!!!!!!!!!
        打包时把此路径改为相对路径，并把图片复制粘贴到打包后的根目录里
        !!!!!!!!!!!!
    """
    icon_path = getcwd() + "\\static\\icon.ico"
    wxgzh_path = getcwd() + "\\static\\wxgzh.jpg"
    cxm_path = getcwd() + "\\static\\cxmstudio-lignt-heng.png"
    # 配置信息start
    # 字号字典
    FONTSIZEDICT = {"八号": 5, "七号": 5.5, "小六": 6.5, "六号": 7.5, "小五": 9, "五号": 10.5, "小四": 12, "四号": 14, "小三号": 15, "三号": 16, "小二": 18, "二号": 22, "小一": 24, "一号": 26, "小初": 36, "初号": 42}
    # 配置信息end
    # tkinter start
    tk = Tk()
    tk.title(f"文档处理工具 {VERSION}")
    screen_width = tk.winfo_screenwidth()
    screen_height = tk.winfo_screenheight()
    tk.iconbitmap(icon_path)
    tk.geometry("1400x820")
    tk.minsize(1265, 805)  # 最小宽高
    # 刷新窗口参数
    tk.update()
    # 计算窗口居中时左上角的坐标
    x = (screen_width - tk.winfo_width()) // 2
    y = (screen_height - tk.winfo_height()) // 2
    tk.geometry(f"+{x}+{y-50}")
    # tk.attributes("-alpha", 0.8)
    windw_width = tk.winfo_width()
    windw_height = tk.winfo_height()
    frm_ = Frame(tk)
    frm_.pack(anchor="center")
    # 文件路径
    path_lf = LabelFrame(frm_, text="选择路径", padx=10, pady=10, width=100)
    path_lf.grid(row=0, column=0, padx=5, pady=5)
    path_entry = Entry(path_lf, width=80, font=("Ya Hei", 10), relief="solid")  # 输入框
    path_entry.grid(row=0, column=0, padx=5, pady=5, ipadx=5, ipady=4, sticky="w")
    path_button = Button(path_lf, font=("Ya Hei", 10), text="选择文件夹", command=inputDir)
    path_button.grid(row=0, column=1, padx=5, pady=5)
    # 保存配置
    ini_frm = Frame(frm_)
    ini_frm.grid(row=0, column=1, padx=5, pady=5)
    separator = Frame(tk, height=2, bd=1, relief=SUNKEN)  # 分隔线
    separator.pack(fill="x", padx=5, pady=5)
    # 主布局
    main_lf = Frame(tk)
    main_lf.pack(padx=5, pady=5)
    main_frm = Frame(main_lf)
    main_frm.grid(row=0, column=0, padx=5, pady=5)
    # 字体
    font_lf = LabelFrame(main_frm, text="设置字体", padx=10, pady=10, width=100)
    font_lf.grid(row=0, column=0, padx=5, pady=5)
    font_title_lf = LabelFrame(font_lf, text="标题", padx=10, pady=10, width=100)  # 标题
    font_title_lf.pack(padx=10, pady=10)
    font_title_name_frm_combox, font_title_size_frm_combox, font_title_ls_frm_combox, font_title_ls_ent, font_ls_frm_lbl_b, font_ls_frm_vlu, font_ls_vlu = MyFrame(font_title_lf, "标题",  0, 0, "").cFontFrame()
    font_title_ls_frm_combox.bind("<<ComboboxSelected>>", lambda event: fontTitleLsFrmCombox(font_title_ls_ent, font_ls_frm_lbl_b, font_ls_frm_vlu))
    font_title_name_frm1_combox, font_title_size_frm1_combox, font_title_ls_frm1_combox, font_title_ls_ent1, font_ls_frm_lbl_b1, font_ls_frm_vlu1, font_ls_vlu1 = MyFrame(font_title_lf, "一级标题", 1, 0, "").cFontFrame()
    font_title_ls_frm1_combox.bind("<<ComboboxSelected>>", lambda event: fontTitleLsFrmCombox(font_title_ls_ent1, font_ls_frm_lbl_b1, font_ls_frm_vlu1))
    font_title_name_frm2_combox, font_title_size_frm2_combox, font_title_ls_frm2_combox, font_title_ls_ent2, font_ls_frm_lbl_b2, font_ls_frm_vlu2, font_ls_vlu2 = MyFrame(font_title_lf, "二级标题", 2, 0, "").cFontFrame()
    font_title_ls_frm2_combox.bind("<<ComboboxSelected>>", lambda event: fontTitleLsFrmCombox(font_title_ls_ent2, font_ls_frm_lbl_b2, font_ls_frm_vlu2))
    font_title_name_frm3_combox, font_title_size_frm3_combox, font_title_ls_frm3_combox, font_title_ls_ent3, font_ls_frm_lbl_b3, font_ls_frm_vlu3, font_ls_vlu3 = MyFrame(font_title_lf, "三级标题", 3, 0, "").cFontFrame()
    font_title_ls_frm3_combox.bind("<<ComboboxSelected>>", lambda event: fontTitleLsFrmCombox(font_title_ls_ent3, font_ls_frm_lbl_b3, font_ls_frm_vlu3))
    # 正文
    font_mb_lf = LabelFrame(font_lf, text="正文", padx=10, pady=10, width=150)
    font_mb_lf.pack(padx=10, pady=10)
    font_mb_frm = Frame(font_mb_lf)
    font_mb_frm.grid(row=0, column=0, padx=5, pady=5)
    font_mb_name_frm_combox, font_mb_size_frm_combox, font_mb_ls_frm_combox, font_mb_ls_ent, font_mb_ls_frm_lbl_b, font_mb_ls_frm_vlu, font_mb_ls_vlu = MyFrame(font_mb_frm, "   正文", 0, 0, "").cFontFrame()  # 正文
    font_mb_ls_frm_combox.bind("<<ComboboxSelected>>", lambda event: fontTitleLsFrmCombox(font_mb_ls_ent, font_mb_ls_frm_lbl_b, font_mb_ls_frm_vlu))
    # 其他
    font_else_lf = LabelFrame(font_lf, text="其他", padx=10, pady=10, width=150)
    font_else_lf.pack(padx=10, pady=10)
    font_num_frm = Frame(font_else_lf)
    font_num_frm.grid(row=0, column=0, padx=5, pady=5)
    font_num_name_frm_combox, font_num_size_frm_combox, font_num_ls_frm_combox, font_num_ls_ent, font_num_ls_frm_lbl_b, font_num_ls_frm_vlu, font_num_ls_vlu = MyFrame(font_num_frm, "数字英文", 0, 0, "").cFontFrame()  # 数字英文
    font_num_ls_frm_combox.bind("<<ComboboxSelected>>", lambda event: fontTitleLsFrmCombox(font_num_ls_ent, font_num_ls_frm_lbl_b, font_num_ls_frm_vlu))
    # 段落设置
    # 对齐
    pgp_lf = LabelFrame(main_frm, text="设置段落", padx=10, pady=10, width=100)
    pgp_lf.grid(row=0, column=1, padx=5, pady=5)
    frm_0 = Frame(pgp_lf)
    frm_0.grid(row=0, column=0, padx=5, pady=5)
    # 缩进
    pgp_indent_lf = LabelFrame(frm_0, text="设置缩进", padx=10, pady=10, width=100)  # 设置缩进
    pgp_indent_lf.grid(row=0, column=0, padx=5, pady=5)
    pgp_indent_frm = Frame(pgp_indent_lf)
    pgp_indent_frm.grid(row=0, column=0, padx=5, pady=4)
    pgp_indent_l_spb, pgp_indent_l_vlu = MyFrame(pgp_indent_frm, "左侧缩进：", 0, 0, "厘米").cIndentSpacingFrame()
    pgp_indent_r_spb, pgp_indent_r_vlu = MyFrame(pgp_indent_frm, "右侧缩进：", 1, 0, "厘米").cIndentSpacingFrame()
    pgp_indent_f_spb, pgp_indent_f_vlu = MyFrame(pgp_indent_frm, "首行缩进：", 2, 0, "字符").cIndentSpacingFrame()
    # 间距
    spacing_lf = LabelFrame(frm_0, text="设置间距", padx=10, pady=10, width=100)  # 设置间距
    spacing_lf.grid(row=1, column=0, padx=5, pady=5)
    spacing_frm = Frame(spacing_lf)
    spacing_frm.grid(row=0, column=0, padx=5, pady=4)
    spacing_b_spb, spacing_b_ent = MyFrame(spacing_frm, "    段前：", 0, 0, "磅 ").cIndentSpacingFrame()
    spacing_a_spb, spacing_a_ent = MyFrame(spacing_frm, "    段后：", 1, 0, "磅 ").cIndentSpacingFrame()
    spacing_l_spb, spacing_l_ent = MyFrame(spacing_frm, "    行距：", 2, 0, "磅 ").cIndentSpacingFrame()
    spacing_l_spb.config(command=lambda: spacingLSpb(spacing_l_ent))
    frm_1 = Frame(pgp_lf)
    frm_1.grid(row=0, column=1, padx=5, pady=5)
    # 对齐方式
    pgp_almt_lf = LabelFrame(frm_1, text="设置对齐方式", padx=10, pady=10)
    pgp_almt_lf.grid(row=0, column=0, padx=5, pady=5)
    pgp_almt_frm = Frame(pgp_almt_lf)
    pgp_almt_frm.grid(row=0, column=0, padx=28, pady=5)
    pgp_almt_frm_lbl = Label(pgp_almt_frm, text="对齐方式：", font=("Ya Hei", 10, "bold"))  # 对齐方式文本
    pgp_almt_frm_lbl.grid(row=0, column=0, padx=2, pady=2)
    pgp_almt_frm_combox = ttk.Combobox(pgp_almt_frm, width=8, font=("Ya Hei", 10))  # 对齐方式下拉框盒子
    pgp_almt_frm_combox.grid(row=0, column=1, padx=2, pady=2)
    pgp_almt_frm_combox['values'] = ("左对齐", "居中", "右对齐", "两端对齐")
    pgp_almt_frm_combox.current(3)
    # 孤行控制
    spacing_sc = LabelFrame(frm_1, text="设置孤行控制", padx=10, pady=10)  # 设置孤行控制
    spacing_sc.grid(row=1, column=0, padx=5, pady=5)
    single_crl_frm = Frame(spacing_sc)
    single_crl_frm.grid(row=1, column=0, padx=16, pady=5)
    single_crl_radio_value, single_crl_radio1, single_crl_radio2 = MyFrame(single_crl_frm, "孤行控制：", 0, 0, "").cRadioFrame()
    single_crl_radio2.select()
    # 设置页边距
    frm_2 = Frame(pgp_lf)
    frm_2.grid(row=1, column=0, padx=5, pady=2, columnspan=2)
    pgp_margin_lf = LabelFrame(frm_2, text="设置页边距", padx=10, pady=10)  # 设置页边距
    pgp_margin_lf.grid(row=0, column=0, padx=5, pady=2)
    pgp_margin_frm = Frame(pgp_margin_lf)
    pgp_margin_frm.grid(row=0, column=0, padx=5, pady=2)
    pgp_margin_t_spb, pgp_margin_t_vlu = MyFrame(pgp_margin_frm, "上：", 0, 0, "").cMarginFrame()
    pgp_margin_b_spb, pgp_margin_b_vlu = MyFrame(pgp_margin_frm, "下：", 0, 1, "").cMarginFrame()
    pgp_margin_l_spb, pgp_margin_l_vlu = MyFrame(pgp_margin_frm, "左：", 0, 2, "").cMarginFrame()
    pgp_margin_r_spb, pgp_margin_r_vlu = MyFrame(pgp_margin_frm, "右：", 0, 3, "").cMarginFrame()
    # 处理信息
    infos_frm = Frame(frm_1)
    infos_frm.grid(row=2, column=0, padx=5, pady=5)
    info_frm = Frame(infos_frm)
    info_frm.grid(row=0, column=0, padx=5, pady=5)
    time_radio_value, time_radio1, time_radio2 = MyFrame(info_frm, "添加时间标记：", 0, 0, "").cRadioFrame()
    page_radio_value, page_radio1, page_radio2 = MyFrame(info_frm, "添加页码：", 1, 0, "").cRadioFrame()
    img_radio_value, img_radio1, img_radio2 = MyFrame(info_frm, "保存文档中的图片：", 2, 0, "").cRadioFrame()
    time_radio2.select()
    page_radio2.select()
    img_radio2.select()
    # 分隔线
    separator = Frame(tk, height=2, bd=1, relief=SUNKEN)
    separator.pack(fill="x", padx=5, pady=5)
    # 处理按钮
    btn_frm = Frame(tk)
    btn_frm.pack(pady=6)
    merge_button = Button(btn_frm, font=("Ya Hei", 10, "bold"), text="开始处理", command=main)
    merge_button.grid(row=0, column=0, padx=5, pady=5)
    label_ = Label(btn_frm, font=("Ya Hei", 10), text=" ")
    label_.grid(row=0, column=1, padx=5, pady=5, sticky="e")
    reset_button = Button(btn_frm, font=("Ya Hei", 10), text="重置", fg="blue", command=reSet)
    reset_button.grid(row=0, column=2, padx=5, pady=5)
    # 分隔线
    separator = Frame(tk, height=2, bd=1, relief=SUNKEN)
    separator.pack(fill="x", padx=2, pady=2)
    # 处理日志
    play_history_frm = Frame(tk)
    play_history_frm.pack()
    play_history_frm_lbl = Label(play_history_frm, text="操 作 日 志", font=("Ya Hei", 10, "bold"), fg="green")
    play_history_frm_lbl.grid(row=0, column=0, padx=5, pady=(10, 5))
    play_history_frm_listbox = Listbox(play_history_frm, width=100, height=6, font=("Ya Hei", 10), border=1, activestyle="none")
    play_history_frm_listbox.grid(row=1, column=0, padx=5, pady=5)
    play_history_scroll_bar_v = ttk.Scrollbar(play_history_frm, orient="vertical", command=play_history_frm_listbox.yview)
    play_history_scroll_bar_v.grid(row=1, column=1, sticky='ns')
    play_history_scroll_bar_h = ttk.Scrollbar(play_history_frm, orient="horizontal", command=play_history_frm_listbox.xview)
    play_history_scroll_bar_h.grid(row=2, column=0, sticky='we')
    play_history_frm_listbox.configure(yscrollcommand=play_history_scroll_bar_v.set, xscrollcommand=play_history_scroll_bar_h.set)
    # 绑定右键点击事件到创建弹出菜单的函数
    play_history_frm_listbox.bind("<Button-3>", create_popup_menu)
    # 绑定双击事件到列表框上
    play_history_frm_listbox.bind("<Double-1>", lambda: open_folder(1))
    ini_save_button = Button(ini_frm, font=("Ya Hei", 10), text="保存配置", command=lambda: InitFile().saveConfig())
    ini_save_button.grid(row=0, column=0, padx=5, pady=5)
    ini_import_button = Button(ini_frm, font=("Ya Hei", 10), text="导入配置", command=lambda: InitFile().importConfig())
    ini_import_button.grid(row=1, column=0, padx=5, pady=5)
    # 底部信息
    # 底部文字
    bottom_frm = Frame(tk)
    bottom_frm.pack(side="bottom")
    # 晨小明工作室
    cxm_frm = Frame(bottom_frm)
    cxm_frm.pack()
    original_image = Image.open(cxm_path)
    resized_image = original_image.resize((round(original_image.width / 21), round(original_image.height / 21)))  # 缩放图片到指定大小
    cxm_image_new = ImageTk.PhotoImage(resized_image)
    cv_cxm = Canvas(cxm_frm, width=110, height=cxm_image_new.height(), highlightthickness=0)
    cv_cxm.create_image(5, 0, image=cxm_image_new, anchor="nw")
    cv_cxm.grid(row=0, column=0)
    bottom_info_frm = Frame(bottom_frm)
    bottom_info_frm.pack()
    bottom_label_a = Label(bottom_info_frm, font=("Ya Hei", 10), fg="#888888", text="作者：晨小明")
    bottom_label_a.grid(row=0, column=0, padx=5, pady=5)
    bottom_label_v = Label(bottom_info_frm, font=("Ya Hei", 10), fg="#888888", text=f"版本：{VERSION}")
    bottom_label_v.grid(row=0, column=1, padx=5, pady=5)
    bottom_label_t = Label(bottom_info_frm, font=("Ya Hei", 10), fg="#888888", text=F"更新时间：{UPDATETIME}")
    bottom_label_t.grid(row=0, column=2, padx=5, pady=5)
    bottom_label_w = Label(bottom_info_frm, font=("Ya Hei", 10), fg="#888888", text="微信公众号：晨小明工作室（CXM-Studio）")
    bottom_label_w.grid(row=0, column=3, padx=5, pady=5)
    bottom_label_w.bind("<Enter>", on_enter)
    bottom_label_w.bind("<Leave>", on_leave)
    bottom_label_w.bind("<Button-1>", wxTk)
    bottom_label_c = Label(bottom_info_frm, font=("Ya Hei", 10), fg="#888888", text="联系作者：3038693133@qq.com")
    bottom_label_c.grid(row=0, column=4, padx=5, pady=5)
    bottom_label_c.bind("<Enter>", on_enter)
    bottom_label_c.bind("<Leave>", on_leave)
    bottom_label_c.bind("<Button-1>", toMail)
    data = getUserInput()
    writeHistory("初始化成功！")
    tk.mainloop()
    # tkinter end
