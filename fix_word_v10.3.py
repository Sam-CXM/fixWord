from docx import Document
from docx.shared import Pt, Cm  # 用来设置字体的大小
from docx.oxml.ns import qn  # 设置字体
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT  # 设置对其方式
from os import listdir, path, makedirs
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from time import localtime, strftime, sleep

"""
开发作者：晨小明
开发日期：2024/01/04
开发版本：v1.3__release
修改日期：2024/08/20
主要功能：一、支持单文件处理或批量文档处理，输入文件路径或文件夹路径，自动判断。
         二、读取.docx文件并设置格式：
            1.页边距：上3.7cm，下3.5cm，左2.8cm，右2.6cm
            2.段落行距：标题：固定值33磅；正文：一般固定值28磅
            3.字体，字号：标题：小二号方正小标宋简体，居中；一级标题：四号黑体；二级标题：四号楷体_GB2312；正文：四号仿宋_GB2312，两端对齐；数字&英文：四号TimesNewRoman字体
            4.支持添加页码（可选）：4号半角宋体阿拉伯数字，数字左右各加一条4号“一字线”，奇数页在右侧左空一字，偶数页在左侧左空一字
            5.识别文档中的图片并输出（可选）：（注：图片可能会被压缩）
         三、替换功能
            1.符号替换
                将英文状态下的符号替换为中文状态下的相同符号，包含如下：
                "(" --> "（"
                ")" --> "）"
                ")、" --> "）"
                "）、" --> "）"
                "," --> "，"
                ":" --> "："
                ";" --> "；"
                "?" --> "？"
                " " --> ""
            2.其他格式
                数字后有顿号替换为点，如："1、" --> "1."
         四、输出文件名称含时间点，方便标记（可选）
         （注，本程序无法处理图片格式，如果图片独立成段，本程序所用API识别到图片会被默认是空段落，为了防止图片删除，只能放弃处理空段落及图片格式）
更新日志：
【优化】解决了首行缩进2字符的问题
【优化】设置基础信息常量，如页边距、行距、字体、字号等，方便后续修改
"""


def margin(docx):
    """ 设置页边距 """
    for s in docx.sections:
        s.top_margin = Cm(PAGETOPMARGIN)
        s.bottom_margin = Cm(PAGEBOTTOMMARGIN)
        s.left_margin = Cm(PAGELEFTMARGIN)
        s.right_margin = Cm(PAGERIGHTMARGIN)


def footer(docx):
    """ 设置页脚，添加页码 """
    # print(len(docx.sections))
    def AddFooterNumber(p):
        t1 = p.add_run("— ")
        font = t1.font
        font.name = PAGENUMBERFONT
        font.size = Pt(PAGENUMBERFONTSIZE)  # 14号字体
        t1._element.rPr.rFonts.set(qn("w:eastAsia"), PAGENUMBERFONT)

        run1 = p.add_run('')
        fldChar1 = OxmlElement('w:fldChar')  # creates a new element
        fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        run1._element.append(fldChar1)

        run2 = p.add_run('')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'PAGE'
        font = run2.font
        font.name = PAGENUMBERFONT
        font.size = Pt(PAGENUMBERFONTSIZE)  # 14号字体
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), PAGENUMBERFONT)
        run2._element.append(instrText)

        run3 = p.add_run('')
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._element.append(fldChar2)

        t2 = p.add_run(" —")
        font = t2.font
        font.name = PAGENUMBERFONT
        font.size = Pt(PAGENUMBERFONTSIZE)  # 14号字体
        t2._element.rPr.rFonts.set(qn("w:eastAsia"), PAGENUMBERFONT)

    for s in docx.sections:
        # print(s.footer)
        footer = s.footer  # 获取第一个节的页脚
        footer.is_linked_to_previous = True  # 编号续前一节
        paragraph = footer.paragraphs[0]  # 获取页脚的第一个段落
        paragraphFun("odd_footer", paragraph)
        AddFooterNumber(paragraph)
        even_footer = s.even_page_footer  # 获取第一个节的页脚
        even_footer.is_linked_to_previous = True  # 编号续前一节
        paragraph = even_footer.paragraphs[0]  # 获取页脚的第一个段落
        paragraphFun("even_footer", paragraph)
        AddFooterNumber(paragraph)


def paragraphFun(is_title, p):
    """ 段落函数 """
    if is_title == "title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = Pt(TITLEMARGIN)  # 行距
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.before_spacing = Pt(0)
        p.paragraph_format.after_spacing = Pt(0)
    elif is_title == "odd_footer":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.paragraph_format.right_indent = Pt(14)
        p.paragraph_format.line_spacing = Pt(TEXTMARGIN)
    elif is_title == "even_footer":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.line_spacing = Pt(TEXTMARGIN)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = Pt(TEXTMARGIN)  # 行距
        p.paragraph_format.before_spacing = Pt(0)
        p.paragraph_format.after_spacing = Pt(0)


def isLevel1(p):
    """ 判断是否是 1 级标题 """
    index1_list = ["一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、",
                   "十一、", "十二、", "十三、", "十四、", "十五、", "十六、", "十七、", "十八、", "十九、", "二十、"]
    for i in range(len(index1_list)):
        if p.text.find(index1_list[i]) != -1:
            if '。' in p.text:
                p.text = p.text.replace('。', '')
            if '？' in p.text:
                p.text = p.text.replace('？', '')
            if '：' in p.text:
                p.text = p.text.replace('：', '')
            if '；' in p.text:
                p.text = p.text.replace('；', '')
            return "level1"
        else:
            continue


def isLevel2(p):
    """ 判断是否是 2 级标题 """
    index2_list = ["（一）", "（二）", "（三）", "（四）", "（五）", "（六）", "（七）", "（八）", "（九）", "（十）",
                   "（十一）", "（十二）", "（十三）", "（十四）", "（十五）", "（十六）", "（十七）", "（十八）", "（十九）", "（二十）"]
    for i in range(len(index2_list)):
        if p.text.find(index2_list[i]) != -1:
            if '。' in p.text:
                p.text = p.text.replace('。', '')
            if '？' in p.text:
                p.text = p.text.replace('？', '')
            if '：' in p.text:
                p.text = p.text.replace('：', '')
            if '；' in p.text:
                p.text = p.text.replace('；', '')
            return "level2"
        else:
            continue


def text(is_title, is_level1, is_level2, is_digit, p, i):
    """ 正文函数 """
    if is_title == "title":
        if is_digit == "num_or_let":
            new_run = p.add_run(i)
            new_run.font.name = NUMBERFONT
            new_run.font.size = Pt(TITLEFONTSIZE)
            new_run.font.bold = False
        else:
            run_title = p.add_run(i)
            run_title.font.name = TITLEFONT
            run_title._element.rPr.rFonts.set(
                qn('w:eastAsia'), TITLEFONT)
            run_title.font.size = Pt(TITLEFONTSIZE)
            run_title.font.bold = False
    else:
        if is_digit == "num_or_let":
            new_run = p.add_run(i)
            new_run.font.name = NUMBERFONT
            new_run.font.size = Pt(TITLEFONTSIZE)
            new_run.font.bold = False
        elif is_level1 == "level1":
            run_level1 = p.add_run(i)
            run_level1.font.name = LEVEL1FONT
            run_level1._element.rPr.rFonts.set(
                qn('w:eastAsia'), LEVEL1FONT)
            run_level1.font.size = Pt(TEXTFONTSIZE)
            run_level1.font.bold = False
        elif is_level2 == "level2":
            run_level2 = p.add_run(i)
            run_level2.font.name = LEVEL2FONT
            run_level2._element.rPr.rFonts.set(
                qn('w:eastAsia'), LEVEL2FONT)
            run_level2.font.size = Pt(TEXTFONTSIZE)
            run_level2.font.bold = False
        else:
            run_content = p.add_run(i)
            run_content.font.name = TEXTFONT
            run_content._element.rPr.rFonts.set(
                qn('w:eastAsia'), TEXTFONT)
            run_content.font.size = Pt(TEXTFONTSIZE)
            run_content.font.bold = False


def replace(p):
    """ 替换函数 """
    # 替换符号
    if ')、' in p.text:
        p.text = p.text.replace(')、', '）')
    if '）、' in p.text:
        p.text = p.text.replace('）、', '）')
    if '(' in p.text:
        p.text = p.text.replace('(', '（')
    if ')' in p.text:
        p.text = p.text.replace(')', '）')
    if ',' in p.text:
        p.text = p.text.replace(',', '，')
    if ':' in p.text:
        p.text = p.text.replace(':', '：')
    if ';' in p.text:
        p.text = p.text.replace(';', '；')
    if '?' in p.text:
        p.text = p.text.replace('?', '？')
    if ' ' in p.text:
        p.text = p.text.replace(' ', '')
    return p


def fixDocx(docx):
    """ 主要格式 """
    lvl = 0
    for p in docx.paragraphs:
        # print(p.text)
        if p.text == '':
            continue
        else:
            lvl += 1
            # print(p.paragraph_format.first_line_indent)
            # print(p.style.font.size)
            p = replace(p)
            is_level1 = isLevel1(p)
            is_level2 = isLevel2(p)
            if lvl == 1:
                paragraphFun("title", p)
                for run_title in p.runs:
                    # print(run_title.text)
                    string = run_title.text
                    # print(string)
                    run_title._element.getparent().remove(run_title._element)
                    for i in string:
                        num_or_let = isNumberOrLetter(i)
                        text("title", is_level1, is_level2, num_or_let, p, i)
            else:
                paragraphFun("text", p)
                for run_content in p.runs:
                    # print(run_content.text)
                    string = run_content.text
                    # print(string)
                    p.paragraph_format.first_line_indent = 0  # 首行缩进
                    p.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '200')
                    run_content._element.getparent().remove(run_content._element)
                    # 替换格式："1、" --> "1."
                    for i in range(len(string)):
                        if i + 1 <= len(string):
                            if string[i:i+1] in '0123456789' and string[i+1:i+2] == "、":
                                # print(string[i:i+1])
                                string = string[i:i+1] + "." + string[i+2:]
                    for i in string:
                        num_or_let = isNumberOrLetter(i)
                        text("notitle", is_level1, is_level2, num_or_let, p, i)


def isNumberOrLetter(char):
    """ 判断是否为数字或字母 """
    number_and_letter_strs = '0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ'
    if char in number_and_letter_strs:
        return "num_or_let"
    else:
        return False


def pic_fix(docx, file, output_path):
    """ 图片处理 """
    img_path = output_path + "\image"
    file_name = path.splitext(file)[0]
    parts = docx.part.related_parts
    parts_values = parts.values()
    parts_keys = parts.keys()
    list_val = list(parts_values)
    list_key = list(parts_keys)
    parts_length = len(parts_values)
    if parts_length > 5:
        # print(type(list(parts_values)[-1]))
        k = 0
        for i in range(parts_length):
            # print(type(list_val[i]))
            if 'image' in str(type(list_val[i])):
                if not path.isdir(img_path):
                    makedirs(img_path)
                # print('找到图片数据')
                k += 1
                try:
                    img_data = parts[list_key[i]].image.blob
                    img_type = parts[list_key[i]].image.ext
                    full_path = f'{img_path}\{file_name}_image{k}.{img_type}'
                    print(f"··>提示<·· 正在输出：{full_path}")
                    with open(full_path, 'wb') as f:
                        f.write(img_data)
                except:
                    print(f"··>错误<·· 图片{k}输出失败！")
        if k == 0:
            print(f"··>提示<·· 未找到图片！")


def sureDo(func):
    """ 确定操作 """
    ipt = input(f"是否{func}？（输入Y或y确定）")
    if ipt.upper() == "Y":
        return True
    else:
        return False


def fixWord(docx_path, save_path, file, output_path, time_ipt, page_ipt, img_ipt):
    """ 文档处理 """
    docx = Document(docx_path)
    # 奇偶页不同
    docx.settings.odd_and_even_pages_header_footer = True

    # 页边距
    margin(docx)

    # 修改格式
    fixDocx(docx)

    # 添加时间后缀
    file_name = path.splitext(file)[0]
    if time_ipt:
        save_time = strftime("%m%d%H%M", localtime())
        save_path = output_path + f"\{file_name}" + save_time + ".docx"
    else:
        save_path = output_path + f"\{file_name}" + ".docx"

    # 设置页码
    if page_ipt:
        footer(docx)

    # 保存文档中的图片
    if img_ipt:
        pic_fix(docx, file, output_path)

    print(f"··>提示<·· 已保存：{save_path}")
    docx.save(save_path)


def inputPath():
    """ 输入路径 """
    print("··>提示<·· 可输入以下文件路径：")
    print("\t1.文件路径：单文件处理")
    print("\t2.文件夹路径：多文件处理（批量处理）")
    input_path = input("请输入路径（文件或文件夹）：")
    if path.isdir(input_path):
        # print("文件夹")
        dir_path = input_path
        path_info = "dir_path"
        return path_info, dir_path
    elif path.isfile(input_path):
        # print("文件")
        if input_path.endswith('.docx'):
            file_path = input_path
            path_info = "file_path"
            return path_info, file_path
        else:
            print("··>错误<·· 请输入.docx文件路径")
            return inputPath()
    else:
        print("··>错误<·· 路径不正确，请重新输入！")
        return inputPath()


def main():
    """ 主函数 """
    is_dir_file_path, input_path = inputPath()
    time_ipt = None
    page_ipt = None
    img_ipt = None
    if is_dir_file_path == "dir_path":
        dir_path = input_path
        output_path = dir_path + "\output"
        have_docx = 0
        for file in listdir(dir_path):
            if '~' in file:
                continue
            elif file.endswith('.docx'):
                if not path.isdir(output_path):
                    makedirs(output_path)
                have_docx += 1
                file_path = path.join(dir_path, file)
                if time_ipt == None:
                    time_ipt = sureDo("在生成文件名后面添加时间标记")
                    page_ipt = sureDo("添加页码")
                    img_ipt = sureDo("保存文档中的图片")
                fixWord(file_path, output_path + f'\{file}', file, output_path, time_ipt, page_ipt, img_ipt)
        if have_docx == 0:
            print("··>错误<·· 没有找到.docx文件")
    elif is_dir_file_path == "file_path":
        file_path = input_path
        # 文件名
        file = file_path.split("\\")[-1]
        # 输出路径
        dir_path = file_path.split("\\")
        dir_path.pop()
        result = '\\'.join(str(x) for x in dir_path)
        output_path = result + "\output"
        if not path.isdir(output_path):
            makedirs(output_path)
        if time_ipt == None:
            time_ipt = sureDo("在生成文件名后面添加时间标记")
            page_ipt = sureDo("添加页码")
            img_ipt = sureDo("保存文档中的图片")
        fixWord(file_path, output_path + f'\{file}', file, output_path, time_ipt, page_ipt, img_ipt)


if __name__ == '__main__':
    # 配置信息start
    # 页边距
    PAGETOPMARGIN = 3.7  # 页边距：上3.7cm
    PAGEBOTTOMMARGIN = 3.5  # 页边距：下3.5cm
    PAGELEFTMARGIN = 2.8  # 页边距：左2.8cm
    PAGERIGHTMARGIN = 2.6  # 页边距：右2.6cm
    # 标题
    TITLEFONT = '方正小标宋简体'  # 标题：方正小标宋简体
    TITLEFONTSIZE = 18  # 标题：小二号
    TITLEMARGIN = 33  # 标题：固定值33磅
    # 正文
    TEXTFONT = '仿宋_GB2312'  # 正文：仿宋_GB2312
    TEXTFONTSIZE = 14  # 正文：四号
    TEXTMARGIN = 28  # 正文：一般固定值28磅
    LEVEL1FONT = '黑体'  # 一级标题：黑体
    LEVEL2FONT = '楷体_GB2312'  # 二级标题：楷体_GB2312
    NUMBERFONT = 'Times New Roman'  # 数字&英文：TimesNewRoman字体
    # 页码
    PAGENUMBERFONTSIZE = 14  # 页码：四号
    PAGENUMBERFONT = '宋体'  # 页码：宋体
    # 配置信息end
    main()
    print("··>提示<·· 处理完成！")
    sleep(0.8)
    # 倒计时退出程序
    t = 5
    while t > 0:
        print(f"{t}秒后自动关闭", end="\r")
        t -= 1
        sleep(1)
