"""
   ______  ____  __   ____  _             _ _
  / ___\ \/ /  \/  | / ___|| |_ _   _  __| (_) ___
 | |    \  /| |\/| | \___ \| __| | | |/ _` | |/ _ \
 | |___ /  \| |  | |  ___) | |_| |_| | (_| | | (_) |
  \____/_/\_\_|  |_| |____/ \__|\__,_|\__,_|_|\___/

开发作者：晨小明
开发日期：2024/09/22
开发版本：v5.4.0.4_Dev
发布版本：v5.4.0.4_Release
修改日期：2026/08/26
主要功能：一、支持批量文档处理，输入文件夹路径，自动判断。
         二、读取.docx文件并设置格式；
         三、支持自定义格式设置：字体、字号、页边距、行距
         四、支持添加页码（可选）：4号半角宋体阿拉伯数字，数字左右各加一条4号“一字线”，奇数页在右侧左空一字，偶数页在左侧左空一字
         五、识别文档中的图片并输出（可选）：（注：图片可能会被压缩）
         六、替换功能
            1.符号替换
                将英文状态下的符号替换为中文状态下的相同符号，包含如下：
                "(" --> "（"
                ")" --> "）"
                "," --> "，"
                ":" --> "："
                ";" --> "；"
                ")、" --> "）"
                "?" --> "？"
                " " --> ""
                "1、" --> "1."
            2.其他格式
         七、输出文件名称含时间点，方便标记（可选）
         （注，本程序无法处理图片格式，如果图片独立成段，本程序所用API识别到图片会被默认是空段落，为了防止图片删除，只能放弃处理空段落及图片格式）
更新日志：
【修复】首行缩进数值单位错误的问题；
【修复】小三字号设置变量错误的问题；
【修复】当文件名中有空格，处理后打开文件提示路径错误的问题；
【优化】编号函数、重置函数的代码结构。
"""

from re import sub
from docx import Document
from docx.shared import Pt, Cm  # 用来设置字体的大小
from docx.oxml.ns import qn  # 控件名称
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT  # 设置对其方式
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from os import listdir, path, makedirs, getcwd, startfile, remove
from tkinter import Tk, filedialog, messagebox, Menu, ttk, Listbox, StringVar, END, Toplevel, Canvas, IntVar
from tkinter import font as tkFont
from datetime import datetime
from time import sleep
from PIL import Image, ImageTk
from webbrowser import open as webopen
from configparser import ConfigParser, NoOptionError
from upGrade import upGrade as update
from requests import get
from zipfile import ZipFile


def upGrade():
    """检测更新函数"""
    def progressBarTk():
        """下载进度窗口 """
        bar_tk = Toplevel(tk)
        bar_tk.geometry(f"+{x}+{y-50}")
        bar_tk.iconbitmap(icon_path)
        # 禁止调整窗口大小
        bar_tk.resizable(False, False)
        # 让进度窗口置顶（可选）
        # bar_tk.attributes('-topmost', True)
        # 关闭主窗口时同步关闭进度窗口
        bar_tk.transient(tk)
        bar_tk.title("正在下载最新版本...")
        bar_title = ttk.Label(bar_tk, text="0 / 0  0%", font=("微软雅黑", 12))
        bar_title.pack(padx=2, pady=2)
        progress_frm = ttk.Progressbar(bar_tk, mode="determinate", length=300, value=0, maximum=0)
        progress_frm.pack(padx=2, pady=2)
        # bar_tk.mainloop()
        return bar_tk, bar_title, progress_frm

    version_info = update()
    if version_info is None:
        writeHistory("更新失败！")
        messagebox.showinfo("提示", "更新失败！")
    else:
        latest_version = version_info['versionName']
        if latest_version == VERSION:
            writeHistory("当前已是最新版本！")
            messagebox.showinfo("提示", "当前已是最新版本！")
        else:
            update_date = version_info['updateDate']
            update_log = "".join(['\n' + i for i in version_info['updateLog']])
            is_up = messagebox.askyesno("提示", f"发现新版本，是否更新？\n当前版本：{VERSION}\n最新版本：{latest_version}\n更新日期：{datetime.fromtimestamp(update_date).strftime('%Y-%m-%d %H:%M:%S')}\n更新内容：{update_log}")
            if is_up:
                save_zip_path = filedialog.asksaveasfile(title="请选择保存路径", initialfile=f"fixWord_{latest_version}.zip", filetypes=[("zip", "*.zip")])
                # webopen(f"https://gitee.com/cxmStudio/fixWord/releases/download/{latest_version}/fixWord_{latest_version}.zip")
                if save_zip_path is None:
                    writeHistory("路径为空，取消更新！")
                    messagebox.showinfo("提示", "路径为空，取消更新！")
                else:
                    save_file_name = save_zip_path.name
                    writeHistory(f"保存路径：{save_file_name}")
                    down_res = get(f"https://gitee.com/cxmStudio/fixWord/releases/download/{latest_version}/fixWord_{latest_version}.zip", stream=True)
                    total_length = int(down_res.headers.get('Content-Length', 0))
                    total = 0
                    bar_tk, bar_title, progress_frm = progressBarTk()
                    total_size = f"{total_length / 1024 / 1024:.2f}MB"
                    progress_frm["maximum"] = total_length
                    if down_res.status_code == 200:
                        with open(save_file_name, 'wb') as f:
                            writeHistory("开始下载...")
                            for chunk in down_res.iter_content(chunk_size=8192):
                                if chunk:
                                    total += len(chunk)
                                    # 计算进度百分比和易读的文件大小
                                    downloaded = f"{total / 1024 / 1024:.2f}MB"
                                    # 更新进度条
                                    progress_frm["value"] = total
                                    bar_title.config(text=f"{downloaded} / {total_size}   {total / total_length * 100:.2f}%")
                                    # 强制刷新进度窗口UI
                                    bar_tk.update_idletasks()
                                    print(f"已下载：{downloaded} / {total_size}   {total / total_length * 100:.2f}%", end="\r")
                                    f.write(chunk)
                            writeHistory(f"下载完成，文件已保存至: {save_file_name}")
                        # 3秒后自动关闭进度窗口
                        bar_tk.after(3000, bar_tk.destroy)
                        messagebox.showinfo("提示", "下载完成！")
                        writeHistory(f"正在解压：{save_file_name}...")
                        try:
                            file_name_list = save_file_name.split('/')
                            file_name_list.pop()
                            zipout_path = '/'.join(file_name_list) + f"/fixWord_{latest_version}"
                            if not path.isdir(zipout_path):
                                makedirs(zipout_path)
                            with ZipFile(save_file_name, 'r') as f:
                                f.extractall(zipout_path)
                            writeHistory(f"解压完成！请重新打开本软件！解压路径：{zipout_path}")
                            messagebox.showinfo("提示", f"解压完成！请重新打开本软件！解压路径：{zipout_path}")
                        except PermissionError:
                            writeHistory(f"解压失败，文件被占用，请关闭")
                            messagebox.showerror("错误", f'解压失败，文件被占用，请关闭"{zipout_path}/fixWord_{latest_version}.exe"或保存至其他路径。')
                        except Exception as e:
                            writeHistory(f"解压失败，请手动解压文件：{save_file_name}\n错误信息：{e}")
                            messagebox.showerror("错误", f'解压失败，请手动解压文件：{save_file_name}\n错误信息：{e}')
                    else:
                        writeHistory(f"下载失败，请检查网络连接！\n{down_res.status_code}")
                        messagebox.showerror("提示", f"下载失败，请检查网络连接！\n{down_res.status_code}")
                    # 删除压缩包
                        # 删除压缩包
                    for attempt in range(3):
                        try:
                            if path.isfile(save_file_name):
                                remove(save_file_name)
                            break
                        except PermissionError:
                            if attempt < 2:
                                sleep(1)
                            else:
                                writeHistory(f"自动删除压缩包失败，请手动删除：{save_file_name}")
                                messagebox.showwarning("警告", f"无法自动删除压缩包，文件可能被占用，请手动删除：\n{save_file_name}")
            else:
                writeHistory("取消更新！")
                messagebox.showinfo("提示", "取消更新！")


class Judge():
    """判断事件"""
    def isLevel(p):
        """判断标题级别"""
        index_list = [
            ["一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、", "十一、", "十二、", "十三、", "十四、", "十五、", "十六、", "十七、", "十八、", "十九、", "二十、"],
            ["（一）", "（二）", "（三）", "（四）", "（五）", "（六）", "（七）", "（八）", "（九）", "（十）", "（十一）", "（十二）", "（十三）", "（十四）", "（十五）", "（十六）", "（十七）", "（十八）", "（十九）", "（二十）"],
            ["1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.", "11.", "12.", "13.", "14.", "15.", "16.", "17.", "18.", "19.", "20."],
            ["（1）", "（2）", "（3）", "（4）", "（5）", "（6）", "（7）", "（8）", "（9）", "（10）", "（11）", "（12）", "（13）", "（14）", "（15）", "（16）", "（17）", "（18）", "（19）", "（20）"]
        ]
        for i in index_list:
            for j in i:
                if j in p.text[:len(j) + 1]:
                    if index_list.index(i) == 0 or index_list.index(i) == 1:
                        if '。' in p.text:
                            p.text = p.text.replace('。', '')
                        if '？' in p.text:
                            p.text = p.text.replace('？', '')
                        if '：' in p.text:
                            p.text = p.text.replace('：', '')
                        if '；' in p.text:
                            p.text = p.text.replace('；', '')
                        return f"level{index_list.index(i) + 1}"
                    elif index_list.index(i) == 2 or index_list.index(i) == 3:
                        return f"level{index_list.index(i) + 1}"
                else:
                    continue

    def isNumberOrLetter(char):
        """判断是否为数字或字母"""
        number_and_letter_strs = '0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ'
        if char in number_and_letter_strs:
            return "num_or_let"
        else:
            return False


class DocxProcessing():
    """文档处理事件"""
    def margin(docx):
        """设置页边距"""
        global data
        for s in docx.sections:
            s.top_margin = Cm(float(data["margin"]["t_value"]))
            s.bottom_margin = Cm(float(data["margin"]["b_value"]))
            s.left_margin = Cm(float(data["margin"]["l_value"]))
            s.right_margin = Cm(float(data["margin"]["r_value"]))

    def footer(docx):
        """设置页脚，添加页码"""
        # print(len(docx.sections))
        def AddFooterNumber(p):
            p.clear()
            t1 = p.add_run("— ")
            font = t1.font
            font.name = '宋体'
            font.size = Pt(14)  # 14号字体
            t1._element.rPr.rFonts.set(qn("w:eastAsia"), '宋体')

            run1 = p.add_run('')
            fldChar1 = OxmlElement('w:fldChar')  # creates a new element
            fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
            run1._element.append(fldChar1)

            run2 = p.add_run('')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
            instrText.text = 'PAGE'
            font = run2.font
            font.name = '宋体'
            font.size = Pt(14)  # 14号字体
            run2._element.rPr.rFonts.set(qn("w:eastAsia"), '宋体')
            run2._element.append(instrText)

            run3 = p.add_run('')
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run3._element.append(fldChar2)

            t2 = p.add_run(" —")
            font = t2.font
            font.name = '宋体'
            font.size = Pt(14)  # 14号字体
            t2._element.rPr.rFonts.set(qn("w:eastAsia"), '宋体')

        for s in docx.sections:
            # print(s.footer)
            footer = s.footer  # 获取第一个节的页脚
            footer.is_linked_to_previous = True  # 编号续前一节
            paragraph = footer.paragraphs[0]  # 获取页脚的第一个段落
            DocxProcessing.paragraphFun("odd_footer", paragraph)
            AddFooterNumber(paragraph)
            even_footer = s.even_page_footer  # 获取第一个节的页脚
            even_footer.is_linked_to_previous = True  # 编号续前一节
            paragraph = even_footer.paragraphs[0]  # 获取页脚的第一个段落
            DocxProcessing.paragraphFun("even_footer", paragraph)
            AddFooterNumber(paragraph)

    def paragraphFun(is_title, p, is_level=""):
        """段落函数"""
        def setLineSpacing(data_font):
            # 设置行距
            if data_font["font_ls_lbl_txt"] == "倍":
                font_ls_vlu = float(data_font["font_ls_vlu"])
                p.paragraph_format.element.pPr.spacing.set(qn("w:line"), f'{font_ls_vlu * 240}')
                p.paragraph_format.element.pPr.spacing.set(qn("w:lineRule"), 'auto')
            else:
                p.paragraph_format.line_spacing = Pt(float(data_font["font_ls_vlu"]))  # 行距
            # 设置段前段后，统一使用底层XML设置间距，防止高层API刷新覆盖行距属性
            pPr_spacing = p.paragraph_format.element.pPr.spacing
            pPr_spacing.set(qn("w:before"), str(int(float(data_font["font_b_s_vlu"]) * 20)))  # 1磅 = 20 twips
            pPr_spacing.set(qn("w:after"), str(int(float(data_font["font_a_s_vlu"]) * 20)))   # 1磅 = 20 twips
            # 设置左右缩进，统一使用底层XML设置缩进，防止高层API生成冲突的Chars属性
            pPr_ind = p.paragraph_format.element.pPr.ind
            pPr_ind.set(qn("w:left"), str(int(float(data_font["font_l_idt_vlu"]) * 20)))  # 1磅 = 20 twips
            pPr_ind.set(qn("w:right"), str(int(float(data_font["font_r_idt_vlu"]) * 20)))
            # 设置首行缩进，显式移除可能引发冲突的Chars属性
            if pPr_ind.get(qn("w:leftChars")) is not None:
                del pPr_ind.attrib[qn("w:leftChars")]
            if pPr_ind.get(qn("w:rightChars")) is not None:
                del pPr_ind.attrib[qn("w:rightChars")]
            # 移除首行缩进属性
            if pPr_ind.get(qn("w:firstLine")) is not None:
                del pPr_ind.attrib[qn("w:firstLine")]
            first_line_idt = float(data_font["font_f_line_vlu"])
            if first_line_idt > 0 and is_title != "odd_footer" and is_title != "even_footer":
                first_line_indent = int(first_line_idt * 240)  # 1字符约等于12磅 = 240 twips
                pPr_ind.set(qn("w:firstLineChars"), f'{first_line_idt * 100}')
                pPr_ind.set(qn("w:firstLine"), str(first_line_indent))
                # 移除可能存在的悬挂缩进属性，防止冲突
                if pPr_ind.get(qn("w:hanging")) is not None:
                    del pPr_ind.attrib[qn("w:hanging")]
                if pPr_ind.get(qn("w:hangingChars")) is not None:
                    del pPr_ind.attrib[qn("w:hangingChars")]
            else:
                pPr_ind.set(qn("w:firstLineChars"), '0')
                pPr_ind.set(qn("w:firstLine"), '0')

        global data
        if p.paragraph_format.element.pPr is None:
            p.paragraph_format.element.append(parse_xml(r'<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        # 判断 ind 是否存在，方便后边设置首行缩进
        if p.paragraph_format.element.pPr.ind is None:
            p.paragraph_format.element.pPr.append(parse_xml(r'<w:ind xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        # 判断 spacing 是否存在，方便后边设置行距
        if p.paragraph_format.element.pPr.spacing is None:
            p.paragraph_format.element.pPr.append(parse_xml(r'<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        # 统一设置对齐方式
        pgp_almt = data["main"]["pgp_almt"]
        if pgp_almt == "居中":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif pgp_almt == "左对齐":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif pgp_almt == "右对齐":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # 各标题设置
        if is_title == "title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            setLineSpacing(data["title_font"])
        elif is_title == "odd_footer":
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p.paragraph_format.right_indent = Pt(14)
            p.paragraph_format.line_spacing = Pt(28)
        elif is_title == "even_footer":
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.line_spacing = Pt(28)
        else:
            setLineSpacing(data["mb_font"])
            setLineSpacing(data["num_font"])
            if is_level == "level1":
                setLineSpacing(data["1title_font"])
            elif is_level == "level2":
                setLineSpacing(data["2title_font"])
            elif is_level == "level3":
                setLineSpacing(data["3title_font"])
            elif is_level == "level4":
                setLineSpacing(data["4title_font"])
        # 统一设置孤行控制
        single_crl_value = data["main"]["single_crl_value"]
        if single_crl_value == "1":
            p.paragraph_format.widow_control = True
        else:
            p.paragraph_format.widow_control = False

    def text(is_title, is_digit, p, i, is_level=""):
        """正文函数"""
        def checkFontSize(size):
            try:
                return FONTSIZEDICT[size]
            except KeyError:
                try:
                    return float(size)
                except Exception as e:
                    messagebox.showwarning("警告", f"字体大小输入错误：{size}， 请检查！{e}\n已重置为14")
                    return 14

        def setFont(run, font_):
            """设置字体"""
            run.font.name = font_["font_name"]
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_["font_name"])
            run.font.size = Pt(checkFontSize(font_["font_size"]))
            if font_["font_bold"] == 1:
                run.bold = True
            else:
                run.bold = False
        global data
        if is_title == "title":
            run = p.add_run(i)
            if is_digit == "num_or_let":
                run.font.name = data["num_font"]["font_name"]
            else:
                run.font.name = data["title_font"]["font_name"]
                run._element.rPr.rFonts.set(qn('w:eastAsia'), data["title_font"]["font_name"])
            run.font.size = Pt(checkFontSize(data["title_font"]["font_size"]))
            if data["title_font"]["font_bold"] == 1:
                run.bold = True
            else:
                run.bold = False
        else:
            run_content = p.add_run(i)
            if is_digit == "num_or_let":
                run_content.font.name = data["num_font"]["font_name"]
                run_content.font.size = Pt(checkFontSize(data["num_font"]["font_size"]))
                if data["num_font"]["font_bold"] == 1:
                    run_content.bold = True
                else:
                    run_content.bold = False
            else:
                setFont(run_content, data["mb_font"])
                if is_level == "level1":
                    setFont(run_content, data["1title_font"])
                elif is_level == "level2":
                    setFont(run_content, data["2title_font"])
                elif is_level == "level3":
                    setFont(run_content, data["3title_font"])
                elif is_level == "level4":
                    setFont(run_content, data["4title_font"])

    def fixDocx(docx):
        """主要格式"""
        lvl = 0
        new_p_cnt = 0
        for idx, p in enumerate(docx.paragraphs):
            if '\n' in p.text:
                p, new_p_cnt = Replace.brFix(p, docx, idx, new_p_cnt)
            if p.text == '':
                DocxProcessing.paragraphFun("", p)
                DocxProcessing.text("", "", p, " ")
            else:
                lvl += 1
                p = Replace.replace(p)
                if lvl == 1:
                    DocxProcessing.paragraphFun("title", p)
                    for run_title in p.runs:
                        # print(run_title.text)
                        run_title._element.getparent().remove(run_title._element)
                        for i in run_title.text:
                            num_or_let = Judge.isNumberOrLetter(i)
                            DocxProcessing.text("title", num_or_let, p, i)
                else:
                    is_level = Judge.isLevel(p)
                    DocxProcessing.paragraphFun("text", p, is_level)
                    for run_content in p.runs:
                        # print(run_content.text)
                        run_content._element.getparent().remove(run_content._element)
                        for i in run_content.text:  # 遍历字符串
                            num_or_let = Judge.isNumberOrLetter(i)
                            DocxProcessing.text("notitle", num_or_let, p, i, is_level)

    def getPic(docx, file, output_path):
        """图片处理"""
        img_path = output_path + "\image"
        file_name = path.splitext(file)[0]
        parts = docx.part.related_parts
        parts_values = parts.values()
        parts_keys = parts.keys()
        list_val = list(parts_values)
        list_key = list(parts_keys)
        parts_length = len(parts_values)
        if parts_length > 5:
            # print(type(list(parts_values)[-1]))
            k = 0
            for i in range(parts_length):
                # print(type(list_val[i]))
                if 'image' in str(type(list_val[i])):
                    if not path.isdir(img_path):
                        makedirs(img_path)
                    # print('找到图片数据')
                    k += 1
                    try:
                        img_data = parts[list_key[i]].image.blob
                        img_type = parts[list_key[i]].image.ext
                        full_path = f'{img_path}\{file_name}_image{k}.{img_type}'
                        writeHistory(f"··>提示<·· 正在输出：{full_path}")
                        with open(full_path, 'wb') as f:
                            f.write(img_data)
                    except:
                        writeHistory(f"··>错误<·· 图片{k}输出失败！")
            if k == 0:
                writeHistory(f"··>提示<·· 未找到图片！")

    def fixWord(docx_path, file, output_path, time_ipt, page_ipt, img_ipt):
        """文档处理"""
        # 自动编号识别 https://www.iotword.com/22828.html 并新建文档
        new_docx = WithNumberDocxReader(docx_path).texts

        # 页边距
        DocxProcessing.margin(new_docx)

        # 修改格式
        DocxProcessing.fixDocx(new_docx)

        # 添加时间后缀
        file_name = path.splitext(file)[0]
        if time_ipt == "1":
            save_time = datetime.now().strftime("%m%d%H%M")
            save_path = output_path + f"\{file_name}" + save_time + ".docx"
        else:
            save_time = ""
            save_path = output_path + f"\{file_name}" + ".docx"

        # 设置页码
        if page_ipt == "1":
            # 奇偶页不同
            new_docx.settings.odd_and_even_pages_header_footer = True
            DocxProcessing.footer(new_docx)

        # 保存文档中的图片
        if img_ipt == "1":
            DocxProcessing.getPic(new_docx, file, output_path)

        # 保存文档
        try:
            new_docx.save(save_path)
            return save_path, True
        except PermissionError:
            return save_path, False


class Replace():
    """替换函数"""
    def brFix(p, docx, idx, new_p_cnt):
        """换行符修复"""
        # 利用xml添加段落处理
        txt_list = p.text.split('\n')
        p.text = txt_list[0]
        idx += 1
        for txt in txt_list[1:]:
            # 创建新段落
            new_paragraph = docx.add_paragraph(txt)
            new_paragraph = Replace.replace(new_paragraph)
            # 获取文档XML结构
            doc_ele = docx._element.body
            # 找到新段落的XML元素
            new_p = new_paragraph._element
            # 从文档中移除新段落（临时）
            doc_ele.remove(new_p)
            # 计算实际插入位置（处理表格等复杂元素）
            for i, child in enumerate(doc_ele):
                if isinstance(child, CT_P) or isinstance(child, CT_Tbl):
                    if i == idx:
                        doc_ele.insert(i + new_p_cnt, new_p)
                        DocxProcessing.paragraphFun("", new_paragraph)
                        is_level = Judge.isLevel(new_paragraph)
                        for run_content in new_paragraph.runs:
                            run_content._element.getparent().remove(run_content._element)
                            for j in run_content.text:  # 遍历字符串
                                DocxProcessing.text("notitle", Judge.isNumberOrLetter(j), new_paragraph, j, is_level)
                        new_p_cnt += 1
                        break
        return p, new_p_cnt

    def replace(p):
        """替换函数"""
        # 替换符号
        if '(' in p.text:
            p.text = p.text.replace('(', '（')
        if ')' in p.text:
            p.text = p.text.replace(')', '）')
        if ',' in p.text:
            p.text = p.text.replace(',', '，')
        if ':' in p.text:
            p.text = p.text.replace(':', '：')
        if ';' in p.text:
            p.text = p.text.replace(';', '；')
        if '?' in p.text:
            p.text = p.text.replace('?', '？')
        if '》、' in p.text:
            p.text = p.text.replace('》、', '》')
        if '．' in p.text:  # U+ff0e
            p.text = p.text.replace('．', '.')
        if ' ' in p.text:  # 空格
            p.text = p.text.replace(' ', '')
        if '　' in p.text:  # U+3000
            p.text = p.text.replace('　', '')
        if ' ' in p.text:  # U+2003
            p.text = p.text.replace(' ', '')
        if '	' in p.text:  # \t
            p.text = p.text.replace('	', '')
        if '\xa0' in p.text:  # \xa0 U+00A0
            p.text = p.text.replace('\xa0', '')
        if ')、' in p.text[:4]:
            p.text = p.text.replace(')、', '）')
        if '）、' in p.text[:4]:
            p.text = p.text.replace('）、', '）')
        # 替换数字后的、为.
        if len(p.text) >= 3 and p.text[:3][0].isdigit() and '、' in p.text[:3]:
            p.text = p.text[:3].replace('、', '.') + p.text[3:]
        return p


class WithNumberDocxReader:
    """识别编号列表并替换"""
    ideographTraditional = "甲乙丙丁戊己庚辛壬癸"
    ideographZodiac = "子丑寅卯辰巳午未申酉戌亥"

    def __init__(self, docx, gap_text="\t"):
        self.docx_path = docx
        self.docx = Document(docx)
        self.numId2style = self.get_style_data()
        self.gap_text = gap_text
        self.cnt = {}
        self.cache = {}

    @property
    def texts(self):
        self.cnt.clear()
        self.cache.clear()
        for paragraph in self.docx.paragraphs:
            try:
                number_text = self.get_number_text(paragraph._element.pPr.numPr)
            except AttributeError as e:
                number_text = ""
            except Exception as e:
                writeLog(f"提示：无法获取文档编号部分：{e if e else '未知原因'}")
                number_text = ""
            # 先获取旧段落文本
            p_text = paragraph.text
            if number_text:
                # 如果有编号，则清除旧段落内容和样式
                paragraph.clear()
                # 如果编号样式存在，则清除旧段落编号样式
                if paragraph._element.pPr.numPr is not None:
                    paragraph._element.pPr.numPr.clear()
                # 拼接编号文本和段落文本
                paragraph.text = number_text + p_text

        return self.docx

    def get_style_data(self):
        try:
            numbering_part = self.docx.part.numbering_part._element
        except Exception as e:
            writeLog(f"未找到编号列表！{f'{e}，' if e else ''}文件：{self.docx_path}")
            return {}
        abstractId2numId = {num.abstractNumId.val: num.numId for num in numbering_part.num_lst}
        numId2style = {}
        for abstractNumIdTag in numbering_part.findall(qn("w:abstractNum")):
            abstractNumId = abstractNumIdTag.get(qn("w:abstractNumId"))
            numId = abstractId2numId[int(abstractNumId)]
            for lvlTag in abstractNumIdTag.findall(qn("w:lvl")):
                ilvl = lvlTag.get(qn("w:ilvl"))
                style = {tag.tag[tag.tag.rfind("}") + 1:]: tag.get(qn("w:val")) for tag in lvlTag.xpath("./*[@w:val]", namespaces=numbering_part.nsmap)}
                if "numFmt" not in style:
                    numFmtVal = lvlTag.xpath("./mc:AlternateContent/mc:Fallback/w:numFmt/@w:val", namespaces=numbering_part.nsmap)
                    if numFmtVal and numFmtVal[0] == "decimal":
                        numFmt_format = lvlTag.xpath("./mc:AlternateContent/mc:Choice/w:numFmt/@w:format", namespaces=numbering_part.nsmap)
                        if numFmt_format:
                            style["numFmt"] = "decimal" + numFmt_format[0].split(",")[0]
                if style.get("numFmt") == "decimalZero":
                    style["numFmt"] = "decimal01"
                numId2style[(numId, int(ilvl))] = style
        return numId2style

    @staticmethod
    def int2upperLetter(num):
        result = []
        while num > 0:
            num -= 1
            remainder = num % 26
            result.append(chr(remainder + ord('A')))
            num //= 26
        return "".join(reversed(result))

    @staticmethod
    def int2upperRoman(num):
        t = [
            (1000, 'M'), (900, 'CM'), (500, 'D'),
            (400, 'CD'), (100, 'C'), (90, 'XC'),
            (50, 'L'), (40, 'XL'), (10, 'X'),
            (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')
        ]
        roman_num = ''
        i = 0
        while num > 0:
            val, syb = t[i]
            for _ in range(num // val):
                roman_num += syb
                num -= val
            i += 1
        return roman_num

    @staticmethod
    def int2cardinalText(num):
        if not isinstance(num, int) or num < 0 or num > 999999999:
            raise ValueError("Invalid number: must be a positive integer within four digits")
        base = ["Zero", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen",
                "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
        tens = ["", "", "Twenty", "Thirty", "Fourty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        thousands = ["", "Thousand", "Million", "Billion"]

        def two_digits(n):
            if n < 20:
                return base[n]
            ten, unit = divmod(n, 10)
            if unit == 0:
                return f"{tens[ten]}"
            else:
                return f"{tens[ten]}-{base[unit]}"

        def three_digits(n):
            hundred, rest = divmod(n, 100)
            if hundred == 0:
                return two_digits(rest)
            result = f"{base[hundred]} hundred "
            if rest > 0:
                result += two_digits(rest)
            return result.strip()

        if num < 99:
            return two_digits(num)
        chunks = []
        while num > 0:
            num, remainder = divmod(num, 1000)
            chunks.append(remainder)
        words = []
        for i in range(len(chunks) - 1, -1, -1):
            if chunks[i] == 0:
                continue
            chunk_word = three_digits(chunks[i])
            if thousands[i]:
                chunk_word += f" {thousands[i]}"
            words.append(chunk_word)
        words = " ".join(words).lower()
        return words[0].upper() + words[1:]

    @staticmethod
    def int2ordinalText(num):
        if not isinstance(num, int) or num < 0 or num > 999999:
            raise ValueError("Invalid number: must be a positive integer within four digits")
        base = ["Zero", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
        baseth = ['Zeroth', 'First', 'Second', 'Third', 'Fourth', 'Fifth', 'Sixth', 'Seventh', 'Eighth', 'Ninth', 'Tenth', 'Eleventh', 'Twelfth', 'Thirteenth', 'Fourteenth', 'Fifteenth', 'Sixteenth', 'Seventeenth', 'Eighteenth', 'Nineteenth', 'Twentieth']
        tens = ["", "", "Twenty", "Thirty", "Fourty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        tensth = ["", "", "Twentieth", "Thirtieth", "Fortieth", "Fiftieth", "Sixtieth", "Seventieth", "Eightieth", "Ninetieth"]

        def two_digits(n):
            if n <= 20:
                return baseth[n]
            ten, unit = divmod(n, 10)
            result = tensth[ten]
            if unit != 0:
                result = f"{tens[ten]}-{baseth[unit]}"
            return result

        thousand, num = divmod(num, 1000)
        result = []
        if thousand > 0:
            if num == 0:
                return f"{WithNumberDocxReader.int2cardinalText(thousand)} thousandth"
            result.append(f"{WithNumberDocxReader.int2cardinalText(thousand)} thousand")
        hundred, num = divmod(num, 100)
        if hundred > 0:
            if num == 0:
                result.append(f"{base[hundred]} hundredth")
                return " ".join(result)
            result.append(f"{base[hundred]} hundred")
        result.append(two_digits(num))
        result = " ".join(result).lower()
        return result[0].upper() + result[1:]

    @staticmethod
    def int2Chinese(num, ch_num, units):
        if not (0 <= num <= 99999999):
            raise ValueError("仅支持小于一亿以内的正整数")

        def int2Chinese_in(num, ch_num, units):
            if not (0 <= num <= 9999):
                raise ValueError("仅支持小于一万以内的正整数")
            result = [ch_num[int(i)] + unit for i, unit in zip(reversed(str(num).zfill(4)), units)]
            result = "".join(reversed(result))
            zero_char = ch_num[0]
            result = sub(f"(?:{zero_char}[{units}])+", zero_char, result)
            result = result.rstrip(units[0])
            if result != zero_char:
                result = result.rstrip(zero_char)
            if result.lstrip(zero_char).startswith("一十"):
                result = result.replace("一", "")
            return result

        if num < 10000:
            result = int2Chinese_in(num, ch_num, units)
        else:
            left = num // 10000
            right = num % 10000
            result = int2Chinese_in(left, ch_num, units) + "万" + int2Chinese_in(right, ch_num, units)
        if result != ch_num[0]:
            result = result.strip(ch_num[0])
        return result

    @staticmethod
    def int2ChineseCounting(num):
        return WithNumberDocxReader.int2Chinese(num, ch_num='〇一二三四五六七八九', units='个十百千')

    @staticmethod
    def int2ChineseLegalSimplified(num):
        return WithNumberDocxReader.int2Chinese(num, ch_num='零壹贰叁肆伍陆柒捌玖', units='个拾佰仟')

    def get_number_text(self, numpr):
        if numpr is None or numpr.numId.val == 0:
            return ""
        numId = numpr.numId.val
        ilvl = numpr.ilvl.val
        style = self.numId2style[(numId, ilvl)]
        numFmt: str = style.get("numFmt")
        lvlText = style.get("lvlText")
        if (numId, ilvl) in self.cnt:
            self.cnt[(numId, ilvl)] += 1
        else:
            self.cnt[(numId, ilvl)] = int(style["start"])
        pos = self.cnt[(numId, ilvl)]
        num_text = str(pos)
        if numFmt.startswith('decimal'):
            num_text = num_text.zfill(numFmt.count("0") + 1)
        elif numFmt == 'upperRoman':
            num_text = self.int2upperRoman(pos)
        elif numFmt == 'lowerRoman':
            num_text = self.int2upperRoman(pos).lower()
        elif numFmt == 'upperLetter':
            num_text = self.int2upperLetter(pos)
        elif numFmt == 'lowerLetter':
            num_text = self.int2upperLetter(pos).lower()
        elif numFmt == 'ordinal':
            num_text = f"{pos}{'th' if 11 <= pos <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(pos % 10, 'th')}"
        elif numFmt == 'cardinalText':
            num_text = self.int2cardinalText(pos)
        elif numFmt == 'ordinalText':
            num_text = self.int2ordinalText(pos)
        elif numFmt == 'ideographTraditional':
            if 1 <= pos <= 10:
                num_text = self.ideographTraditional[pos - 1]
        elif numFmt == 'ideographZodiac':
            if 1 <= pos <= 12:
                num_text = self.ideographZodiac[pos - 1]
        elif numFmt == 'chineseCounting':
            num_text = self.int2ChineseCounting(pos)
        elif numFmt == 'chineseCountingThousand':
            num_text = self.int2ChineseCounting(pos)
        elif numFmt == 'chineseLegalSimplified':
            num_text = self.int2ChineseLegalSimplified(pos)
        elif numFmt == 'decimalEnclosedCircleChinese':
            pass
        self.cache[(numId, ilvl)] = num_text
        for i in range(0, ilvl + 1):
            lvlText = lvlText.replace(f'%{i + 1}', self.cache.get((numId, i), ""))
        suff_text = {"space": " ", "nothing": ""}.get(style.get("suff"), self.gap_text)
        lvlText += suff_text
        return lvlText


class PathEvents():
    """路径事件"""
    def inputPath():
        """输入路径 """
        input_path = type_radio_value.get()
        if input_path == "file_path":
            file_path = filedialog.askopenfile(title="请选择文件", filetypes=[("docx文件", "*.docx")])
            if file_path != None:
                path_entry.delete(0, END)
                path_entry.insert(0, file_path.name)
        elif input_path == "dir_path":
            dir_path = filedialog.askdirectory(title="请选择文件夹")
            if dir_path != "":
                path_entry.delete(0, END)
                path_entry.insert(0, dir_path)

    def inputFile():
        """选择文件 """
        path_button.config(text="选择文件")
        path_entry.delete(0, END)
        PathEvents.inputPath()

    def inputDir():
        """选择文件夹 """
        path_button.config(text="选择文件夹")
        path_entry.delete(0, END)
        PathEvents.inputPath()


class InitFile():
    """配置文件"""

    def __init__(self, is_auto_import=False):
        self.is_auto_import = is_auto_import
        self.config_path = path.join(path.dirname(__file__), "fixWord_config.ini")
        self.config = ConfigParser()

        # print(data)
        # 确保配置文件存在
        if not path.exists(self.config_path):
            self.config['DEFAULT'] = {}
            with open(self.config_path, 'w') as configfile:
                self.config.write(configfile)

    def saveConfig(self):
        """保存当前配置"""
        # 检查是否有配置文件
        if path.exists(self.config_path):
            self.config.read(self.config_path, encoding="utf-8")
            if self.config.sections() != []:
                isgoon = messagebox.askyesno("提示", "默认配置文件已存在，点击【是】会替换原有配置，是否继续？")
                if not isgoon:
                    writeHistory("取消保存配置！")
                    return
        # 解析数据
        global data
        data = SystemEvents.getUserInput()
        for section, options in data.items():
            if not self.config.has_section(section):
                self.config.add_section(section)
            if isinstance(options, dict):
                for key, value in options.items():
                    self.config.set(section, key, str(value))
            else:
                self.config.set("DEFAULT", section, str(options))

        # 保存配置
        with open(self.config_path, "w", encoding="utf-8") as f:
            self.config.write(f)
        writeHistory("保存配置成功！")

    def esaveConfig(self):
        """配置另存为"""
        # 检查是否有配置文件
        config_path = filedialog.asksaveasfilename(title="请选择保存配置路径", filetypes=[("配置文件", "*.ini")], initialfile=f"fixWord_config_{datetime.now().strftime('%Y%m%d%H%M%S')}.ini")
        if config_path == "":
            writeHistory("取消保存配置！")
            return
        if path.exists(config_path):
            isgoon = messagebox.askyesno("提示", "配置文件已存在，点击【是】会替换原有配置，是否继续？")
            if not isgoon:
                writeHistory("取消保存配置！")
                return

        # 解析数据
        global data
        data = SystemEvents.getUserInput()
        for section, options in data.items():
            if not self.config.has_section(section):
                self.config.add_section(section)
            if isinstance(options, dict):
                for key, value in options.items():
                    self.config.set(section, key, str(value))
            else:
                self.config.set("DEFAULT", section, str(options))

        # 保存配置
        with open(config_path, "w", encoding="utf-8") as f:
            self.config.write(f)
        writeHistory(f"保存配置成功！路径：{config_path}")
        messagebox.showinfo("提示", f"保存配置成功！\n路径：{config_path}")

    def importConfig(self):
        """导入配置"""
        if self.is_auto_import:
            config_path = self.config_path
        else:
            config_path = filedialog.askopenfilename(title="请选择配置文件", filetypes=[("配置文件", "*.ini")], initialdir=path.dirname(__file__))
        if config_path == "":
            writeHistory("取消导入配置！")
            return
        # 读取配置
        self.config.read(config_path, encoding="utf-8")
        # 获取配置
        for section in self.config.sections():
            if section not in data:
                data[section] = {}
            for key, value in self.config.items(section):
                data[section][key] = value
        # print(data)
        auto_import_ini_vlu.set(data['main']['auto_import'])
        font_title_name_frm_combox.set(data['title_font']['font_name']), font_title_size_frm_combox.set(data['title_font']['font_size']), font_title_ls_frm_combox.set(data['title_font']['font_ls']), font_ls_vlu.set(data['title_font']['font_ls_vlu']), font_ls_frm_lbl_b.config(text=f"{data['title_font']['font_ls_lbl_txt']}"), font_spacing_b_vlu.set(
            data['title_font']['font_b_s_vlu']), font_spacing_a_vlu.set(data['title_font']['font_a_s_vlu']), font_indent_left_vlu.set(data['title_font']['font_l_idt_vlu']), font_indent_right_vlu.set(data['title_font']['font_r_idt_vlu']), font_first_line_vlu.set(data['title_font']['font_f_line_vlu']), font_bold_frm_vlu.set(data['title_font']['font_bold'])
        font_title_name_frm1_combox.set(data['1title_font']['font_name']), font_title_size_frm1_combox.set(data['1title_font']['font_size']), font_title_ls_frm1_combox.set(data['1title_font']['font_ls']), font_ls_vlu1.set(data['1title_font']['font_ls_vlu']), font_ls_frm_lbl_b1.config(text=f"{data['1title_font']['font_ls_lbl_txt']}"), font_spacing_b_vlu1.set(
            data['1title_font']['font_b_s_vlu']), font_spacing_a_vlu1.set(data['1title_font']['font_a_s_vlu']), font_indent_left_vlu1.set(data['1title_font']['font_l_idt_vlu']), font_indent_right_vlu1.set(data['1title_font']['font_r_idt_vlu']), font_first_line_vlu1.set(data['1title_font']['font_f_line_vlu']), font_bold_frm_vlu1.set(data['1title_font']['font_bold'])
        font_title_name_frm2_combox.set(data['2title_font']['font_name']), font_title_size_frm2_combox.set(data['2title_font']['font_size']), font_title_ls_frm2_combox.set(data['2title_font']['font_ls']), font_ls_vlu2.set(data['2title_font']['font_ls_vlu']), font_ls_frm_lbl_b2.config(text=f"{data['2title_font']['font_ls_lbl_txt']}"), font_spacing_b_vlu2.set(
            data['2title_font']['font_b_s_vlu']), font_spacing_a_vlu2.set(data['2title_font']['font_a_s_vlu']), font_indent_left_vlu2.set(data['2title_font']['font_l_idt_vlu']), font_indent_right_vlu2.set(data['2title_font']['font_r_idt_vlu']), font_first_line_vlu2.set(data['2title_font']['font_f_line_vlu']), font_bold_frm_vlu2.set(data['2title_font']['font_bold'])
        font_title_name_frm3_combox.set(data['3title_font']['font_name']), font_title_size_frm3_combox.set(data['3title_font']['font_size']), font_title_ls_frm3_combox.set(data['3title_font']['font_ls']), font_ls_vlu3.set(data['3title_font']['font_ls_vlu']), font_ls_frm_lbl_b3.config(text=f"{data['3title_font']['font_ls_lbl_txt']}"), font_spacing_b_vlu3.set(
            data['3title_font']['font_b_s_vlu']), font_spacing_a_vlu3.set(data['3title_font']['font_a_s_vlu']), font_indent_left_vlu3.set(data['3title_font']['font_l_idt_vlu']), font_indent_right_vlu3.set(data['3title_font']['font_r_idt_vlu']), font_first_line_vlu3.set(data['3title_font']['font_f_line_vlu']), font_bold_frm_vlu3.set(data['3title_font']['font_bold'])
        font_title_name_frm4_combox.set(data['4title_font']['font_name']), font_title_size_frm4_combox.set(data['4title_font']['font_size']), font_title_ls_frm4_combox.set(data['4title_font']['font_ls']), font_ls_vlu4.set(data['4title_font']['font_ls_vlu']), font_ls_frm_lbl_b4.config(text=f"{data['4title_font']['font_ls_lbl_txt']}"), font_spacing_b_vlu4.set(
            data['4title_font']['font_b_s_vlu']), font_spacing_a_vlu4.set(data['4title_font']['font_a_s_vlu']), font_indent_left_vlu4.set(data['4title_font']['font_l_idt_vlu']), font_indent_right_vlu4.set(data['4title_font']['font_r_idt_vlu']), font_first_line_vlu4.set(data['4title_font']['font_f_line_vlu']), font_bold_frm_vlu4.set(data['4title_font']['font_bold'])
        font_mb_name_frm_combox.set(data['mb_font']['font_name']), font_mb_size_frm_combox.set(data['mb_font']['font_size']), font_mb_ls_frm_combox.set(data['mb_font']['font_ls']), font_mb_ls_vlu.set(data['mb_font']['font_ls_vlu']), font_mb_ls_frm_lbl_b.config(text=f"{data['mb_font']['font_ls_lbl_txt']}"), font_mb_spacing_b_vlu.set(
            data['mb_font']['font_b_s_vlu']), font_mb_spacing_a_vlu.set(data['mb_font']['font_a_s_vlu']), font_mb_indent_left_vlu.set(data['mb_font']['font_l_idt_vlu']), font_mb_indent_right_vlu.set(data['mb_font']['font_r_idt_vlu']), font_mb_first_line_vlu.set(data['mb_font']['font_f_line_vlu']), font_mb_bold_frm_vlu.set(data['mb_font']['font_bold'])
        font_num_name_frm_combox.set(data['num_font']['font_name']), font_num_size_frm_combox.set(data['num_font']['font_size']), font_num_bold_frm_vlu.set(data['num_font']['font_bold'])
        # 行距下拉框逻辑处理
        if len(self.config.sections()) > 0:
            if data['title_font']['font_ls_lbl_txt'] == "倍":
                font_title_ls_ent.config(state="disabled")
            else:
                font_title_ls_ent.config(state="normal")
            if data['1title_font']['font_ls_lbl_txt'] == "倍":
                font_title_ls_ent1.config(state="disabled")
            else:
                font_title_ls_ent1.config(state="normal")
            if data['2title_font']['font_ls_lbl_txt'] == "倍":
                font_title_ls_ent2.config(state="disabled")
            else:
                font_title_ls_ent2.config(state="normal")
            if data['3title_font']['font_ls_lbl_txt'] == "倍":
                font_title_ls_ent3.config(state="disabled")
            else:
                font_title_ls_ent3.config(state="normal")
            if data['4title_font']['font_ls_lbl_txt'] == "倍":
                font_title_ls_ent4.config(state="disabled")
            else:
                font_title_ls_ent4.config(state="normal")
            if data['mb_font']['font_ls_lbl_txt'] == "倍":
                font_mb_ls_ent.config(state="disabled")
            else:
                font_mb_ls_ent.config(state="normal")
            pgp_almt_frm_combox.set(data['main']['pgp_almt'])
            single_crl_radio_value.set(data['main']['single_crl_value'])
            time_radio_value.set(data['main']['time_ipt'])
            page_radio_value.set(data['main']['page_ipt'])
            img_radio_value.set(data['main']['img_ipt'])
            pgp_margin_t_vlu.set(data['margin']['t_value']), pgp_margin_b_vlu.set(data['margin']['b_value']), pgp_margin_l_vlu.set(data['margin']['l_value']), pgp_margin_r_vlu.set(data['margin']['r_value'])
            if self.is_auto_import:
                writeHistory("配置自动导入成功！")
            else:
                writeHistory("配置导入成功！")
        else:
            writeHistory("配置文件内容为空！")
            messagebox.showinfo("提示", "配置文件内容为空！")


class eventBottom():
    """窗口底部交互事件"""
    def on_enter(event):
        # ttk 控件需要配置样式或使用 configure
        widget = event.widget
        if isinstance(widget, ttk.Label):
            widget.configure(cursor="hand2", background="#DDDDDD")

    def on_leave(event):
        widget = event.widget
        if isinstance(widget, ttk.Label):
            widget.configure(cursor="", background="#F0F0F0")

    def toMail(event):
        """打开邮箱"""
        webopen("mailto:3038693133@qq.com")

    def toFeedback():
        """打开邮箱"""
        webopen("http://shzsyyey.mikecrm.com/tZ1iuRy")

    def wxTk(event):
        wx_tk = Toplevel(tk)
        original_image = Image.open(wxgzh_path)
        wx_tk.geometry(f"{original_image.width}x460+0+0")
        wx_tk.iconbitmap(icon_path)
        wx_tk.title("微信公众号：晨小明工作室")
        wx_title = ttk.Label(wx_tk, text="微信扫一扫关注公众号", font=("微软雅黑", 14))
        wx_title.grid(row=0, column=0, padx=2, pady=2)
        # 创建Canvas
        cv = Canvas(wx_tk, width=original_image.width, height=original_image.height + 30, highlightthickness=0)
        cv.grid(row=1, column=0, padx=2, pady=0)
        # 加载图片
        time_icon = original_image.resize((round(original_image.width / 1), round(original_image.height / 1)))  # 缩放图片到指定大小
        time_icon_new = ImageTk.PhotoImage(time_icon)
        cv.create_image(0, 0, image=time_icon_new, anchor="nw")
        wx_tk.mainloop()


class CreateFrame():
    def __init__(self, frm, title_txt, row, col, last_txt):
        self.frm = frm
        self.title_txt = title_txt
        self.row = row + 1
        self.col = col
        self.last_txt = last_txt

    def cFontFrame(self):
        """字体标题"""

        def cIndentSpacingFrame(self, col, l_txt="磅"):
            """左侧缩进、右侧缩进、段前、段后布局通用"""
            spacing_frm = ttk.Frame(self.frm)
            spacing_frm.grid(row=self.row, column=col, sticky="n")
            spacing_vlu = StringVar()  # 创建一个StringVar变量来存储数值
            spacing_vlu.set("0")  # 初始值设置为0
            spacing_spinbox = ttk.Spinbox(spacing_frm, from_=0, to=100, increment=0.1, textvariable=spacing_vlu, width=5, font=("Ya Hei", 10), wrap=True)
            spacing_spinbox.grid(row=self.row, column=1, padx=(10, 2), pady=5)
            ttk.Label(spacing_frm, text=l_txt, font=("Ya Hei", 10)).grid(row=self.row, column=2, padx=(2, 10), pady=5)  # 磅文本
            return spacing_vlu
        font_label = ttk.Label(self.frm, font=("Ya Hei", 10, "bold"), text=self.title_txt)
        font_label.grid(row=self.row, column=0, padx=2, pady=5, sticky="e")
        # 字体选择下拉框
        font_name_frm = ttk.Frame(self.frm)
        font_name_frm.grid(row=self.row, column=1, sticky="n")
        font_name_frm_combox = ttk.Combobox(font_name_frm, width=22, font=("Ya Hei", 10), name=self.title_txt, state="readonly")  # 字体下拉框盒子
        font_name_frm_combox.grid(row=self.row, column=1, padx=10, pady=5)
        font_name_frm_combox['values'] = sorted(FONTS)
        font_name_frm_combox.current(sorted(FONTS).index("宋体"))
        # 字号选择下拉框
        font_size_frm = ttk.Frame(self.frm)
        font_size_frm.grid(row=self.row, column=2, sticky="n")
        font_size_frm_combox = ttk.Combobox(font_size_frm, width=4, font=("Ya Hei", 10))  # 字号下拉框盒子
        font_size_frm_combox.grid(row=self.row, column=1, padx=10, pady=5)
        font_size_frm_combox['values'] = [i for i in FONTSIZEDICT.keys()]
        font_size_frm_combox.current(0)
        # 加粗，用复选框
        font_bold_frm = ttk.Frame(self.frm)
        font_bold_frm.grid(row=self.row, column=3, sticky="n")
        font_bold_frm_vlu = IntVar()  # 创建一个BooleanVar变量来存储数值
        font_bold_frm_vlu.set(0)  # 初始值设置为0
        font_bold_frm_check = ttk.Checkbutton(font_bold_frm, text="加粗", variable=font_bold_frm_vlu)  # 加粗复选框
        font_bold_frm_check.grid(row=self.row, column=1, padx=10, pady=5)

        if self.title_txt != "数字英文":
            # 左侧缩进
            indent_left_vlu = cIndentSpacingFrame(self, 4)
            # 右侧缩进
            indent_right_vlu = cIndentSpacingFrame(self, 5)
            # 首行缩进
            first_line_vlu = cIndentSpacingFrame(self, 6, "字符")
            # 段前
            spacing_b_vlu = cIndentSpacingFrame(self, 7)
            # 段后
            spacing_a_vlu = cIndentSpacingFrame(self, 8)
            # 行距选择下拉框
            font_ls_frm = ttk.Frame(self.frm)
            font_ls_frm.grid(row=self.row, column=9, sticky="n")
            font_ls_frm_vlu = StringVar()  # 创建一个StringVar变量来存储数值
            font_ls_frm_combox = ttk.Combobox(font_ls_frm, width=6, font=("Ya Hei", 10), textvariable=font_ls_frm_vlu, state="readonly")  # 行距下拉框盒子
            font_ls_frm_combox.grid(row=self.row, column=1, padx=(10, 2), pady=5)
            font_ls_frm_combox['values'] = ("单倍", "1.5倍", "2倍", "最小值", "固定值", "多倍")
            font_ls_frm_combox.current(0)
            font_ls_vlu = StringVar()  # 创建一个StringVar变量来存储数值
            font_ls_vlu.set("1")  # 初始值设置为0
            font_ls_ent = ttk.Entry(font_ls_frm, width=4, font=("Ya Hei", 10), textvariable=font_ls_vlu, state="disabled")  # 输入框
            font_ls_ent.grid(row=self.row, column=2, padx=2, pady=5)
            font_ls_frm_lbl_b = ttk.Label(font_ls_frm, text="倍", font=("Ya Hei", 10))  # 磅文本
            font_ls_frm_lbl_b.grid(row=self.row, column=3, padx=(2, 10), pady=5)
            return font_name_frm_combox, font_size_frm_combox, font_ls_frm_combox, font_ls_ent, font_ls_frm_lbl_b, font_ls_frm_vlu, font_ls_vlu, spacing_b_vlu, spacing_a_vlu, indent_left_vlu, indent_right_vlu, first_line_vlu, font_bold_frm_vlu
        return font_name_frm_combox, font_size_frm_combox, "", "", "", "", "", "", "", "", "", "", font_bold_frm_vlu

    def cRadioFrame(self):
        """是否"""
        crl_label = ttk.Label(self.frm, font=("Ya Hei", 10, "bold"), text=self.title_txt)
        crl_label.grid(row=self.row, column=0, padx=0, pady=0, sticky="e")
        crl_radio_value = StringVar()
        crl_radio1 = ttk.Radiobutton(self.frm, text="是", variable=crl_radio_value, value="1")
        crl_radio1.grid(row=self.row, column=1, padx=2, pady=0)
        crl_radio2 = ttk.Radiobutton(self.frm, text="否", variable=crl_radio_value, value="0")
        crl_radio2.grid(row=self.row, column=2, padx=5, pady=0)
        return crl_radio_value, crl_radio1, crl_radio2

    def cMarginFrame(self):
        """页边距"""
        if self.col > 0:
            self.col = self.col + self.col * 2
        pgp_margin_frm_lbl = ttk.Label(self.frm, text=self.title_txt, font=("Ya Hei", 10, "bold"))  # 文本
        pgp_margin_frm_lbl.grid(row=self.row, column=self.col, padx=(5, 2), pady=9, sticky="e")
        pgp_margin_vlu = StringVar()  # 创建一个IntVar变量来存储数值
        pgp_margin_spinbox = ttk.Spinbox(self.frm, from_=0, to=100, increment=0.01, textvariable=pgp_margin_vlu, width=5, font=("Ya Hei", 10), wrap=True)
        pgp_margin_spinbox.grid(row=self.row, column=self.col + 1, padx=2, pady=9)
        pgp_margin_lbl_b = ttk.Label(self.frm, text="cm", font=("Ya Hei", 10))  # 磅文本
        pgp_margin_lbl_b.grid(row=self.row, column=self.col + 2, padx=(2, 5), pady=9)
        return pgp_margin_spinbox, pgp_margin_vlu


class LogicalEvents():
    """逻辑处理事件"""
    def fontTitleLsFrmCombox(ls_ent, ls_frm_lbl_b, ls_frm_vlu):
        """标题行距选择事件"""
        # print(ls_frm_vlu.get())
        if ls_frm_vlu.get() == "单倍":
            ls_ent.config(state='normal')
            ls_ent.delete(0, END)  # 删除所有文本
            ls_ent.insert(0, "1")  # 设置新的默认值
            ls_ent.config(state='disabled')
            ls_frm_lbl_b.config(text="倍")
        elif ls_frm_vlu.get() == "1.5倍":
            ls_ent.config(state='normal')
            ls_ent.delete(0, END)  # 删除所有文本
            ls_ent.insert(0, "1.5")  # 设置新的默认值
            ls_ent.config(state='disabled')
            ls_frm_lbl_b.config(text="倍")
        elif ls_frm_vlu.get() == "2倍":
            ls_ent.config(state='normal')
            ls_ent.delete(0, END)  # 删除所有文本
            ls_ent.insert(0, "2")  # 设置新的默认值
            ls_ent.config(state='disabled')
            ls_frm_lbl_b.config(text="倍")
        elif ls_frm_vlu.get() == "最小值":
            ls_ent.config(state='normal')
            ls_ent.delete(0, END)  # 删除所有文本
            ls_ent.insert(0, "12")  # 设置新的默认值
            ls_ent.config(state='normal')
            ls_frm_lbl_b.config(text="磅")
        elif ls_frm_vlu.get() == "固定值":
            ls_ent.config(state='normal')
            if float(ls_ent.get()) < 12:
                ls_ent.delete(0, END)  # 删除所有文本
                ls_ent.insert(0, "12")  # 设置新的默认值
            ls_ent.config(state='normal')
            ls_frm_lbl_b.config(text="磅")
        elif ls_frm_vlu.get() == "多倍":
            ls_ent.config(state='normal')
            ls_ent.delete(0, END)  # 删除所有文本
            ls_ent.insert(0, "1")  # 设置新的默认值
            ls_frm_lbl_b.config(text="倍")
        else:
            ls_ent.config(state='disabled')
            ls_frm_lbl_b.config(text="磅")

    def checkSpinboxValue(data):
        """检查spinbox的值是否合法"""
        t_f = []
        dict_list = [data["title_font"], data["1title_font"], data["2title_font"], data["3title_font"], data["4title_font"], data["mb_font"], data["num_font"]]
        key_name = ["font_ls_vlu", "font_b_s_vlu", "font_a_s_vlu", "font_l_idt_vlu", "font_r_idt_vlu", "font_f_line_vlu"]
        for dict_ in dict_list:
            idx = 0
            for key, value in dict_.items():
                if key != "font_name" and key != "font_size" and key != "font_ls" and key != "font_ls_lbl_txt" and key != "font_bold":
                    try:
                        float(value)
                        t_f.append(True)
                    except:
                        t_f.append(False)
                        writeHistory(f"控件名称：{dict_list.index(dict_)}-{key_name[idx]}，错误值：{value}，请重新输入有效数字！")
                        # messagebox.showerror("错误", f"控件名称：{key}，错误值：{value}，请重新输入有效数字！")
                    idx += 1
        return t_f

    def importIni(auto_import_ini_vlu):
        """导入ini文件"""
        if auto_import_ini_vlu == "":
            return
        cfg_path = path.join(path.dirname(__file__), "fixWord_config.ini")
        config = ConfigParser()
        if not path.isfile(cfg_path):
            config.add_section("main")
        else:
            config.read(cfg_path, encoding="utf-8")
            if config.sections() == []:
                config.add_section("main")
        if auto_import_ini_vlu == "0":
            config.set("main", "auto_import", str(0))
            with open(cfg_path, "w", encoding="utf-8") as f:
                config.write(f)
            data["main"]["auto_import"] = 0
            writeHistory("自动导入已取消！下次打开程序生效！")
        else:
            config.set("main", "auto_import", str(1))
            with open(cfg_path, "w", encoding="utf-8") as f:
                config.write(f)
            data["main"]["auto_import"] = 1
            writeHistory("已开启自动导入！下次打开程序将自动导入默认配置！")


class SystemEvents():
    """系统级事件"""
    def getSysFonts(font_name_frm_combox):
        """获取系统字体"""
        ft = font_name_frm_combox.get()
        ft_name = font_name_frm_combox.winfo_name()
        if ft not in FONTS:
            messagebox.showerror("警告", f"控件名称：{ft_name}, 系统没有 {ft} 的字体！\n已重置为【宋体】或安装相应字体后重试。")  # 改为ask方法
            ft = "宋体"
            font_name_frm_combox.set(ft)
        else:
            pass
            # print(f"··>提示<·· 控件名称：{ft_name}，字体名称：{ft}，检查成功！")
        return ft

    def getUserInput():
        """获取用户输入"""
        global data
        input_path = path_entry.get().replace("/", "\\")
        output_path = input_path + "\output"
        # 获取数值
        ini_impt_ipt = auto_import_ini_vlu.get() or "0"
        font_title_name, font_title_size, font_title_ls, font_ls_value, font_ls_lbl_txt, font_space_b_value, font_space_a_value, font_l_idt_value, font_r_idt_value, font_f_line_value, font_bold_vlu = SystemEvents.getSysFonts(font_title_name_frm_combox), font_title_size_frm_combox.get(
        ), font_title_ls_frm_combox.get(), font_ls_vlu.get(), font_ls_frm_lbl_b.cget("text"), font_spacing_b_vlu.get(), font_spacing_a_vlu.get(), font_indent_left_vlu.get(), font_indent_right_vlu.get(), font_first_line_vlu.get(), font_bold_frm_vlu.get()

        font_title_name1, font_title_size1, font_title_ls1, font_ls_value1, font_ls_lbl_txt1, font_space_b_value1, font_space_a_value1, font_l_idt_value1, font_r_idt_value1, font_f_line_value1, font_bold_vlu1 = SystemEvents.getSysFonts(font_title_name_frm1_combox), font_title_size_frm1_combox.get(
        ), font_title_ls_frm1_combox.get(), font_ls_vlu1.get(), font_ls_frm_lbl_b1.cget("text"), font_spacing_b_vlu1.get(), font_spacing_a_vlu1.get(), font_indent_left_vlu1.get(), font_indent_right_vlu1.get(), font_first_line_vlu1.get(), font_bold_frm_vlu1.get()

        font_title_name2, font_title_size2, font_title_ls2, font_ls_value2, font_ls_lbl_txt2, font_space_b_value2, font_space_a_value2, font_l_idt_value2, font_r_idt_value2, font_f_line_value2, font_bold_vlu2 = SystemEvents.getSysFonts(font_title_name_frm2_combox), font_title_size_frm2_combox.get(
        ), font_title_ls_frm2_combox.get(), font_ls_vlu2.get(), font_ls_frm_lbl_b2.cget("text"), font_spacing_b_vlu2.get(), font_spacing_a_vlu2.get(), font_indent_left_vlu2.get(), font_indent_right_vlu2.get(), font_first_line_vlu2.get(), font_bold_frm_vlu2.get()

        font_title_name3, font_title_size3, font_title_ls3, font_ls_value3, font_ls_lbl_txt3, font_space_b_value3, font_space_a_value3, font_l_idt_value3, font_r_idt_value3, font_f_line_value3, font_bold_vlu3 = SystemEvents.getSysFonts(font_title_name_frm3_combox), font_title_size_frm3_combox.get(
        ), font_title_ls_frm3_combox.get(), font_ls_vlu3.get(), font_ls_frm_lbl_b3.cget("text"), font_spacing_b_vlu3.get(), font_spacing_a_vlu3.get(), font_indent_left_vlu3.get(), font_indent_right_vlu3.get(), font_first_line_vlu3.get(), font_bold_frm_vlu3.get()

        font_title_name4, font_title_size4, font_title_ls4, font_ls_value4, font_ls_lbl_txt4, font_space_b_value4, font_space_a_value4, font_l_idt_value4, font_r_idt_value4, font_f_line_value4, font_bold_vlu4 = SystemEvents.getSysFonts(font_title_name_frm4_combox), font_title_size_frm4_combox.get(
        ), font_title_ls_frm4_combox.get(), font_ls_vlu4.get(), font_ls_frm_lbl_b4.cget("text"), font_spacing_b_vlu4.get(), font_spacing_a_vlu4.get(), font_indent_left_vlu4.get(), font_indent_right_vlu4.get(), font_first_line_vlu4.get(), font_bold_frm_vlu4.get()

        font_mb_name, font_mb_size, font_mb_ls, font_mb_ls_value, font_mb_ls_txt, font_mb_space_b_value, font_mb_space_a_value, font_mb_l_idt_value, font_mb_r_idt_value, font_mb_f_line_value, font_mb_bold_vlu = SystemEvents.getSysFonts(font_mb_name_frm_combox), font_mb_size_frm_combox.get(
        ), font_mb_ls_frm_combox.get(), font_mb_ls_vlu.get(), font_mb_ls_frm_lbl_b.cget("text"), font_mb_spacing_b_vlu.get(), font_mb_spacing_a_vlu.get(), font_mb_indent_left_vlu.get(), font_mb_indent_right_vlu.get(), font_mb_first_line_vlu.get(), font_mb_bold_frm_vlu.get()

        font_num_name, font_num_size, font_num_bold_vlu = SystemEvents.getSysFonts(font_num_name_frm_combox), font_num_size_frm_combox.get(), font_num_bold_frm_vlu.get()

        pgp_almt = pgp_almt_frm_combox.get()
        single_crl_value = single_crl_radio_value.get()
        pgp_margin_t_value, pgp_margin_b_value, pgp_margin_l_value, pgp_margin_r_value = pgp_margin_t_vlu.get(), pgp_margin_b_vlu.get(), pgp_margin_l_vlu.get(), pgp_margin_r_vlu.get()
        time_ipt = time_radio_value.get()
        page_ipt = page_radio_value.get()
        img_ipt = img_radio_value.get()
        data = {
            "title_font": {
                "font_name": font_title_name,
                "font_size": font_title_size,
                "font_bold": font_bold_vlu,
                "font_ls": font_title_ls,
                "font_ls_vlu": font_ls_value,
                "font_ls_lbl_txt": font_ls_lbl_txt,
                "font_b_s_vlu": font_space_b_value,
                "font_a_s_vlu": font_space_a_value,
                "font_l_idt_vlu": font_l_idt_value,
                "font_r_idt_vlu": font_r_idt_value,
                "font_f_line_vlu": font_f_line_value
            },
            "1title_font": {
                "font_name": font_title_name1,
                "font_size": font_title_size1,
                "font_bold": font_bold_vlu1,
                "font_ls": font_title_ls1,
                "font_ls_vlu": font_ls_value1,
                "font_ls_lbl_txt": font_ls_lbl_txt1,
                "font_b_s_vlu": font_space_b_value1,
                "font_a_s_vlu": font_space_a_value1,
                "font_l_idt_vlu": font_l_idt_value1,
                "font_r_idt_vlu": font_r_idt_value1,
                "font_f_line_vlu": font_f_line_value1
            },
            "2title_font": {
                "font_name": font_title_name2,
                "font_size": font_title_size2,
                "font_bold": font_bold_vlu2,
                "font_ls": font_title_ls2,
                "font_ls_vlu": font_ls_value2,
                "font_ls_lbl_txt": font_ls_lbl_txt2,
                "font_b_s_vlu": font_space_b_value2,
                "font_a_s_vlu": font_space_a_value2,
                "font_l_idt_vlu": font_l_idt_value2,
                "font_r_idt_vlu": font_r_idt_value2,
                "font_f_line_vlu": font_f_line_value2
            },
            "3title_font": {
                "font_name": font_title_name3,
                "font_size": font_title_size3,
                "font_bold": font_bold_vlu3,
                "font_ls": font_title_ls3,
                "font_ls_vlu": font_ls_value3,
                "font_ls_lbl_txt": font_ls_lbl_txt3,
                "font_b_s_vlu": font_space_b_value3,
                "font_a_s_vlu": font_space_a_value3,
                "font_l_idt_vlu": font_l_idt_value3,
                "font_r_idt_vlu": font_r_idt_value3,
                "font_f_line_vlu": font_f_line_value3
            },
            "4title_font": {
                "font_name": font_title_name4,
                "font_size": font_title_size4,
                "font_bold": font_bold_vlu4,
                "font_ls": font_title_ls4,
                "font_ls_vlu": font_ls_value4,
                "font_ls_lbl_txt": font_ls_lbl_txt4,
                "font_b_s_vlu": font_space_b_value4,
                "font_a_s_vlu": font_space_a_value4,
                "font_l_idt_vlu": font_l_idt_value4,
                "font_r_idt_vlu": font_r_idt_value4,
                "font_f_line_vlu": font_f_line_value4
            },
            "mb_font": {
                "font_name": font_mb_name,
                "font_size": font_mb_size,
                "font_bold": font_mb_bold_vlu,
                "font_ls": font_mb_ls,
                "font_ls_vlu": font_mb_ls_value,
                "font_ls_lbl_txt": font_mb_ls_txt,
                "font_b_s_vlu": font_mb_space_b_value,
                "font_a_s_vlu": font_mb_space_a_value,
                "font_l_idt_vlu": font_mb_l_idt_value,
                "font_r_idt_vlu": font_mb_r_idt_value,
                "font_f_line_vlu": font_mb_f_line_value
            },
            "num_font": {
                "font_name": font_num_name,
                "font_size": font_num_size,
                "font_bold": font_num_bold_vlu,
                "font_ls": font_mb_ls,
                "font_ls_vlu": font_mb_ls_value,
                "font_ls_lbl_txt": font_mb_ls_txt,
                "font_b_s_vlu": font_mb_space_b_value,
                "font_a_s_vlu": font_mb_space_a_value,
                "font_l_idt_vlu": font_mb_l_idt_value,
                "font_r_idt_vlu": font_mb_r_idt_value,
                "font_f_line_vlu": font_mb_f_line_value
            },
            "margin": {
                "t_value": pgp_margin_t_value,
                "b_value": pgp_margin_b_value,
                "l_value": pgp_margin_l_value,
                "r_value": pgp_margin_r_value,
            },
            "main": {
                "pgp_almt": pgp_almt,
                "single_crl_value": single_crl_value,
                "time_ipt": time_ipt,
                "page_ipt": page_ipt,
                "img_ipt": img_ipt,
                "input_path": input_path,
                "output_path": output_path,
                "auto_import": ini_impt_ipt}
        }
        # print(data)
        return data

    def isAutoImport():
        """判断是否自动导入 """
        cfg_path = path.join(path.dirname(__file__), "fixWord_config.ini")
        config = ConfigParser()
        if not path.isfile(cfg_path):
            return False
        config.read(cfg_path, encoding="utf-8")
        if config.sections() == []:
            return False
        try:
            config.getboolean("main", "auto_import")
            return config.getboolean("main", "auto_import")
        except NoOptionError:
            return False


class LogEvents():
    """操作日志内交互事件"""
    def create_popup_menu(event):
        """创建右键菜单"""
        # 获取当前选中的条目索引和内容
        selected = play_history_frm_listbox.curselection()
        if selected:
            # 创建一个菜单
            popup_menu = Menu(play_history_frm, tearoff=0, font=("Ya Hei", 10))
            # 添加菜单项
            popup_menu.add_command(label="打开文件", command=lambda: LogEvents.open_folder(1))
            popup_menu.add_command(label="复制路径", command=lambda: LogEvents.copy_selected(play_history_frm_listbox))
            popup_menu.add_command(label="在文件夹中显示", command=lambda: LogEvents.open_folder(2))
            # popup_menu.add_separator()  # 添加分隔线
            # popup_menu.add_command(label="退出", command=tk.quit)  # 添加退出命令（可选）
            # 显示菜单
            popup_menu.tk_popup(event.x_root, event.y_root)

    def open_folder(type):
        # 获取当前选中的条目索引和内容
        selected_index = play_history_frm_listbox.curselection()[0]  # 获取当前选中项的索引
        selected_folder = play_history_frm_listbox.get(selected_index)  # 获取当前选中项的内容
        selected_ = selected_folder.split("  ")[-1].split("* ")[-1]
        if path.exists(selected_):
            if type == 1:  # 打开文件
                # 使用系统默认的文件浏览器打开文件夹
                startfile(selected_)  # Windows系统使用此方法
            elif type == 2:    # 打开文件夹
                folder_path = path.dirname(selected_)
                # 使用系统默认的文件浏览器打开文件夹
                startfile(folder_path)  # Windows系统使用此方法
        else:
            messagebox.showwarning("警告", "请指向正确路径！")

    def copy_selected(listbox):
        # 获取选中的项
        # 获取当前选中的条目索引和内容
        selected_index = listbox.curselection()[0]  # 获取当前选中项的索引
        selected_folder = listbox.get(selected_index)  # 获取当前选中项的内容
        selected_ = selected_folder.split("  ")[-1].split("* ")[-1]
        if path.exists(selected_):
            # 这里可以添加复制到剪贴板的代码，例如使用tkinter的clipboard模块
            if tk.clipboard_get():
                tk.clipboard_clear()  # 清空剪贴板
            tk.clipboard_append(selected_)
            messagebox.showinfo("提示", "已复制到剪贴板！\n使用 【Ctrl+V】 粘贴即可！")
        else:
            messagebox.showwarning("警告", "未检测到有效路径！")


def aboutTk():
    about_tk = Toplevel(tk)
    about_tk.title("关于")
    about_tk.iconbitmap(icon_path)
    about_tk.geometry("300x240")
    about_tk.resizable(False, False)
    # 刷新窗口参数
    about_tk.update_idletasks()
    # 计算窗口居中时左上角的坐标
    x = (about_tk.winfo_screenwidth() - about_tk.winfo_width()) // 2
    y = (about_tk.winfo_screenheight() - about_tk.winfo_height()) // 2
    about_tk.geometry(f"+{x}+{y-50}")
    about_tk.focus_set()
    about_tk.protocol("WM_DELETE_WINDOW", lambda: about_tk.destroy())
    about_tk_lbl = ttk.Label(about_tk, text="关于", font=("Ya Hei", 15, "bold"))
    about_tk_lbl.pack(pady=(10, 0))
    # 底部信息
    # 底部文字
    bottom_frm = ttk.Frame(about_tk)
    bottom_frm.pack()
    # 晨小明工作室
    cxm_frm = ttk.Frame(bottom_frm)
    cxm_frm.pack()
    global cxm_image_new  # 声明为全局变量，防止图片被垃圾回收
    original_image = Image.open(cxm_path)
    resized_image = original_image.resize((round(original_image.width / 21), round(original_image.height / 21)))  # 缩放图片到指定大小
    cxm_image_new = ImageTk.PhotoImage(resized_image)
    cv_cxm = Canvas(cxm_frm, width=cxm_image_new.width(), height=cxm_image_new.height(), highlightthickness=0)
    cv_cxm.create_image(0, 0, image=cxm_image_new, anchor="nw")
    cv_cxm.grid(row=0, column=0, pady=(20, 10))
    bottom_info_frm = ttk.Frame(bottom_frm)
    bottom_info_frm.pack()
    bottom_label_a = ttk.Label(bottom_info_frm, text="作者：晨小明")
    bottom_label_a.grid(row=1, column=0, padx=5, pady=5)
    bottom_label_v = ttk.Label(bottom_info_frm, text=f"版本：{VERSION}")
    bottom_label_v.grid(row=2, column=0, padx=5, pady=5)
    bottom_label_t = ttk.Label(bottom_info_frm, text=F"更新时间：{UPDATETIME}")
    bottom_label_t.grid(row=3, column=0, padx=5, pady=5)
    bottom_label_w = ttk.Label(bottom_info_frm, text="微信公众号：晨小明工作室（CXM-Studio）")
    bottom_label_w.grid(row=4, column=0, padx=5, pady=5)
    bottom_label_w.bind("<Enter>", eventBottom.on_enter)
    bottom_label_w.bind("<Leave>", eventBottom.on_leave)
    bottom_label_w.bind("<Button-1>", eventBottom.wxTk)
    bottom_label_c = ttk.Label(bottom_info_frm, text="联系作者：3038693133@qq.com")
    bottom_label_c.grid(row=5, column=0, padx=5, pady=5)
    bottom_label_c.bind("<Enter>", eventBottom.on_enter)
    bottom_label_c.bind("<Leave>", eventBottom.on_leave)
    bottom_label_c.bind("<Button-1>", eventBottom.toMail)


def writeHistory(text=""):
    """写入历史记录"""
    time_stamp = datetime.now().strftime("%m-%d %H:%M:%S.%f")[:-3]
    output_txt = time_stamp + "  " + text
    play_history_frm_listbox.insert(END, output_txt)
    play_history_frm_listbox.update()
    print(f"··>提示<·· {output_txt}")
    # 设置滚动条位置到最大值，即拖动到最底部
    play_history_frm_listbox.yview_moveto(1)


def writeLog(log=""):
    """写入日志"""
    now = datetime.now()
    time_stamp = now.strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
    error_log_path = getcwd() + f"\\error_log.txt"
    if not path.isfile(error_log_path):
        open(error_log_path, 'w', encoding="utf-8").close()
    with open(error_log_path, "a", encoding="utf-8") as f:
        f.write(time_stamp + "    " + log + "\n")


def done():
    """处理完成"""
    merge_button.config(state="normal", cursor="", text="开始处理")
    reset_button.config(state="normal")


def reSet():
    """重置"""
    global data
    font_title_name_frm_combox.current(sorted(FONTS).index("宋体")), font_title_size_frm_combox.current(0), font_title_ls_frm_combox.current(0), font_ls_vlu.set("1"), font_ls_frm_lbl_b.config(text="倍"), font_title_ls_ent.config(
        state="disabled"), font_spacing_b_vlu.set(0), font_spacing_a_vlu.set(0), font_indent_left_vlu.set(0), font_indent_right_vlu.set(0), font_first_line_vlu.set(0), font_bold_frm_vlu.set(0)

    font_title_name_frm1_combox.current(sorted(FONTS).index("宋体")), font_title_size_frm1_combox.current(0), font_title_ls_frm1_combox.current(0), font_ls_vlu1.set("1"), font_ls_frm_lbl_b1.config(
        text="倍"), font_title_ls_ent1.config(state="disabled"), font_spacing_b_vlu1.set(0), font_spacing_a_vlu1.set(0), font_indent_left_vlu1.set(0), font_indent_right_vlu1.set(0), font_first_line_vlu1.set(0), font_bold_frm_vlu1.set(0)

    font_title_name_frm2_combox.current(sorted(FONTS).index("宋体")), font_title_size_frm2_combox.current(0), font_title_ls_frm2_combox.current(0), font_ls_vlu2.set("1"), font_ls_frm_lbl_b2.config(
        text="倍"), font_title_ls_ent2.config(state="disabled"), font_spacing_b_vlu2.set(0), font_spacing_a_vlu2.set(0), font_indent_left_vlu2.set(0), font_indent_right_vlu2.set(0), font_first_line_vlu2.set(0), font_bold_frm_vlu2.set(0)

    font_title_name_frm3_combox.current(sorted(FONTS).index("宋体")), font_title_size_frm3_combox.current(0), font_title_ls_frm3_combox.current(0), font_ls_vlu3.set("1"), font_ls_frm_lbl_b3.config(
        text="倍"), font_title_ls_ent3.config(state="disabled"), font_spacing_b_vlu3.set(0), font_spacing_a_vlu3.set(0), font_indent_left_vlu3.set(0), font_indent_right_vlu3.set(0), font_first_line_vlu3.set(0), font_bold_frm_vlu3.set(0)

    font_title_name_frm4_combox.current(sorted(FONTS).index("宋体")), font_title_size_frm4_combox.current(0), font_title_ls_frm4_combox.current(0), font_ls_vlu4.set("1"), font_ls_frm_lbl_b4.config(
        text="倍"), font_title_ls_ent4.config(state="disabled"), font_spacing_b_vlu4.set(0), font_spacing_a_vlu4.set(0), font_indent_left_vlu4.set(0), font_indent_right_vlu4.set(0), font_first_line_vlu4.set(0), font_bold_frm_vlu4.set(0)

    font_mb_name_frm_combox.current(sorted(FONTS).index("宋体")), font_mb_size_frm_combox.current(0), font_mb_ls_frm_combox.current(0), font_mb_ls_vlu.set("1"), font_mb_ls_frm_lbl_b.config(
        text="倍"), font_mb_ls_ent.config(state="disabled"), font_mb_spacing_b_vlu.set(0), font_mb_spacing_a_vlu.set(0), font_mb_indent_left_vlu.set(0), font_mb_indent_right_vlu.set(0), font_mb_first_line_vlu.set(0), font_mb_bold_frm_vlu.set(0)

    font_num_name_frm_combox.current(sorted(FONTS).index("宋体")), font_num_size_frm_combox.current(0), font_num_bold_frm_vlu.set(0)
    try:
        if font_title_ls_frm_combox.cget("state").string == "disabled":
            font_title_ls_frm_combox.configure(state="readonly")
            font_title_ls_frm1_combox.configure(state="readonly")
            font_title_ls_frm2_combox.configure(state="readonly")
            font_title_ls_frm3_combox.configure(state="readonly")
            font_mb_ls_frm_combox.configure(state="readonly")
    except:
        pass
    pgp_almt_frm_combox.current(3)
    pgp_margin_t_vlu.set("2.54"), pgp_margin_b_vlu.set("2.54"), pgp_margin_l_vlu.set("3.17"), pgp_margin_r_vlu.set("3.17")
    # 修改为使用 set 方法设置值
    single_crl_radio_value.set("0")
    time_radio_value.set("0")
    page_radio_value.set("0")
    img_radio_value.set("0")
    play_history_frm_listbox.delete(0, END)
    writeHistory("重置成功！")


def main():
    """主函数"""
    try:
        global data
        input_path = path_entry.get().replace("/", "\\")
        if input_path == "":
            messagebox.showinfo("提示", "请选择文件或文件夹路径！")
        else:
            file_type = type_radio_value.get()
            if file_type == "file_path":
                if not path.isfile(input_path):
                    messagebox.showerror("错误", "文件路径错误！")
                    return
            elif file_type == "dir_path":
                if not path.isdir(input_path):
                    messagebox.showerror("错误", "文件夹路径错误！")
                    return
            data = SystemEvents.getUserInput()
            if not all(LogicalEvents.checkSpinboxValue(data)) or not all(LogicalEvents.checkSpinboxValue(data)):
                return
            output_path = data["main"]["output_path"]
            time_ipt = data["main"]["time_ipt"]
            page_ipt = data["main"]["page_ipt"]
            img_ipt = data["main"]["img_ipt"]
            merge_button.config(state="disabled", cursor="wait", text="正在处理")
            reset_button.config(state="disabled")
            merge_button.update_idletasks()
            reset_button.update_idletasks()
            writeHistory("开始处理...")
            if file_type == "dir_path":
                have_docx = 0
                done_list = []
                for file in listdir(input_path):
                    if '~' in file:
                        continue
                    elif file.endswith('.docx'):
                        if not path.isdir(output_path):
                            makedirs(output_path)
                        have_docx += 1
                        file_path = path.join(input_path, file)
                        save_path, is_done = DocxProcessing.fixWord(file_path, file, output_path, time_ipt, page_ipt, img_ipt)
                        if is_done:
                            writeHistory(str(have_docx) + " * " + save_path)
                            done_list.append(file_path)
                        else:
                            writeHistory(f"{path.splitext(file)[0]}.docx 保存失败！文件已打开，请关闭后重试！")
                            messagebox.showerror("错误", f"{path.splitext(file)[0]}.docx 保存失败！\n文件已打开，请关闭后重试！")
                if have_docx == 0:
                    print("··>错误<·· 没有找到.docx文件")
                    messagebox.showinfo("提示", "没有找到.docx文件！")
                else:
                    if len(done_list) == have_docx:
                        messagebox.showinfo("提示", "全部处理完成！\n输出路径：" + output_path)
                    else:
                        messagebox.showinfo("提示", f"处理完成！\n共 {have_docx} 个文件，成功 {len(done_list)} 个，失败 {have_docx - len(done_list)} 个\n输出路径：" + output_path)
            elif file_type == "file_path":
                # 文件名
                file = input_path.split("\\")[-1]
                # 输出路径
                dir_path = input_path.split("\\")
                dir_path.pop()
                result = '\\'.join(str(x) for x in dir_path)
                output_path = result + "\output"
                if not path.isdir(output_path):
                    makedirs(output_path)
                save_path, is_done = DocxProcessing.fixWord(input_path, file, output_path, time_ipt, page_ipt, img_ipt)
                if is_done:
                    writeHistory(save_path)
                    messagebox.showinfo("提示", "处理完成！\n输出路径：" + save_path)
                else:
                    writeHistory(f"{path.splitext(file)[0]}.docx 保存失败！文件已打开，请关闭后重试！")
                    messagebox.showerror("错误", f"{path.splitext(file)[0]}.docx 保存失败！\n文件已打开，请关闭后重试！")
    except Exception as e:
        tb_next = e.__traceback__
        err_out = ""
        while tb_next:
            error_log = f"Function: {tb_next.tb_frame.f_code.co_name}，Line: {tb_next.tb_lineno}"
            writeLog(error_log)
            err_out += f"Filename：{tb_next.tb_frame.f_code.co_filename}" + "，" + error_log + "\n"
            tb_next_ = tb_next
            tb_next = tb_next.tb_next
        writeLog(f"Info: {e}")
        writeHistory(f"程序出错！请截图并联系作者！Filename：{tb_next_.tb_frame.f_code.co_filename}，Function：{tb_next_.tb_frame.f_code.co_name}，Line：{tb_next_.tb_lineno}，Info：{e}")
        messagebox.showerror("错误", f"程序出错！请截图并联系作者！\n{err_out + 'Info: '+ str(e)}")
    finally:
        done()


if __name__ == '__main__':
    VERSION = "v5.4.0.4"
    UPDATETIME = "2026年8月26日"
    """
        !!!!!!!!!!!!
        打包时把此路径改为相对路径，并把图片复制粘贴到打包后的根目录里
        !!!!!!!!!!!!
        pyinstaller -D -w fix_word.py -i static/icon.ico -n fixWord_v5.4.0.4
    """
    icon_path = getcwd() + "\\static\\icon.ico"
    wxgzh_path = getcwd() + "\\static\\wxgzh.jpg"
    cxm_path = getcwd() + "\\static\\cxmstudio-lignt-heng.png"
    # 配置信息start
    # 字号字典
    FONTSIZEDICT = {"初号": 42, "小初": 36, "一号": 26, "小一": 24, "二号": 22, "小二": 18, "三号": 16, "小三": 15, "四号": 14, "小四": 12, "五号": 10.5, "小五": 9, "六号": 7.5, "小六": 6.5, "七号": 5.5, "八号": 5, "5":5, "5.5":5.5, "6.5":6.5, "7.5":7.5, "8":8, "9":9, "10":10, "10.5":10.5, "11":11, "12":12, "14":14, "16":16, "18":18, "20":20, "22":22, "24":24, "26":26, "28":28, "36":36, "48":48, "72":72}
    # 配置信息end
    # tkinter start
    tk = Tk()
    tk.title(f"文档处理工具 {VERSION} - 微信公众号：晨小明工作室")
    tk.iconbitmap(icon_path)
    tk.geometry("1600x750")
    # 获取数值前刷新
    tk.update_idletasks()
    # 计算窗口居中时左上角的坐标
    x = (tk.winfo_screenwidth() - tk.winfo_width()) // 2
    y = (tk.winfo_screenheight() - tk.winfo_height()) // 2
    tk.geometry(f"+{x}+{y-30}")
    # 调整位置后再刷新
    tk.update()
    # tk.attributes("-alpha", 0.8)
    tk.minsize(1486, 690)  # 最小宽高
    FONTS = [font_ for font_ in tkFont.families() if "@" not in font_]
    # 设置ttk样式
    # style.theme_use('clam')
    # 文件路径
    frm_ = ttk.Frame(tk)
    frm_.pack(anchor="center")
    path_lf = ttk.LabelFrame(frm_, text="选择路径", padding=10)
    path_lf.grid(row=0, column=0, padx=2, pady=5)
    type_label = ttk.Label(path_lf, font=("Ya Hei", 10, "bold"), text="请选择输入类型：")
    type_label.grid(row=0, column=0, padx=2, pady=5, sticky="e")
    type_radio_value = StringVar()
    type_radio1 = ttk.Radiobutton(path_lf, text="文件", value="file_path", variable=type_radio_value, command=PathEvents.inputFile)
    type_radio1.grid(row=0, column=1, padx=2, pady=2)
    type_radio2 = ttk.Radiobutton(path_lf, text="文件夹", value="dir_path", variable=type_radio_value, command=PathEvents.inputDir)
    type_radio2.grid(row=0, column=2, padx=2, pady=2)
    type_radio_value.set("file_path")  # 使用 set 方法设置默认值
    path_entry = ttk.Entry(path_lf, width=80, font=("Ya Hei", 12))
    path_entry.grid(row=0, column=3, padx=2, pady=5, sticky="w")
    path_button = ttk.Button(path_lf, text="选择文件", command=PathEvents.inputPath, width=12)
    path_button.grid(row=0, column=4, padx=2, pady=5)
    separator = ttk.Separator(tk, orient='horizontal')
    separator.pack(fill="x", padx=5, pady=5)
    # 主布局
    main_lf = ttk.Frame(tk)
    main_lf.pack(padx=5, pady=5)
    main_frm = ttk.Frame(main_lf)
    main_frm.grid(row=0, column=0, padx=5, pady=5)
    # 字体
    font_lf = ttk.LabelFrame(main_frm, text="设置字体/段落", padding=10)
    font_lf.grid(row=0, column=0, padx=5, pady=5)
    font_frm = ttk.Frame(font_lf)
    font_frm.pack(padx=10, pady=10)
    font_name_label = ttk.Label(font_frm, text="字体名称", font=("Ya Hei", 10, "bold"))
    font_name_label.grid(row=0, column=1, padx=5, pady=5)
    font_size_label = ttk.Label(font_frm, text="字号", font=("Ya Hei", 10, "bold"))
    font_size_label.grid(row=0, column=2, padx=5, pady=5)
    font_bold_label = ttk.Label(font_frm, text="加粗", font=("Ya Hei", 10, "bold"))
    font_bold_label.grid(row=0, column=3, padx=5, pady=5)
    font_indent_left_label = ttk.Label(font_frm, text="左侧缩进", font=("Ya Hei", 10, "bold"))
    font_indent_left_label.grid(row=0, column=4, padx=5, pady=5)
    font_indent_right_label = ttk.Label(font_frm, text="右侧缩进", font=("Ya Hei", 10, "bold"))
    font_indent_right_label.grid(row=0, column=5, padx=5, pady=5)
    font_first_line_indent_label = ttk.Label(font_frm, text="首行缩进", font=("Ya Hei", 10, "bold"))
    font_first_line_indent_label.grid(row=0, column=6, padx=5, pady=5)
    font_spacing_before_label = ttk.Label(font_frm, text="段前", font=("Ya Hei", 10, "bold"))
    font_spacing_before_label.grid(row=0, column=7, padx=5, pady=5)
    font_spacing_after_label = ttk.Label(font_frm, text="段后", font=("Ya Hei", 10, "bold"))
    font_spacing_after_label.grid(row=0, column=8, padx=5, pady=5)
    font_ls_label = ttk.Label(font_frm, text="行距", font=("Ya Hei", 10, "bold"))
    font_ls_label.grid(row=0, column=9, padx=5, pady=5)
    # 标题
    font_title_name_frm_combox, font_title_size_frm_combox, font_title_ls_frm_combox, font_title_ls_ent, font_ls_frm_lbl_b, font_ls_frm_vlu, font_ls_vlu, font_spacing_b_vlu, font_spacing_a_vlu, font_indent_left_vlu, font_indent_right_vlu, font_first_line_vlu, font_bold_frm_vlu = CreateFrame(font_frm, "标题",  0, 0, "磅").cFontFrame()
    font_title_ls_frm_combox.bind("<<ComboboxSelected>>", lambda event: LogicalEvents.fontTitleLsFrmCombox(font_title_ls_ent, font_ls_frm_lbl_b, font_ls_frm_vlu))
    # 一级标题
    font_title_name_frm1_combox, font_title_size_frm1_combox, font_title_ls_frm1_combox, font_title_ls_ent1, font_ls_frm_lbl_b1, font_ls_frm_vlu1, font_ls_vlu1, font_spacing_b_vlu1, font_spacing_a_vlu1, font_indent_left_vlu1, font_indent_right_vlu1, font_first_line_vlu1, font_bold_frm_vlu1 = CreateFrame(font_frm, "一级标题", 1, 0, "磅").cFontFrame()
    font_title_ls_frm1_combox.bind("<<ComboboxSelected>>", lambda event: LogicalEvents.fontTitleLsFrmCombox(font_title_ls_ent1, font_ls_frm_lbl_b1, font_ls_frm_vlu1))
    # 二级标题
    font_title_name_frm2_combox, font_title_size_frm2_combox, font_title_ls_frm2_combox, font_title_ls_ent2, font_ls_frm_lbl_b2, font_ls_frm_vlu2, font_ls_vlu2, font_spacing_b_vlu2, font_spacing_a_vlu2, font_indent_left_vlu2, font_indent_right_vlu2, font_first_line_vlu2, font_bold_frm_vlu2 = CreateFrame(font_frm, "二级标题", 2, 0, "磅").cFontFrame()
    font_title_ls_frm2_combox.bind("<<ComboboxSelected>>", lambda event: LogicalEvents.fontTitleLsFrmCombox(font_title_ls_ent2, font_ls_frm_lbl_b2, font_ls_frm_vlu2))
    # 三级标题
    font_title_name_frm3_combox, font_title_size_frm3_combox, font_title_ls_frm3_combox, font_title_ls_ent3, font_ls_frm_lbl_b3, font_ls_frm_vlu3, font_ls_vlu3, font_spacing_b_vlu3, font_spacing_a_vlu3, font_indent_left_vlu3, font_indent_right_vlu3, font_first_line_vlu3, font_bold_frm_vlu3 = CreateFrame(font_frm, "三级标题", 3, 0, "磅").cFontFrame()
    font_title_ls_frm3_combox.bind("<<ComboboxSelected>>", lambda event: LogicalEvents.fontTitleLsFrmCombox(font_title_ls_ent3, font_ls_frm_lbl_b3, font_ls_frm_vlu3))
    # 四级标题
    font_title_name_frm4_combox, font_title_size_frm4_combox, font_title_ls_frm4_combox, font_title_ls_ent4, font_ls_frm_lbl_b4, font_ls_frm_vlu4, font_ls_vlu4, font_spacing_b_vlu4, font_spacing_a_vlu4, font_indent_left_vlu4, font_indent_right_vlu4, font_first_line_vlu4, font_bold_frm_vlu4 = CreateFrame(font_frm, "四级标题", 4, 0, "磅").cFontFrame()
    font_title_ls_frm4_combox.bind("<<ComboboxSelected>>", lambda event: LogicalEvents.fontTitleLsFrmCombox(font_title_ls_ent4, font_ls_frm_lbl_b4, font_ls_frm_vlu4))
    # 正文
    font_mb_name_frm_combox, font_mb_size_frm_combox, font_mb_ls_frm_combox, font_mb_ls_ent, font_mb_ls_frm_lbl_b, font_mb_ls_frm_vlu, font_mb_ls_vlu, font_mb_spacing_b_vlu, font_mb_spacing_a_vlu, font_mb_indent_left_vlu, font_mb_indent_right_vlu, font_mb_first_line_vlu, font_mb_bold_frm_vlu = CreateFrame(font_frm, "   正文", 5, 0, "磅").cFontFrame()
    font_mb_ls_frm_combox.bind("<<ComboboxSelected>>", lambda event: LogicalEvents.fontTitleLsFrmCombox(font_mb_ls_ent, font_mb_ls_frm_lbl_b, font_mb_ls_frm_vlu))
    # 其他
    font_num_name_frm_combox, font_num_size_frm_combox, font_num_ls_frm_combox, font_num_ls_ent, font_num_ls_frm_lbl_b, font_num_ls_frm_vlu, font_num_ls_vlu, font_num_spacing_b_vlu, font_num_spacing_a_vlu, font_num_indent_left_vlu, font_num_indent_right_vlu, font_num_first_line_vlu, font_num_bold_frm_vlu = CreateFrame(
        font_frm, "数字英文", 6, 0, "磅").cFontFrame()
    # 其他段落设置
    _almt = ttk.Frame(main_frm)
    _almt.grid(row=1, column=0, padx=5, pady=17)
    frm_almt = ttk.Frame(main_frm)
    frm_almt.grid(row=2, column=0, padx=5, pady=5)
    pgp_almt_lf = ttk.LabelFrame(frm_almt, text="其他段落设置", padding=10)
    pgp_almt_lf.grid(row=0, column=0, padx=(0, 50), pady=5, ipady=6)
    pgp_almt_frm = ttk.Frame(pgp_almt_lf)
    pgp_almt_frm.grid(row=0, column=0, padx=28, pady=5)
    pgp_almt_frm_lbl = ttk.Label(pgp_almt_frm, text="对齐方式：", font=("Ya Hei", 10, "bold"))
    pgp_almt_frm_lbl.grid(row=0, column=0, padx=2, pady=2)
    pgp_almt_frm_combox = ttk.Combobox(pgp_almt_frm, width=8, font=("Ya Hei", 10), state="readonly")
    pgp_almt_frm_combox.grid(row=0, column=1, padx=2, pady=2)
    pgp_almt_frm_combox['values'] = ("左对齐", "居中", "右对齐", "两端对齐")
    pgp_almt_frm_combox.current(3)
    single_crl_frm = ttk.Frame(pgp_almt_lf)  # 孤行控制
    single_crl_frm.grid(row=1, column=0, padx=5, pady=5)
    single_crl_radio_value, single_crl_radio1, single_crl_radio2 = CreateFrame(single_crl_frm, "孤行控制：", 0, 0, "").cRadioFrame()
    single_crl_radio_value.set("0")  # 使用 set 方法设置默认值
    # 设置页边距
    pgp_margin_lf = ttk.LabelFrame(frm_almt, text="设置页边距", padding=10)
    pgp_margin_lf.grid(row=0, column=2, padx=95, pady=2)
    pgp_margin_frm = ttk.Frame(pgp_margin_lf)
    pgp_margin_frm.grid(row=0, column=0, padx=5, pady=2)
    pgp_margin_t_spb, pgp_margin_t_vlu = CreateFrame(pgp_margin_frm, "上：", 0, 0, "").cMarginFrame()
    pgp_margin_b_spb, pgp_margin_b_vlu = CreateFrame(pgp_margin_frm, "下：", 0, 1, "").cMarginFrame()
    pgp_margin_l_spb, pgp_margin_l_vlu = CreateFrame(pgp_margin_frm, "左：", 1, 0, "").cMarginFrame()
    pgp_margin_r_spb, pgp_margin_r_vlu = CreateFrame(pgp_margin_frm, "右：", 1, 1, "").cMarginFrame()
    pgp_margin_t_vlu.set("2.54"), pgp_margin_b_vlu.set("2.54"), pgp_margin_l_vlu.set("3.17"), pgp_margin_r_vlu.set("3.17")
    # 自定义选项
    if_frm = ttk.LabelFrame(frm_almt, text="自定义选项", padding=10)
    if_frm.grid(row=0, column=3, padx=(50, 0), pady=2)
    info_frm = ttk.Frame(if_frm)
    info_frm.grid(row=0, column=0, padx=5, pady=5)
    time_radio_value, time_radio1, time_radio2 = CreateFrame(info_frm, "添加时间标记：", 0, 0, "").cRadioFrame()
    page_radio_value, page_radio1, page_radio2 = CreateFrame(info_frm, "添加页码：", 1, 0, "").cRadioFrame()
    img_radio_value, img_radio1, img_radio2 = CreateFrame(info_frm, "保存文档中的图片：", 2, 0, "").cRadioFrame()
    time_radio_value.set("0")  # 使用 set 方法设置默认值
    page_radio_value.set("0")   # 使用 set 方法设置默认值
    img_radio_value.set("0")    # 使用 set 方法设置默认值
    # 处理日志 - Listbox 没有 ttk 版本，继续使用 tk.Listbox
    play_history_lfrm = ttk.LabelFrame(main_frm, text="操作日志", padding=10)
    play_history_lfrm.grid(row=0, column=1, padx=5, pady=0, rowspan=3)
    play_history_frm = ttk.Frame(play_history_lfrm)
    play_history_frm.grid(row=0, column=1, padx=(20, 5), pady=5, rowspan=2)
    play_history_frm_listbox = Listbox(play_history_frm, width=50, height=30, font=("Ya Hei", 10), border=1, activestyle="none")
    play_history_frm_listbox.grid(row=1, column=0, padx=(0, 0), pady=(0, 0))
    play_history_scroll_bar_v = ttk.Scrollbar(play_history_frm, orient="vertical", command=play_history_frm_listbox.yview)
    play_history_scroll_bar_v.grid(row=1, column=1, sticky='ns')
    play_history_scroll_bar_h = ttk.Scrollbar(play_history_frm, orient="horizontal", command=play_history_frm_listbox.xview)
    play_history_scroll_bar_h.grid(row=2, column=0, sticky='we')
    play_history_frm_listbox.configure(yscrollcommand=play_history_scroll_bar_v.set, xscrollcommand=play_history_scroll_bar_h.set)
    # 绑定右键点击事件到创建弹出菜单的函数
    play_history_frm_listbox.bind("<Button-3>", LogEvents.create_popup_menu)
    # 绑定双击事件到列表框上
    play_history_frm_listbox.bind("<Double-1>", lambda event: LogEvents.open_folder(1))
    # 分隔线
    separator = ttk.Separator(tk, orient='horizontal')
    separator.pack(fill="x", padx=5, pady=5)
    # 处理按钮
    btn_frm = ttk.Frame(tk)
    btn_frm.pack(pady=6)
    style = ttk.Style()
    style.configure("reset.TButton", foreground="blue")
    reset_button = ttk.Button(btn_frm, text="重  置", style="reset.TButton", command=reSet)
    reset_button.grid(row=0, column=0, padx=5, pady=5)
    label_ = ttk.Label(btn_frm, text=" ")
    label_.grid(row=0, column=1, padx=5, pady=5, sticky="e")
    style.configure("merge.TButton", foreground="green")
    merge_button = ttk.Button(btn_frm, text="开始处理", style="merge.TButton", command=main)
    merge_button.grid(row=0, column=2, padx=5, pady=5)
    # tkinter end
    # 创建菜单 - Menu 没有 ttk 版本，继续使用 tk.Menu
    menu = Menu(tk)
    tk.config(menu=menu)
    # 创建文件菜单
    file_menu = Menu(menu, tearoff=0)
    menu.add_cascade(label="文件", menu=file_menu)
    file_menu.add_command(label="打开文件", command=PathEvents.inputFile)
    file_menu.add_command(label="打开文件夹", command=PathEvents.inputDir)
    file_menu.add_separator()
    file_menu.add_command(label="退出", command=tk.quit)
    # 创建工具菜单
    tool_menu = Menu(menu, tearoff=0)
    menu.add_cascade(label="配置", menu=tool_menu)
    auto_import_ini_vlu = StringVar()
    tool_menu.add_command(label="导入配置", command=InitFile().importConfig)
    tool_menu.add_command(label="配置另存为", command=InitFile().esaveConfig)
    tool_menu.add_command(label="保存当前配置", command=InitFile().saveConfig)
    tool_menu.add_separator()
    auto_import_ini = Menu(tool_menu, tearoff=0)
    auto_import_ini.add_radiobutton(label="打开", variable=auto_import_ini_vlu, value="1", command=lambda: LogicalEvents.importIni(auto_import_ini_vlu.get()))
    auto_import_ini.add_radiobutton(label="关闭", variable=auto_import_ini_vlu, value="0", command=lambda: LogicalEvents.importIni(auto_import_ini_vlu.get()))
    tool_menu.add_cascade(label="自动导入", menu=auto_import_ini)
    # 创建帮助菜单
    help_menu = Menu(menu, tearoff=0)
    menu.add_cascade(label="帮助", menu=help_menu)
    help_menu.add_command(label="反馈", command=eventBottom.toFeedback)
    help_menu.add_command(label="检查更新", command=upGrade)
    help_menu.add_separator()
    help_menu.add_command(label="关于", command=aboutTk)

    data = SystemEvents.getUserInput()
    if SystemEvents.isAutoImport():
        InitFile(True).importConfig()
        auto_import_ini_vlu.set("1")
    else:
        auto_import_ini_vlu.set("0")
    tk.bind("<Return>", lambda event: main())
    # 初始化成功
    writeHistory("初始化成功！")
    tk.mainloop()
    # tkinter end
