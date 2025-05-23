from docx import Document
from docx.shared import Pt, Cm  # 用来设置字体的大小
from docx.oxml.ns import qn  # 设置字体
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT  # 设置对其方式
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from os import listdir, path, makedirs, getcwd
from tkinter import Tk, Entry, Button, Label, filedialog, messagebox, SUNKEN, Radiobutton, Frame, ttk, Listbox, StringVar, END
from time import localtime, strftime

"""
开发作者：晨小明
开发日期：2024/09/22
开发版本：v3.0__release
修改日期：2025/05/06
主要功能：一、支持单文件处理或批量文档处理，输入文件路径或文件夹路径，自动判断。
         二、读取.docx文件并设置格式：
        三、支持添加页码（可选）：4号半角宋体阿拉伯数字，数字左右各加一条4号“一字线”，奇数页在右侧左空一字，偶数页在左侧左空一字
        四、识别文档中的图片并输出（可选）：（注：图片可能会被压缩）
        五、替换功能
            1.符号替换
                将英文状态下的符号替换为中文状态下的相同符号，包含如下：
                "(" --> "（"
                ")" --> "）"
                ")、" --> "）"
                "）、" --> "）"
                "," --> "，"
                ":" --> "："
                ";" --> "；"
                "?" --> "？"
                " " --> ""
            2.其他格式
                数字后有顿号替换为点，如："1、" --> "1."
         六、输出文件名称含时间点，方便标记（可选）
         （注，本程序无法处理图片格式，如果图片独立成段，本程序所用API识别到图片会被默认是空段落，为了防止图片删除，只能放弃处理空段落及图片格式）
更新日志：
【新增】字体常量，便于统一；
【新增】两个版本：学校留存；上交上报；
【新增】当前格式显示；
【优化】其他内容；
【修复】弹窗的路径不准确的情况。
"""


def margin(docx):
    """ 设置页边距 """
    for s in docx.sections:
        s.top_margin = Cm(PAGETOPMARGIN)
        s.bottom_margin = Cm(PAGEBOTTOMMARGIN)
        s.left_margin = Cm(PAGELEFTMARGIN)
        s.right_margin = Cm(PAGERIGHTMARGIN)


def footer(docx):
    """ 设置页脚，添加页码 """
    # print(len(docx.sections))
    def AddFooterNumber(p):
        t1 = p.add_run("— ")
        font = t1.font
        font.name = PAGENUMBERFONT
        font.size = Pt(PAGENUMBERFONTSIZE)  # 14号字体
        t1._element.rPr.rFonts.set(qn("w:eastAsia"), PAGENUMBERFONT)

        run1 = p.add_run('')
        fldChar1 = OxmlElement('w:fldChar')  # creates a new element
        fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        run1._element.append(fldChar1)

        run2 = p.add_run('')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'PAGE'
        font = run2.font
        font.name = PAGENUMBERFONT
        font.size = Pt(PAGENUMBERFONTSIZE)  # 14号字体
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), PAGENUMBERFONT)
        run2._element.append(instrText)

        run3 = p.add_run('')
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._element.append(fldChar2)

        t2 = p.add_run(" —")
        font = t2.font
        font.name = PAGENUMBERFONT
        font.size = Pt(PAGENUMBERFONTSIZE)  # 14号字体
        t2._element.rPr.rFonts.set(qn("w:eastAsia"), PAGENUMBERFONT)

    for s in docx.sections:
        # print(s.footer)
        footer = s.footer  # 获取第一个节的页脚
        footer.is_linked_to_previous = True  # 编号续前一节
        paragraph = footer.paragraphs[0]  # 获取页脚的第一个段落
        paragraphFun("odd_footer", paragraph)
        AddFooterNumber(paragraph)
        even_footer = s.even_page_footer  # 获取第一个节的页脚
        even_footer.is_linked_to_previous = True  # 编号续前一节
        paragraph = even_footer.paragraphs[0]  # 获取页脚的第一个段落
        paragraphFun("even_footer", paragraph)
        AddFooterNumber(paragraph)


def paragraphFun(is_title, p):
    """ 段落函数 """
    if is_title == "title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = Pt(TITLEMARGIN)  # 行距
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.before_spacing = Pt(0)
        p.paragraph_format.after_spacing = Pt(0)
    elif is_title == "odd_footer":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.paragraph_format.right_indent = Pt(14)
        p.paragraph_format.line_spacing = Pt(TEXTMARGIN)
    elif is_title == "even_footer":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.line_spacing = Pt(TEXTMARGIN)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = Pt(TEXTMARGIN)  # 行距
        p.paragraph_format.before_spacing = Pt(0)
        p.paragraph_format.after_spacing = Pt(0)


def isLevel1(p):
    """ 判断是否是 1 级标题 """
    index1_list = ["一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、",
                   "十一、", "十二、", "十三、", "十四、", "十五、", "十六、", "十七、", "十八、", "十九、", "二十、"]
    for i in range(len(index1_list)):
        if p.text.find(index1_list[i]) != -1:
            if '。' in p.text:
                p.text = p.text.replace('。', '')
            if '？' in p.text:
                p.text = p.text.replace('？', '')
            if '：' in p.text:
                p.text = p.text.replace('：', '')
            if '；' in p.text:
                p.text = p.text.replace('；', '')
            return "level1"
        else:
            continue


def isLevel2(p):
    """ 判断是否是 2 级标题 """
    index2_list = ["（一）", "（二）", "（三）", "（四）", "（五）", "（六）", "（七）", "（八）", "（九）", "（十）",
                   "（十一）", "（十二）", "（十三）", "（十四）", "（十五）", "（十六）", "（十七）", "（十八）", "（十九）", "（二十）"]
    for i in range(len(index2_list)):
        if p.text.find(index2_list[i]) != -1:
            if '。' in p.text:
                p.text = p.text.replace('。', '')
            if '？' in p.text:
                p.text = p.text.replace('？', '')
            if '：' in p.text:
                p.text = p.text.replace('：', '')
            if '；' in p.text:
                p.text = p.text.replace('；', '')
            return "level2"
        else:
            continue


def text(is_title, is_level1, is_level2, is_digit, p, i, version_ipt):
    """ 正文函数 """
    if is_title == "title":
        if is_digit == "num_or_let":
            title_num = p.add_run(i)
            title_num.font.name = NUMBERFONT
            if version_ipt == "school":
                title_num.font.size = Pt(FONTSIZEDICT["小二号"])
            else:
                title_num.font.size = Pt(FONTSIZEDICT["二号"])
            title_num.font.bold = False
        else:
            run_title = p.add_run(i)
            run_title.font.name = TITLEFONT
            run_title._element.rPr.rFonts.set(
                qn('w:eastAsia'), TITLEFONT)
            if version_ipt == "school":
                run_title.font.size = Pt(FONTSIZEDICT["小二号"])
            else:
                run_title.font.size = Pt(FONTSIZEDICT["二号"])
            run_title.font.bold = False
    else:
        if is_digit == "num_or_let":
            text_num = p.add_run(i)
            text_num.font.name = NUMBERFONT
            if version_ipt == "school":
                text_num.font.size = Pt(FONTSIZEDICT["四号"])
            else:
                text_num.font.size = Pt(FONTSIZEDICT["三号"])
            text_num.font.bold = False
        elif is_level1 == "level1":
            run_level1 = p.add_run(i)
            run_level1.font.name = LEVEL1FONT
            run_level1._element.rPr.rFonts.set(
                qn('w:eastAsia'), LEVEL1FONT)
            if version_ipt == "school":
                run_level1.font.size = Pt(FONTSIZEDICT["四号"])
            else:
                run_level1.font.size = Pt(FONTSIZEDICT["三号"])
            run_level1.font.bold = False
        elif is_level2 == "level2":
            run_level2 = p.add_run(i)
            run_level2.font.name = LEVEL2FONT
            run_level2._element.rPr.rFonts.set(
                qn('w:eastAsia'), LEVEL2FONT)
            if version_ipt == "school":
                run_level2.font.size = Pt(FONTSIZEDICT["四号"])
            else:
                run_level2.font.size = Pt(FONTSIZEDICT["三号"])
            run_level2.font.bold = False
        else:
            run_content = p.add_run(i)
            run_content.font.name = TEXTFONT
            run_content._element.rPr.rFonts.set(
                qn('w:eastAsia'), TEXTFONT)
            if version_ipt == "school":
                run_content.font.size = Pt(FONTSIZEDICT["四号"])
            else:
                run_content.font.size = Pt(FONTSIZEDICT["三号"])
            run_content.font.bold = False


def replace(p):
    """ 替换函数 """
    # 替换符号
    if ')、' in p.text:
        p.text = p.text.replace(')、', '）')
    if '）、' in p.text:
        p.text = p.text.replace('）、', '）')
    if '(' in p.text:
        p.text = p.text.replace('(', '（')
    if ')' in p.text:
        p.text = p.text.replace(')', '）')
    if ',' in p.text:
        p.text = p.text.replace(',', '，')
    if ':' in p.text:
        p.text = p.text.replace(':', '：')
    if ';' in p.text:
        p.text = p.text.replace(';', '；')
    if '?' in p.text:
        p.text = p.text.replace('?', '？')
    if ' ' in p.text:
        p.text = p.text.replace(' ', '')
    return p


def fixDocx(docx, version_ipt):
    """ 主要格式 """
    lvl = 0
    for p in docx.paragraphs:
        # print(p.text)
        if p.text == '':
            continue
        else:
            lvl += 1
            # print(p.paragraph_format.first_line_indent)
            # print(p.style.font.size)
            p = replace(p)
            is_level1 = isLevel1(p)
            is_level2 = isLevel2(p)
            if lvl == 1:
                paragraphFun("title", p)
                for run_title in p.runs:
                    # print(run_title.text)
                    string = run_title.text
                    # print(string)
                    run_title._element.getparent().remove(run_title._element)
                    for i in string:
                        num_or_let = isNumberOrLetter(i)
                        text("title", is_level1, is_level2, num_or_let, p, i, version_ipt)
            else:
                paragraphFun("text", p)
                for run_content in p.runs:
                    # print(run_content.text)
                    string = run_content.text
                    # print(string)
                    p.paragraph_format.first_line_indent = 0  # 首行缩进
                    p.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '200')
                    run_content._element.getparent().remove(run_content._element)
                    # 替换格式："1、" --> "1."
                    for i in range(len(string)):
                        if i + 1 <= len(string):
                            if string[i:i+1] in '0123456789' and string[i+1:i+2] == "、":
                                # print(string[i:i+1])
                                string = string[i:i+1] + "." + string[i+2:]
                    for i in string:
                        num_or_let = isNumberOrLetter(i)
                        text("notitle", is_level1, is_level2, num_or_let, p, i, version_ipt)


def isNumberOrLetter(char):
    """ 判断是否为数字或字母 """
    number_and_letter_strs = '0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ'
    if char in number_and_letter_strs:
        return "num_or_let"
    else:
        return False


def pic_fix(docx, file, output_path):
    """ 图片处理 """
    img_path = output_path + "\image"
    file_name = path.splitext(file)[0]
    parts = docx.part.related_parts
    parts_values = parts.values()
    parts_keys = parts.keys()
    list_val = list(parts_values)
    list_key = list(parts_keys)
    parts_length = len(parts_values)
    if parts_length > 5:
        # print(type(list(parts_values)[-1]))
        k = 0
        for i in range(parts_length):
            # print(type(list_val[i]))
            if 'image' in str(type(list_val[i])):
                if not path.isdir(img_path):
                    makedirs(img_path)
                # print('找到图片数据')
                k += 1
                try:
                    img_data = parts[list_key[i]].image.blob
                    img_type = parts[list_key[i]].image.ext
                    full_path = f'{img_path}\{file_name}_image{k}.{img_type}'
                    print(f"··>提示<·· 正在输出：{full_path}")
                    with open(full_path, 'wb') as f:
                        f.write(img_data)
                except:
                    print(f"··>错误<·· 图片{k}输出失败！")
        if k == 0:
            print(f"··>提示<·· 未找到图片！")


def fixWord(docx_path, save_path, file, output_path, time_ipt, page_ipt, img_ipt, version_ipt):
    """ 文档处理 """
    docx = Document(docx_path)

    # 页边距
    margin(docx)

    # 修改格式
    fixDocx(docx, version_ipt)

    # 添加时间后缀
    file_name = path.splitext(file)[0]
    if time_ipt == "1":
        save_time = strftime("%m%d%H%M", localtime())
        save_path = output_path + f"\{file_name}" + save_time + ".docx"
    else:
        save_time = ""
        save_path = output_path + f"\{file_name}" + ".docx"

    # 设置页码
    if page_ipt == "1":
        # 奇偶页不同
        docx.settings.odd_and_even_pages_header_footer = True
        footer(docx)

    # 保存文档中的图片
    if img_ipt == "1":
        pic_fix(docx, file, output_path)

    docx.save(save_path)
    play_history_frm_listbox.insert(END, save_path)
    print(f"··>提示<·· 已保存：{save_path}")
    # 设置滚动条位置到最大值，即拖动到最底部
    play_history_frm_listbox.yview_moveto(1)
    # play_history_frm_listbox.xview_moveto(1)
    return save_time


def inputPath():
    """ 输入路径 """
    input_path = type_radio_value.get()
    if input_path == "file_path":
        file_path = filedialog.askopenfile(title="请选择文件", filetypes=[("docx文件", "*.docx")])
        if file_path != None:
            path_label["text"] = file_path.name
    elif input_path == "dir_path":
        dir_path = filedialog.askdirectory(title="请选择文件夹")
        if dir_path != "":
            path_label["text"] = dir_path


def versionText1():
    """ 版本文本控制1 """
    version_text_label["text"] = VERSIONTEXT1


def versionText2():
    """ 版本文本控制2 """
    version_text_label["text"] = VERSIONTEXT2


def inputFile():
    """ 选择文件 """
    browse_path_button.config(text="选择文件")
    path_label["text"] = ""


def inputDir():
    """ 选择文件夹 """
    browse_path_button.config(text="选择文件夹")
    path_label["text"] = ""


def reSet():
    """ 重置 """
    type_radio1.select()
    inputFile()
    version_radio1.select()
    versionText1()
    time_radio2.select()
    page_radio2.select()
    img_radio2.select()
    play_history_frm_listbox.delete(0, END)


def main():
    """ 主函数 """
    input_path = path_label.cget("text")
    if input_path == "":
        messagebox.showerror("错误", "请选择文件或文件夹路径！")
        inputPath()
    else:
        input_path = input_path.replace("/", "\\")
        version_ipt = version_radio_value.get()
        file_type = type_radio_value.get()
        time_ipt = time_radio_value.get()
        page_ipt = page_radio_value.get()
        img_ipt = img_radio_value.get()
        if file_type == "dir_path":
            output_path = input_path + "\output"
            have_docx = 0
            for file in listdir(input_path):
                if '~' in file:
                    continue
                elif file.endswith('.docx'):
                    if not path.isdir(output_path):
                        makedirs(output_path)
                    have_docx += 1
                    file_path = path.join(input_path, file)
                    fixWord(file_path, output_path + f'\{file}', file, output_path, time_ipt, page_ipt, img_ipt, version_ipt)
            if have_docx == 0:
                print("··>错误<·· 没有找到.docx文件")
                messagebox.showwarning("警告", "没有找到.docx文件！")
            else:
                messagebox.showinfo("提示", "全部处理完成！\n输出路径：" + output_path)
        elif file_type == "file_path":
            # 文件名
            file = input_path.split("\\")[-1]
            # 输出路径
            dir_path = input_path.split("\\")
            dir_path.pop()
            result = '\\'.join(str(x) for x in dir_path)
            output_path = result + "\output"
            if not path.isdir(output_path):
                makedirs(output_path)
            save_time = fixWord(input_path, output_path + f'\{file}', file, output_path, time_ipt, page_ipt, img_ipt, version_ipt)
            messagebox.showinfo("提示", "处理完成！\n输出路径：" + output_path + "\\" + file.split(".")[0] + save_time + ".docx")


if __name__ == '__main__':
    # 配置信息start
    # 字号字典
    FONTSIZEDICT = {
        "八号": 5, "七号": 5.5, "小六号": 6.5, "六号": 7.5, "小五号": 9, "五号": 10.5, "小四号": 12, "四号": 14, "小三号": 15, "三号": 16, "小二号": 18, "二号": 22, "小一号": 24, "一号": 26, "小初号": 36, "初号": 42
    }
    # 版本文本
    VERSIONTEXT1 = """当前配置：
    版    本：学校留存
    页 边 距：上3.7cm 下3.5cm 左2.8cm 右2.6cm
    标    题：小二号 方正小标宋简体  33磅
    正    文：四号 仿宋_GB2312 28磅
    一级标题：四号 黑体 28磅
    二级标题：四号 楷体_GB2312 28磅
    数字&英文：四号 TimesNewRoman字体
    页    码：四号 宋体
"""
    VERSIONTEXT2 = """当前配置：
    版    本：上交上报
    页 边 距：上3.7cm 下3.5cm 左2.8cm 右2.6cm
    标    题：二号 方正小标宋简体  33磅
    正    文：三号 仿宋_GB2312 28磅
    一级标题：三号 黑体 28磅
    二级标题：三号 楷体_GB2312 28磅
    数字&英文：三号 TimesNewRoman字体
    页    码：四号 宋体
"""
    # 页边距-CM
    PAGETOPMARGIN = 3.7  # 页边距：上3.7cm
    PAGEBOTTOMMARGIN = 3.5  # 页边距：下3.5cm
    PAGELEFTMARGIN = 2.8  # 页边距：左2.8cm
    PAGERIGHTMARGIN = 2.6  # 页边距：右2.6cm
    # 标题-磅值
    TITLEFONT = '方正小标宋简体'  # 标题：方正小标宋简体
    TITLEMARGIN = 33  # 标题：固定值33磅
    # 正文-磅值
    TEXTFONT = '仿宋_GB2312'  # 正文：仿宋_GB2312
    TEXTMARGIN = 28  # 正文：一般固定值28磅
    LEVEL1FONT = '黑体'  # 一级标题：黑体
    LEVEL2FONT = '楷体_GB2312'  # 二级标题：楷体_GB2312
    NUMBERFONT = 'Times New Roman'  # 数字&英文：TimesNewRoman字体
    # 页码-磅值
    PAGENUMBERFONTSIZE = FONTSIZEDICT["四号"]  # 页码：四号
    PAGENUMBERFONT = '宋体'  # 页码：宋体
    # 配置信息end
    # tkinter start
    tk = Tk()
    tk.title("文档处理工具-晨小明工作室 v3.0 （学校定制版）")
    screen_width = tk.winfo_screenwidth()
    screen_height = tk.winfo_screenheight()
    """
        !!!!!!!!!!!!
        打包时把此路径改为[\\icon.ico]，并把图标复制粘贴到打包后的根目录里
        !!!!!!!!!!!!
    """
    tk.iconbitmap(getcwd() + "\\test\\fix_word\\icon.ico")
    tk.geometry("800x540")
    # 刷新窗口参数
    tk.update()
    # 计算窗口居中时左上角的坐标
    x = (screen_width - tk.winfo_width()) // 2
    y = (screen_height - tk.winfo_height()) // 2
    tk.geometry(f"+{x}+{y-30}")
    # tk.attributes("-alpha", 0.8)
    windw_width = tk.winfo_width()
    windw_height = tk.winfo_height()
    # 输入类型单选按钮
    type_frm = Frame(tk)
    type_frm.pack()
    type_label = Label(type_frm, font=("Ya Hei", 10), text="请选择输入类型：")
    type_label.grid(row=0, column=0, padx=5, pady=5, sticky="e")
    type_radio_value = StringVar()
    type_radio1 = Radiobutton(type_frm, text="文件", font=("Ya Hei", 10), value="file_path", variable=type_radio_value, command=inputFile)
    type_radio1.grid(row=0, column=1, padx=5, pady=2)
    type_radio2 = Radiobutton(type_frm, text="文件夹", font=("Ya Hei", 10), value="dir_path", variable=type_radio_value, command=inputDir)
    type_radio2.grid(row=0, column=2, padx=5, pady=2)
    type_radio1.select()
    # 文件路径
    path_frm = Frame(tk)
    path_frm.pack()
    path_label = Label(path_frm, width=50, height=1, font=("Ya Hei", 10), border=1, relief="solid")
    path_label.grid(row=1, column=0, padx=5, pady=2, ipadx=5, ipady=5, sticky="w")
    browse_path_button = Button(path_frm, font=("Ya Hei", 10), text="选择文件", command=inputPath)
    browse_path_button.grid(row=1, column=1, padx=5, pady=2)
    # 输入框--后续修复
    # path_entry = Entry(path_frm, width=50, font=("Ya Hei", 10), border=1, relief="solid")
    # path_entry.grid(row=1, column=2, padx=10, pady=10, ipadx=5, ipady=5, sticky="w")
    # 分隔线
    separator = Frame(tk, height=2, bd=1, relief=SUNKEN)
    separator.pack(fill="x", padx=5, pady=5)
    # 选择版本
    version_frm = Frame(tk)
    version_frm.pack(side="top", padx=2, pady=2)
    version_label = Label(version_frm, font=("Ya Hei", 10), text="请选择版本：")
    version_label.grid(row=2, column=0, padx=5, pady=5, sticky="e")
    version_radio_value = StringVar()
    version_radio1 = Radiobutton(version_frm, font=("Ya Hei", 10), text="学校留存", variable=version_radio_value, value="school", command=versionText1)
    version_radio1.grid(row=2, column=1, padx=5, pady=5)
    version_radio2 = Radiobutton(version_frm, font=("Ya Hei", 10), text="上交上报", variable=version_radio_value, value="report", command=versionText2)
    version_radio2.grid(row=2, column=2, padx=5, pady=5)
    version_radio1.select()
    version_text_label = Label(version_frm, width=46, height=12, font=("Ya Hei", 10), text=VERSIONTEXT1, justify="left", border=1, relief="solid")
    version_text_label.grid(row=2, column=3, sticky="w")
    # 分隔线
    separator = Frame(tk, height=2, bd=1, relief=SUNKEN)
    separator.pack(fill="x", padx=5, pady=5)
    # 处理信息
    infos_frm = Frame(tk)
    infos_frm.pack(side="top", padx=2, pady=2)
    info_frm = Frame(infos_frm)
    info_frm.grid(row=0, column=0, padx=5, pady=5, sticky="n")
    time_label = Label(info_frm, font=("Ya Hei", 10), text="添加时间标记：")
    time_label.grid(row=2, column=0, padx=5, pady=5, sticky="e")
    time_radio_value = StringVar()
    time_radio1 = Radiobutton(info_frm, font=("Ya Hei", 10), text="是", variable=time_radio_value, value=True)
    time_radio1.grid(row=2, column=1, padx=5, pady=5)
    time_radio2 = Radiobutton(info_frm, font=("Ya Hei", 10), text="否", variable=time_radio_value, value=False)
    time_radio2.grid(row=2, column=2, padx=5, pady=5)
    time_radio2.select()
    page_label = Label(info_frm, font=("Ya Hei", 10), text="添加页码：")
    page_label.grid(row=3, column=0, padx=5, pady=5, sticky="e")
    page_radio_value = StringVar()
    page_radio1 = Radiobutton(info_frm, font=("Ya Hei", 10), text="是", variable=page_radio_value, value=True)
    page_radio1.grid(row=3, column=1, padx=5, pady=5)
    page_radio2 = Radiobutton(info_frm, font=("Ya Hei", 10), text="否", variable=page_radio_value, value=False)
    page_radio2.grid(row=3, column=2, padx=5, pady=5)
    page_radio2.select()
    img_label = Label(info_frm, font=("Ya Hei", 10), text="保存文档中的图片：")
    img_label.grid(row=4, column=0, padx=5, pady=5, sticky="e")
    img_radio_value = StringVar()
    img_radio1 = Radiobutton(info_frm, font=("Ya Hei", 10), text="是", variable=img_radio_value, value=True)
    img_radio1.grid(row=4, column=1, padx=5, pady=5)
    img_radio2 = Radiobutton(info_frm, font=("Ya Hei", 10), text="否", variable=img_radio_value, value=False)
    img_radio2.grid(row=4, column=2, padx=5, pady=5)
    img_radio2.select()
    # 分隔线
    separator = Frame(tk, height=2, bd=1, relief=SUNKEN)
    separator.pack(fill="x", padx=2, pady=2)
    # 处理日志
    list_frm = Frame(tk)
    list_frm.pack()
    play_history_frm_lbl = Label(list_frm, text="处理日志", font=("Ya Hei", 10, "bold"), fg="green")
    play_history_frm_lbl.grid(row=0, column=0, padx=5, pady=5)
    play_history_frm_listbox = Listbox(list_frm, width=80, height=4, font=("Ya Hei", 10), border=1, activestyle="none")
    play_history_frm_listbox.grid(row=1, column=0, padx=5, pady=5, ipadx=5, ipady=5)
    play_history_scroll_bar_v = ttk.Scrollbar(list_frm, orient="vertical", command=play_history_frm_listbox.yview)
    play_history_scroll_bar_v.grid(row=1, column=1, sticky='ns')
    play_history_scroll_bar_h = ttk.Scrollbar(list_frm, orient="horizontal", command=play_history_frm_listbox.xview)
    play_history_scroll_bar_h.grid(row=2, column=0, sticky='we')
    play_history_frm_listbox.configure(yscrollcommand=play_history_scroll_bar_v.set, xscrollcommand=play_history_scroll_bar_h.set)
    # 分隔线
    separator = Frame(tk, height=2, bd=1, relief=SUNKEN)
    separator.pack(fill="x", padx=2, pady=2)
    # 处理按钮
    btn_frm = Frame(tk)
    btn_frm.pack()
    merge_button = Button(btn_frm, font=("Ya Hei", 10), text="开始处理", command=main)
    merge_button.grid(row=5, column=0, padx=5, pady=5)
    merge_button = Button(btn_frm, font=("Ya Hei", 10), text="重置", fg="blue", command=reSet)
    merge_button.grid(row=5, column=1, padx=5, pady=5)
    tk.mainloop()
    # tkinter end
